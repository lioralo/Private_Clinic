import os
import sqlite3
import socket
import json
import importlib
import sys
import hmac
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass
# OAUTHLIB_INSECURE_TRANSPORT=1 in .env allows OAuth over plain HTTP in local dev.
# Never set this in production — the production .env does not include it.
import hashlib
import threading
import smtplib
from email.message import EmailMessage
from io import BytesIO
import shutil
import secrets
import zipfile
from collections import Counter, defaultdict
from pathlib import Path
from urllib.parse import quote, quote_plus, parse_qs, urlparse
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.interval import IntervalTrigger
from flask import Flask, render_template, request, redirect, url_for, flash, g, send_from_directory, jsonify, session, Response, send_file
from werkzeug.utils import secure_filename
from flask_wtf.csrf import CSRFProtect
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash
import re
import pyotp
from docx import Document
from datetime import datetime, timedelta, timezone
from clinic_app.routes.health import health_bp
from clinic_app.routes.auth import register_auth_routes
from clinic_app.routes.patients import patients_bp
from clinic_app.routes.calendar import calendar_bp
from clinic_app.routes.billing import billing_bp
from clinic_app.routes.google_calendar import google_calendar_bp
from clinic_app.routes.messaging import messaging_bp
from clinic_app.routes.admin import admin_bp
from clinic_app.routes.google_docs import google_docs_bp
from clinic_app.routes.treatment_plans import treatment_plans_bp
from clinic_app.routes.assessments import assessments_bp
from clinic_app.utils import (
    parse_recurrence_days,
    recurring_occurrences_between,
    has_time_conflict,
    ensure_ongoing_recurrence_from_previous_week,
    _ensure_patient_has_upcoming_booking,
    ensure_ongoing_patients_have_upcoming_bookings,
    ensure_default_recurring_vacancies,
    build_booking_management_payload,
    _validate_appointment_duration,
    overlaps,
    redirect_to_patient_tab,
    _check_public_rate_limit,
)
from clinic_app.config import (
    _import_optional_module,
    gcal, gdocs,
    DATABASE, DUMMY_PASSWORD_HASH, BACKUP_DIR, KEY_DIR, BACKUP_INTERVAL_HOURS,
    ALLOWED_UPLOAD_EXTENSIONS, ALLOWED_DIAGNOSIS_EXTENSIONS,
    GDOC_AUTO_SYNC_INTERVAL_SECONDS, GDOC_AUTO_SYNC_GROUP_MODES,
    _GDOC_AUTO_SYNC_LOCK, _GDOC_AUTO_SYNC_LAST_CHECK_TS,
    _GDOC_AUTO_SYNC_WORKER_STATE_LOCK, _GDOC_AUTO_SYNC_WORKER_STARTED,
    _GDOC_AUTO_SYNC_STOP_EVENT,
    _GDOC_MANUAL_SYNC_JOB_LOCK, _GDOC_MANUAL_SYNC_JOBS,
    _GDOC_MANUAL_SYNC_ACTIVE_JOB_ID, _GDOC_MANUAL_SYNC_MAX_JOBS,
    _REMINDER_WORKER_STATE_LOCK, _REMINDER_WORKER_STARTED,
    _REMINDER_WORKER_STOP_EVENT,
    _SECURITY_RETENTION_LOCK, _SECURITY_RETENTION_LAST_CHECK_TS,
    DEFAULT_SITE_SETTINGS, TRANSLATION_OVERRIDES_FILE,
    HEBREW_TRANSLATIONS,
    HEBREW_NUMBER_WORDS, BACKGROUND_REASON_TOPICS, BACKGROUND_THEME_TOPICS,
    LEGACY_WAITING_STATUSES,
)
from clinic_app.backup import (
    _resolve_backup_artifact_sources, _snapshot_artifact_tree,
    _artifact_backup_fingerprint, _write_backup_bundle,
    _is_encrypted_zip_backup, _restore_artifact_tree, _backup_live_artifacts,
    _get_or_create_backup_key, _database_backup_fingerprint,
    perform_encrypted_backup, _export_json_backup,
    _ensure_backup_key_consistency, perform_routine_encrypted_backup,
    list_encrypted_backups, perform_encrypted_restore, _perform_restore,
)


app = Flask(__name__)
app.register_blueprint(health_bp)
app.register_blueprint(patients_bp)
app.register_blueprint(calendar_bp)
app.register_blueprint(billing_bp)
app.register_blueprint(messaging_bp)
app.register_blueprint(google_calendar_bp)
app.register_blueprint(admin_bp)
app.register_blueprint(google_docs_bp)
app.register_blueprint(treatment_plans_bp)
app.register_blueprint(assessments_bp)

# Keep pre-refactor endpoint names stable for tests and backward compatibility.
# Map: old_endpoint_name -> (blueprint_endpoint_name, url_rule, methods)
_legacy_admin_endpoints = [
    ('admin_dashboard', 'admin_dashboard', '/admin/dashboard', ['GET']),
    ('export_calendar', 'export_calendar', '/api/admin/export_calendar', ['GET']),
    ('export_appointments_csv', 'export_appointments_csv', '/api/admin/export_appointments.csv', ['GET']),
    ('bulk_complete_past_appointments', 'bulk_complete_past_appointments', '/api/admin/bulk_complete_past_appointments', ['POST']),
    ('seed_data', 'seed_data', '/admin/seed_data', ['POST']),
    ('reset_test_patients', 'reset_test_patients', '/admin/reset_test_patients', ['POST']),
    ('import_calendar', 'import_calendar', '/api/admin/import_calendar', ['POST']),
    ('admin_global_search', 'admin_global_search', '/admin/search', ['GET']),
    ('list_cancel_requests', 'list_cancel_requests', '/cancel_requests', ['GET']),
    ('approve_cancel_request', 'approve_cancel_request', '/cancel_requests/<int:request_id>/approve', ['POST']),
    ('reject_cancel_request', 'reject_cancel_request', '/cancel_requests/<int:request_id>/reject', ['POST']),
    ('google_docs_auto_sync_now', 'google_docs_auto_sync_now', '/admin/google-docs/auto-sync-now', ['POST']),
    ('google_docs_auto_sync_status', 'google_docs_auto_sync_status', '/admin/google-docs/auto-sync-status/<job_id>', ['GET']),
    ('admin_profile', 'admin_profile', '/admin/profile', ['GET', 'POST']),
    ('admin_smtp_health', 'admin_smtp_health', '/admin/smtp/health', ['GET']),
    ('admin_security_log', 'admin_security_log', '/admin/security-log', ['GET']),
    ('admin_security_log_export', 'admin_security_log_export', '/admin/security-log/export', ['GET']),
    ('admin_smtp_test', 'admin_smtp_test', '/admin/smtp/test', ['POST']),
    ('admin_smtp_send_test', 'admin_smtp_test', '/admin/smtp/test', ['POST']),
    ('setup_authenticator', 'setup_authenticator', '/admin/setup_authenticator', ['POST']),
    ('admin_questionnaire_options', 'admin_questionnaire_options', '/admin/questionnaires/options', ['GET']),
    ('admin_change_password', 'admin_change_password', '/admin/change_password', ['POST']),
    ('backup_now', 'admin_backup_now', '/admin/backup_now', ['POST']),
    ('admin_backup_now', 'admin_backup_now', '/admin/backup_now', ['POST']),
    ('restore_backup_now', 'admin_restore_backup', '/admin/restore_backup', ['POST']),
    ('admin_restore_backup', 'admin_restore_backup', '/admin/restore_backup', ['POST']),
    ('admin_profile_name', 'admin_profile_name', '/admin/profile/name', ['POST']),
    ('manage_resources', 'manage_resources', '/admin/resources', ['GET', 'POST']),
    ('edit_resource', 'edit_resource', '/admin/resources/<int:resource_id>/edit', ['POST']),
    ('delete_resource', 'delete_resource', '/admin/resources/<int:resource_id>/delete', ['POST']),
    ('unassign_resource', 'unassign_resource', '/patient/<int:patient_id>/unassign_resource/<int:resource_id>', ['POST']),
    ('assign_resource', 'assign_resource', '/patient/<int:patient_id>/assign_resource', ['POST']),
    ('admin_users', 'admin_users', '/admin/users', ['GET', 'POST']),
]
for _old_ep, _bp_ep, _rule, _methods in _legacy_admin_endpoints:
    _view = app.view_functions.get(f'admin.{_bp_ep}')
    if _view and _old_ep not in app.view_functions:
        app.add_url_rule(_rule, endpoint=_old_ep, view_func=_view, methods=_methods)
del _legacy_admin_endpoints

_legacy_messaging_aliases = [
    ('api_get_messages', '/api/messages', ['GET']),
    ('api_send_message', '/api/messages/send', ['POST']),
    ('send_message', '/patient/<int:patient_id>/send_message', ['POST']),
    ('admin_reply_message', '/admin_reply_message/<int:patient_id>', ['POST']),
    ('contact_admin', '/contact_admin', ['POST']),
    ('contact_inquiry', '/contact-inquiry', ['POST']),
    ('admin_contact_inquiries', '/admin/contact-inquiries', ['GET']),
    ('mark_contact_inquiry_read', '/admin/contact-inquiries/<int:inquiry_id>/read', ['POST']),
    ('delete_contact_inquiry', '/admin/contact-inquiries/<int:inquiry_id>/delete', ['POST']),
    ('notification_recipients', '/api/notification_recipients', ['GET']),
    ('send_admin_notification', '/admin/notifications/send', ['POST']),
    ('get_notifications', '/api/notifications', ['GET']),
    ('mark_notifications_read', '/api/notifications/mark_read', ['POST']),
]
for name, path, methods in _legacy_messaging_aliases:
    app.add_url_rule(path, endpoint=name, view_func=app.view_functions.get(f'messaging.{name}'), methods=methods)

_legacy_billing_aliases = [
    ('download_receipt', '/patient/receipt/<int:receipt_id>/download', ['GET']),
    ('manage_service_types', '/service_types/manage', ['GET', 'POST']),
    ('toggle_service_type', '/service_types/<int:service_id>/toggle', ['POST']),
    ('add_receipt', '/patient/<int:patient_id>/add_receipt', ['POST']),
    ('set_appointment_status', '/appointment/<int:appointment_id>/set_status', ['POST']),
    ('api_set_appointment_status', '/api/appointment/<int:appointment_id>/status', ['POST']),
]
for name, path, methods in _legacy_billing_aliases:
    app.add_url_rule(path, endpoint=name, view_func=app.view_functions.get(f'billing.{name}'), methods=methods)

# Keep pre-refactor endpoint names stable for tests and backward compatibility.
_legacy_calendar_aliases = [
    ('api_calendar_snapshot', '/api/calendar/snapshot', ['GET']),
    ('api_calendar_block', '/api/calendar/block', ['POST']),
    ('api_calendar_block_update', '/api/calendar/block/<int:block_id>/update', ['POST']),
    ('api_calendar_block_delete', '/api/calendar/block/<int:block_id>/delete', ['POST']),
    ('api_calendar_bookings', '/api/calendar/bookings', ['GET']),
    ('api_calendar_vacancy', '/api/calendar/vacancy', ['POST']),
    ('api_calendar_vacancies', '/api/calendar/vacancies', ['GET']),
    ('api_calendar_vacancy_occupy', '/api/calendar/vacancy/<int:override_id>/occupy', ['POST']),
    ('api_calendar_vacancy_delete', '/api/calendar/vacancy/<int:override_id>/delete', ['POST']),
    ('api_calendar_book', '/api/calendar/book', ['POST']),
    ('weekly_calendar', '/calendar', ['GET']),
    ('api_upcoming_appointments', '/api/appointments/upcoming', ['GET']),
    ('api_create_public_booking_link', '/api/calendar/public-link', ['POST']),
    ('open_public_booking_calendar', '/calendar/public/<token>', ['GET']),
    ('api_public_calendar_book', '/api/calendar/public/<token>/book', ['POST']),
    ('open_booking_page', '/calendar/open/<token>', ['GET']),
    ('api_open_slot_book', '/api/calendar/open/<token>/book', ['POST']),
    ('api_calendar_appointment_delete', '/api/calendar/appointment/<int:appointment_id>/delete', ['POST']),
    ('api_calendar_appointment_update', '/api/calendar/appointment/<int:appointment_id>/update', ['POST']),
    ('add_appointment', '/patient/<int:patient_id>/add_appointment', ['POST']),
    ('export_ics', '/export_ics/<int:appointment_id>', ['GET']),
    ('export_ical', '/appointment/<int:appointment_id>/ical', ['GET']),
    ('delete_appointment', '/appointment/<int:appointment_id>/delete', ['POST']),
]
for _endpoint, _rule, _methods in _legacy_calendar_aliases:
    _view = app.view_functions.get(f'calendar.{_endpoint}')
    if _view and _endpoint not in app.view_functions:
        app.add_url_rule(_rule, endpoint=_endpoint, view_func=_view, methods=_methods)
del _legacy_calendar_aliases

# Keep pre-refactor endpoint names stable for tests and backward compatibility.
_legacy_google_docs_aliases = [
    ('link_gdoc', '/patient/<int:patient_id>/link-gdoc', ['POST']),
    ('attach_gdoc', '/patient/<int:patient_id>/attach-gdoc', ['POST']),
    ('detach_gdoc', '/patient/<int:patient_id>/detach-gdoc', ['POST']),
    ('open_gdoc', '/patient/<int:patient_id>/open-gdoc', ['GET']),
    ('sync_from_gdoc', '/patient/<int:patient_id>/sync-from-gdoc', ['POST']),
    ('link_group_gdoc', '/groups/<int:group_id>/link-gdoc', ['POST']),
    ('attach_group_gdoc', '/groups/<int:group_id>/attach-gdoc', ['POST']),
    ('detach_group_gdoc', '/groups/<int:group_id>/detach-gdoc', ['POST']),
    ('open_group_gdoc', '/groups/<int:group_id>/open-gdoc', ['GET']),
    ('sync_group_gdoc', '/groups/<int:group_id>/sync-gdoc', ['POST']),
    ('pull_group_gdoc', '/groups/<int:group_id>/pull-gdoc', ['POST']),
    ('push_group_gdoc', '/groups/<int:group_id>/push-gdoc', ['POST']),
    ('gdoc_webhook', '/api/gdoc/webhook', ['POST']),
]
for _endpoint, _rule, _methods in _legacy_google_docs_aliases:
    _view = app.view_functions.get(f'google_docs.{_endpoint}')
    if _view and _endpoint not in app.view_functions:
        app.add_url_rule(_rule, endpoint=_endpoint, view_func=_view, methods=_methods)
del _legacy_google_docs_aliases

# Keep pre-refactor endpoint names stable for tests and backward compatibility.
_legacy_google_calendar_aliases = [
    ('google_calendar_status', '/admin/google-calendar/status', ['GET']),
    ('api_google_calendar_status', '/api/google_calendar/status', ['GET']),
    ('google_calendar_connect', '/admin/google-calendar/connect', ['GET', 'POST']),
    ('google_calendar_callback', '/admin/google-calendar/callback', ['GET']),
    ('google_calendar_disconnect', '/admin/google-calendar/disconnect', ['POST']),
    ('google_calendar_set_calendar', '/admin/google-calendar/set-calendar', ['POST']),
]
for _endpoint, _rule, _methods in _legacy_google_calendar_aliases:
    _view = app.view_functions.get(f'google_calendar.{_endpoint}')
    if _view and _endpoint not in app.view_functions:
        app.add_url_rule(_rule, endpoint=_endpoint, view_func=_view, methods=_methods)
del _legacy_google_calendar_aliases

# Keep pre-refactor endpoint names stable for tests and backward compatibility.
_legacy_treatment_plan_aliases = [
    ('view_patient_plans', '/treatment-plans/patient/<int:patient_id>', ['GET']),
    ('create_plan', '/treatment-plans/patient/<int:patient_id>/create', ['GET', 'POST']),
    ('edit_plan', '/treatment-plans/<int:plan_id>/edit', ['GET', 'POST']),
    ('view_plan', '/treatment-plans/<int:plan_id>/view', ['GET']),
    ('delete_plan', '/treatment-plans/<int:plan_id>/delete', ['POST']),
    ('update_goal_progress', '/treatment-plans/api/goal/<int:goal_id>/update-progress', ['POST']),
]
for _endpoint, _rule, _methods in _legacy_treatment_plan_aliases:
    _view = app.view_functions.get(f'treatment_plans.{_endpoint}')
    if _view and _endpoint not in app.view_functions:
        app.add_url_rule(_rule, endpoint=_endpoint, view_func=_view, methods=_methods)
del _legacy_treatment_plan_aliases

_legacy_assessment_aliases = [
    ('view_patient_assessments', '/assessments/patient/<int:patient_id>', ['GET']),
    ('take_assessment', '/assessments/patient/<int:patient_id>/take', ['GET', 'POST']),
    ('assessment_progress_api', '/assessments/api/patient/<int:patient_id>/progress', ['GET']),
    ('delete_assessment', '/assessments/<int:assessment_id>/delete', ['POST']),
]
for _endpoint, _rule, _methods in _legacy_assessment_aliases:
    _view = app.view_functions.get(f'assessments.{_endpoint}')
    if _view and _endpoint not in app.view_functions:
        app.add_url_rule(_rule, endpoint=_endpoint, view_func=_view, methods=_methods)
del _legacy_assessment_aliases

# Re-export moved helpers for test backward compatibility
from clinic_app.routes.google_calendar import (
    _generate_google_oauth_state,
    _is_valid_google_oauth_state,
    _store_google_oauth_pending_state,
    _pop_google_oauth_pending_state,
    _load_active_admin_user,
)

# Re-export moved helpers for test backward compatibility
from clinic_app.routes.calendar import (
    calendar_allowed_windows,
    build_recurrence_group_id,
    canonical_recurrence_days,
    estimate_recurring_series_end,
    find_related_recurring_appointments,
    ensure_recurrence_group_id,
    build_group_recurrence_dates,
    _combine_google_sync_messages,
    _delete_google_events,
    _sync_appointment_with_google,
    _sync_multiple_appointments_with_google,
    _handle_appointment_update_one,
    _handle_appointment_update_upcoming,
    _handle_appointment_update_all,
    _insert_appointment_db,
)

# Re-export moved helpers for test backward compatibility
from clinic_app.routes.google_docs import (
    _extract_google_doc_id,
    _google_docs_dependency_error,
    _extract_google_sheet_id,
    _extract_google_activation_url,
    _friendly_google_sheets_error,
    _google_sheets_dependency_error,
    _get_google_sheets_credentials,
    _list_questionnaire_tabs,
    _list_spreadsheet_tab_titles,
    _create_diagnosee_questionnaires_sheet,
    _copy_questionnaire_tabs_to_spreadsheet,
    _pull_gdoc_notes,
    _pull_group_gdoc_notes,
    _sync_group_gdoc_sessions,
)

app.jinja_env.add_extension('jinja2.ext.do')
app.config['UPLOAD_FOLDER'] = os.environ.get('UPLOAD_FOLDER', 'static/uploads')
app.config['PUBLIC_BASE_URL'] = os.environ.get('PUBLIC_BASE_URL', '').strip()

_secret_key = os.environ.get('SECRET_KEY', '').strip()
if not _secret_key or len(_secret_key) < 32:
    if os.environ.get('FLASK_ENV') == 'production' or not app.debug:
        raise RuntimeError(
            "SECRET_KEY environment variable is not set or is too short. "
            "Set a strong random value of at least 32 characters before starting the app in production."
        )
    import warnings
    warnings.warn("SECRET_KEY is not set or is too short — using an insecure default. Do not use this in production.", stacklevel=2)
    _secret_key = 'dev-insecure-placeholder'
app.secret_key = _secret_key
del _secret_key
app.config['INACTIVITY_TIMEOUT_MINUTES'] = int(os.environ.get('INACTIVITY_TIMEOUT_MINUTES', '5') or 5)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10 MB upload limit
app.config['SESSION_COOKIE_HTTPONLY'] = True
_session_cookie_samesite = (os.environ.get('SESSION_COOKIE_SAMESITE', 'Lax') or 'Lax').strip() or 'Lax'
# OAuth callbacks arrive from accounts.google.com; SameSite=Strict drops the
# session cookie on that cross-site return and breaks the flow.
if _session_cookie_samesite.lower() == 'strict':
    _session_cookie_samesite = 'Lax'
app.config['SESSION_COOKIE_SAMESITE'] = _session_cookie_samesite
app.config['SESSION_COOKIE_SECURE'] = str(os.environ.get('SESSION_COOKIE_SECURE', '0')).strip().lower() in {'1', 'true', 'yes', 'on'}
app.config['LOGIN_RATE_LIMIT_MAX_ATTEMPTS'] = int(os.environ.get('LOGIN_RATE_LIMIT_MAX_ATTEMPTS', '5') or 5)
app.config['LOGIN_RATE_LIMIT_WINDOW_SECONDS'] = int(os.environ.get('LOGIN_RATE_LIMIT_WINDOW_SECONDS', '300') or 300)
app.config['LOGIN_RATE_LIMIT_LOCKOUT_SECONDS'] = int(os.environ.get('LOGIN_RATE_LIMIT_LOCKOUT_SECONDS', '900') or 900)
app.config['PASSWORD_RESET_TOKEN_TTL_SECONDS'] = int(os.environ.get('PASSWORD_RESET_TOKEN_TTL_SECONDS', '1800') or 1800)
app.config['PASSWORD_RESET_RATE_LIMIT_MAX'] = int(os.environ.get('PASSWORD_RESET_RATE_LIMIT_MAX', '5') or 5)
app.config['PASSWORD_RESET_RATE_LIMIT_WINDOW_SECONDS'] = int(os.environ.get('PASSWORD_RESET_RATE_LIMIT_WINDOW_SECONDS', '900') or 900)
app.config['SMTP_HOST'] = (os.environ.get('SMTP_HOST') or '').strip()
app.config['SMTP_PORT'] = int(os.environ.get('SMTP_PORT', '587') or 587)
app.config['SMTP_USERNAME'] = (os.environ.get('SMTP_USERNAME') or '').strip()
app.config['SMTP_PASSWORD'] = os.environ.get('SMTP_PASSWORD', '')
app.config['SMTP_FROM_EMAIL'] = (os.environ.get('SMTP_FROM_EMAIL') or '').strip()
app.config['SMTP_USE_TLS'] = str(os.environ.get('SMTP_USE_TLS', '1')).strip().lower() in {'1', 'true', 'yes', 'on'}
app.config['SECURITY_RETENTION_CHECK_INTERVAL_SECONDS'] = int(os.environ.get('SECURITY_RETENTION_CHECK_INTERVAL_SECONDS', '3600') or 3600)
app.config['AUDIT_LOG_RETENTION_DAYS'] = int(os.environ.get('AUDIT_LOG_RETENTION_DAYS', '365') or 365)
app.config['PASSWORD_RESET_TOKEN_RETENTION_DAYS'] = int(os.environ.get('PASSWORD_RESET_TOKEN_RETENTION_DAYS', '30') or 30)
app.config['PUBLIC_BOOKING_RATE_LIMIT_MAX'] = int(os.environ.get('PUBLIC_BOOKING_RATE_LIMIT_MAX', '20') or 20)
app.config['PUBLIC_BOOKING_RATE_LIMIT_WINDOW_SECONDS'] = int(os.environ.get('PUBLIC_BOOKING_RATE_LIMIT_WINDOW_SECONDS', '60') or 60)
app.config['REMINDER_HOURS_BEFORE'] = int(os.environ.get('REMINDER_HOURS_BEFORE', '24') or 24)
app.config['REMINDER_SCHEDULER_INTERVAL'] = int(os.environ.get('REMINDER_SCHEDULER_INTERVAL', '300') or 300)
app.config['TWILIO_ACCOUNT_SID'] = (os.environ.get('TWILIO_ACCOUNT_SID') or '').strip()
app.config['TWILIO_AUTH_TOKEN'] = os.environ.get('TWILIO_AUTH_TOKEN', '')
app.config['TWILIO_FROM_NUMBER'] = (os.environ.get('TWILIO_FROM_NUMBER') or '').strip()
scheduler = BackgroundScheduler(daemon=True)
csrf = CSRFProtect(app)
app.config['DATABASE'] = DATABASE

def _allowed_upload(filename, allowed_set):
    from clinic_app.utils import _allowed_upload as _impl
    return _impl(filename, allowed_set)

def ensure_runtime_paths():
    db_path = Path(app.config.get('DATABASE', DATABASE))
    if db_path.parent != Path('.'):
        db_path.parent.mkdir(parents=True, exist_ok=True)

    upload_path = Path(app.config.get('UPLOAD_FOLDER', 'static/uploads'))
    upload_path.mkdir(parents=True, exist_ok=True)

    backup_path = Path(BACKUP_DIR)
    backup_path.mkdir(parents=True, exist_ok=True)


import functools
import time as _time

def _db_retry(max_attempts=3, delay=0.1):
    def decorator(f):
        @functools.wraps(f)
        def wrapper(*args, **kwargs):
            last_exc = None
            for attempt in range(max_attempts):
                try:
                    return f(*args, **kwargs)
                except sqlite3.OperationalError as exc:
                    if 'database is locked' not in str(exc):
                        raise
                    last_exc = exc
                    if attempt < max_attempts - 1:
                        _time.sleep(delay * (2 ** attempt))
            raise last_exc
        return wrapper
    return decorator


def _request_client_ip():
    from clinic_app.utils import _request_client_ip as _impl
    return _impl()


def _smtp_settings_summary():
    from clinic_app.utils import _smtp_settings_summary as _impl
    return _impl(app)


def _send_smtp_email(recipient_email, subject, body_text, html_body=None):
    from clinic_app.utils import _send_smtp_email as _impl
    return _impl(recipient_email, subject, body_text, app=app)


def _validate_patient_fields(name, phone=None, birth_date=None, email=None):
    from clinic_app.utils import _validate_patient_fields as _impl
    return _impl(name, phone=phone, birth_date=birth_date, email=email)



def _validate_password_strength(password, username=None, email=None):
    from clinic_app.utils import _validate_password_strength as _impl
    return _impl(password, username=username, email=email)

    lowered = candidate.lower()
    for source in (username or '', email or ''):
        token = str(source).strip().lower()
        if token and len(token) >= 3 and token in lowered:
            return False, 'Password should not include your username or email.'

    return True, ''


def _security_retention_cleanup(db):
    global _SECURITY_RETENTION_LAST_CHECK_TS

    if app.config.get('TESTING'):
        return

    now_ts = datetime.now(timezone.utc).timestamp()
    interval_seconds = int(app.config.get('SECURITY_RETENTION_CHECK_INTERVAL_SECONDS', 3600) or 3600)
    if now_ts - float(_SECURITY_RETENTION_LAST_CHECK_TS or 0.0) < max(60, interval_seconds):
        return

    if not _SECURITY_RETENTION_LOCK.acquire(blocking=False):
        return

    try:
        audit_days = int(app.config.get('AUDIT_LOG_RETENTION_DAYS', 365) or 365)
        reset_days = int(app.config.get('PASSWORD_RESET_TOKEN_RETENTION_DAYS', 30) or 30)
        candidate_days = int(app.config.get('CANDIDATE_PURGE_DAYS', 0) or 0)

        db.execute(
            "DELETE FROM audit_logs WHERE created_at < datetime('now', ?)",
            (f'-{max(1, audit_days)} days',),
        )
        db.execute(
            """
            DELETE FROM password_reset_tokens
            WHERE expires_at < datetime('now', ?)
               OR (used_at IS NOT NULL AND used_at < datetime('now', ?))
            """,
            (f'-{max(1, reset_days)} days', f'-{max(1, reset_days)} days'),
        )
        # Purge stale candidate patients (opt-in — only when CANDIDATE_PURGE_DAYS > 0)
        if candidate_days > 0:
            db.execute(
                """
                DELETE FROM patients
                WHERE status = 'candidate'
                  AND created_at IS NOT NULL
                  AND created_at < datetime('now', ?)
                """,
                (f'-{candidate_days} days',),
            )
        db.commit()
        _SECURITY_RETENTION_LAST_CHECK_TS = now_ts
    except Exception:
        app.logger.exception('Security retention cleanup failed')
    finally:
        _SECURITY_RETENTION_LOCK.release()



def _validate_gdoc_webhook_request():
    channel_id = (request.headers.get('X-Goog-Channel-ID') or '').strip()
    resource_state = (request.headers.get('X-Goog-Resource-State') or '').strip()
    if not channel_id:
        return False
    if not resource_state:
        return False

    configured_secret = (
        app.config.get('GOOGLE_DOCS_WEBHOOK_SECRET')
        or os.environ.get('GOOGLE_DOCS_WEBHOOK_SECRET')
        or ''
    ).strip()
    if configured_secret:
        provided_secret = (request.headers.get('X-Webhook-Secret') or '').strip()
        if not hmac.compare_digest(provided_secret, configured_secret):
            return False

    return True


@app.template_filter('rjust')
def rjust_filter(s, width, fillchar=' '):
    return str(s).rjust(width, fillchar)

@app.template_filter('from_iso_date')
def from_iso_date(value):
    try:
        if value is None:
            return value
        return datetime.fromisoformat(value).date()
    except (ValueError, TypeError, AttributeError):
        return value

@app.template_filter('from_iso_datetime')
def from_iso_datetime(value):
    try:
        if value is None:
            return value
        return datetime.fromisoformat(value.replace('Z', '+00:00'))
    except (ValueError, TypeError, AttributeError):
        return value

@app.template_filter('strftime')
def strftime_filter(value, format_string):
    try:
        return value.strftime(format_string)
    except AttributeError:
        return value

@app.template_filter('date')
def date_filter(value):
    try:
        return value.date()
    except AttributeError:
        return value

@app.template_filter('israeli_date')
def israeli_date_filter(value):
    try:
        return value.strftime('%d.%m.%Y')
    except AttributeError:
        return value

@app.template_filter('from_json')
def from_json_filter(value):
    if value is None:
        return []
    try:
        if isinstance(value, str):
            return json.loads(value)
        return value
    except (json.JSONDecodeError, TypeError):
        return []

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

def load_hebrew_translation_overrides():
    overrides_file = Path(__file__).resolve().parent / 'translations' / 'he.json'
    if not overrides_file.exists():
        return {}
    try:
        payload = json.loads(overrides_file.read_text(encoding='utf-8'))
        if not isinstance(payload, dict):
            return {}
        return {
            str(k): str(v)
            for k, v in payload.items()
            if isinstance(k, str) and isinstance(v, str)
        }
    except Exception:
        app.logger.exception('Failed to load Hebrew translation overrides')
        return {}


HEBREW_TRANSLATIONS.update(load_hebrew_translation_overrides())



def get_site_settings(db=None):
    settings = DEFAULT_SITE_SETTINGS.copy()
    try:
        active_db = db or get_db()
        rows = active_db.execute('SELECT setting_key, setting_value FROM site_settings').fetchall()
        for row in rows:
            key = row['setting_key']
            if key in settings:
                settings[key] = row['setting_value'] or ''
    except Exception:
        app.logger.exception('get_site_settings failed')
        return settings

    settings['about_enabled'] = '1' if str(settings.get('about_enabled') or '0') in {'1', 'true', 'yes', 'on'} else '0'
    return settings


def _build_about_map_urls(raw_url):
    map_url = (raw_url or '').strip()
    if not map_url:
        return {'open_url': '', 'embed_url': ''}

    if not re.match(r'^https?://', map_url, flags=re.IGNORECASE):
        map_url = 'https://' + map_url.lstrip('/')

    open_url = map_url
    parsed = urlparse(map_url)
    lower_url = map_url.lower()

    if '/maps/embed' in lower_url or 'output=embed' in lower_url:
        return {'open_url': open_url, 'embed_url': map_url}

    query = parse_qs(parsed.query or '', keep_blank_values=True)
    location_hint = ''
    for key in ('q', 'query', 'destination'):
        values = query.get(key) or []
        if values and str(values[0]).strip():
            location_hint = str(values[0]).strip()
            break

    path = parsed.path or ''
    if not location_hint and '/maps/place/' in path:
        location_hint = path.split('/maps/place/', 1)[1].split('/', 1)[0].replace('+', ' ').strip()

    if location_hint:
        embed_url = f"https://www.google.com/maps?q={quote_plus(location_hint)}&output=embed"
        return {'open_url': open_url, 'embed_url': embed_url}

    return {'open_url': open_url, 'embed_url': ''}


def save_site_settings(db, updates):
    if not updates:
        return
    for key, default_value in DEFAULT_SITE_SETTINGS.items():
        if key not in updates:
            continue
        value = updates.get(key)
        if key == 'about_enabled':
            value = '1' if str(value or '0') in {'1', 'true', 'yes', 'on'} else '0'
        elif value is None:
            value = default_value
        db.execute(
            'INSERT INTO site_settings (setting_key, setting_value) VALUES (?, ?) '
            'ON CONFLICT(setting_key) DO UPDATE SET setting_value=excluded.setting_value',
            (key, str(value))
        )


def _format_gdoc_target_key(target_type, target_id):
    normalized_type = (target_type or '').strip().lower()
    try:
        normalized_id = int(target_id)
    except (TypeError, ValueError):
        return None
    if normalized_type not in {'patient', 'group'} or normalized_id <= 0:
        return None
    return f'{normalized_type}:{normalized_id}'


def _parse_gdoc_target_key(raw_value):
    text = (raw_value or '').strip().lower()
    match = re.fullmatch(r'(patient|group):(\d+)', text)
    if not match:
        return None, None
    return match.group(1), int(match.group(2))


def _safe_parse_gdoc_targets_json(raw_json):
    try:
        parsed = json.loads(raw_json or '[]')
    except (TypeError, ValueError, json.JSONDecodeError):
        return []

    if not isinstance(parsed, list):
        return []

    normalized = []
    seen = set()
    for item in parsed:
        target_type, target_id = _parse_gdoc_target_key(str(item))
        if not target_type:
            continue
        key = f'{target_type}:{target_id}'
        if key in seen:
            continue
        seen.add(key)
        normalized.append(key)
    return normalized


def _safe_parse_gdoc_targets_config_json(raw_json):
    try:
        parsed = json.loads(raw_json or '[]')
    except (TypeError, ValueError, json.JSONDecodeError):
        return []

    if not isinstance(parsed, list):
        return []

    normalized = []
    seen = set()
    for item in parsed:
        if not isinstance(item, dict):
            continue
        target_type, target_id = _parse_gdoc_target_key(item.get('target_key'))
        if not target_type:
            continue
        mode = str(item.get('mode') or 'pull').strip().lower()
        if target_type == 'group':
            mode = mode if mode in GDOC_AUTO_SYNC_GROUP_MODES else 'pull'
        else:
            mode = 'pull'
        target_key = f'{target_type}:{target_id}'
        if target_key in seen:
            continue
        seen.add(target_key)
        normalized.append({'target_key': target_key, 'mode': mode})
    return normalized


def _list_connected_google_docs(db):
    docs = []

    try:
        patient_rows = db.execute('''
            SELECT id, name, gdoc_id
            FROM patients
            WHERE COALESCE(gdoc_id, '') <> ''
              AND COALESCE(is_deleted, 0) = 0
            ORDER BY name COLLATE NOCASE ASC, id ASC
        ''').fetchall()
    except sqlite3.OperationalError:
        patient_rows = []
    for row in patient_rows:
        target_key = _format_gdoc_target_key('patient', row['id'])
        if not target_key:
            continue
        docs.append({
            'target_key': target_key,
            'target_type': 'patient',
            'target_id': int(row['id']),
            'label': row['name'] or f"Patient #{row['id']}",
            'doc_id': row['gdoc_id'],
            'doc_url': f"https://docs.google.com/document/d/{row['gdoc_id']}/edit",
        })

    try:
        group_rows = db.execute('''
            SELECT id, name, gdoc_id
            FROM groups
            WHERE COALESCE(gdoc_id, '') <> ''
            ORDER BY name COLLATE NOCASE ASC, id ASC
        ''').fetchall()
    except sqlite3.OperationalError:
        group_rows = []
    for row in group_rows:
        target_key = _format_gdoc_target_key('group', row['id'])
        if not target_key:
            continue
        docs.append({
            'target_key': target_key,
            'target_type': 'group',
            'target_id': int(row['id']),
            'label': row['name'] or f"Group #{row['id']}",
            'doc_id': row['gdoc_id'],
            'doc_url': f"https://docs.google.com/document/d/{row['gdoc_id']}/edit",
        })

    docs.sort(key=lambda item: (item['target_type'], (item['label'] or '').lower(), item['target_id']))
    return docs


def _get_google_docs_auto_sync_state(db, connected_docs=None):
    settings = get_site_settings(db)
    interval_key = (settings.get('gdocs_auto_sync_interval') or 'daily').strip().lower()
    if interval_key not in GDOC_AUTO_SYNC_INTERVAL_SECONDS:
        interval_key = 'daily'

    connected_docs = connected_docs if connected_docs is not None else _list_connected_google_docs(db)
    connected_by_key = {item['target_key']: item for item in connected_docs}

    selected_config = _safe_parse_gdoc_targets_config_json(settings.get('gdocs_auto_sync_targets_config_json'))
    if not selected_config:
        selected_config = [{'target_key': key, 'mode': 'pull'} for key in _safe_parse_gdoc_targets_json(settings.get('gdocs_auto_sync_targets_json'))]

    selected_targets = []
    for item in selected_config:
        target_key = item.get('target_key')
        if target_key not in connected_by_key:
            continue
        target_type, target_id = _parse_gdoc_target_key(target_key)
        if not target_type:
            continue
        mode = str(item.get('mode') or 'pull').strip().lower()
        if target_type == 'group':
            mode = mode if mode in GDOC_AUTO_SYNC_GROUP_MODES else 'pull'
        else:
            mode = 'pull'
        target_doc = connected_by_key[target_key]
        selected_targets.append({
            'target_key': target_key,
            'target_type': target_type,
            'target_id': target_id,
            'mode': mode,
            'label': target_doc.get('label') or target_key,
        })
    last_run_at_raw = (settings.get('gdocs_auto_sync_last_run_at') or '').strip()
    last_run_at = None
    if last_run_at_raw:
        try:
            last_run_at = datetime.fromisoformat(last_run_at_raw)
        except ValueError:
            last_run_at = None

    return {
        'enabled': str(settings.get('gdocs_auto_sync_enabled') or '0') in {'1', 'true', 'yes', 'on'},
        'interval_key': interval_key,
        'selected_target_keys': [item['target_key'] for item in selected_targets],
        'selected_targets': selected_targets,
        'last_run_at': last_run_at,
        'last_run_at_raw': last_run_at_raw,
        'connected_docs': connected_docs,
    }


def _is_transient_sync_error(error_text):
    text = str(error_text or '').lower()
    transient_signals = (
        'timeout',
        'timed out',
        'temporarily unavailable',
        'rate limit',
        'too many requests',
        'connection reset',
        'connection aborted',
        'connection refused',
        'database is locked',
        'database locked',
        '503',
        '502',
        '500',
        '429',
    )
    return any(signal in text for signal in transient_signals)


def _run_sync_with_retry(sync_callable, max_attempts=5, base_delay_seconds=2.0):
    attempt = 0
    while attempt < max_attempts:
        attempt += 1
        count, sync_err = sync_callable()
        if not sync_err:
            return int(count or 0), None, attempt
        if attempt >= max_attempts or not _is_transient_sync_error(sync_err):
            return int(count or 0), str(sync_err), attempt
        _time.sleep(base_delay_seconds * (2 ** (attempt - 1)))
    return 0, 'Sync failed after retries', max_attempts


def _record_gdocs_sync_history(db, trigger_source, status, interval_key, selected_targets, processed_targets,
                               synced_total, synced_patients, synced_groups, pushed_groups, errors, details):
    try:
        cursor = db.execute(
            '''INSERT INTO gdocs_sync_history (
                   trigger_source, status, interval_key,
                   targets_total, targets_processed,
                   synced_total, synced_patients, synced_groups, pushed_groups,
                   errors_json, details_json
               ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
            (
                trigger_source,
                status,
                interval_key,
                len(selected_targets),
                len(processed_targets),
                int(synced_total or 0),
                int(synced_patients or 0),
                int(synced_groups or 0),
                int(pushed_groups or 0),
                json.dumps(errors or []),
                json.dumps(details or []),
            )
        )
        return int(cursor.lastrowid)
    except sqlite3.OperationalError:
        return None


def _run_google_docs_auto_sync(db, force=False, trigger_source='auto', progress_callback=None):
    def emit_progress(**payload):
        if not callable(progress_callback):
            return
        try:
            progress_callback(payload)
        except Exception:
            app.logger.exception('progress_callback failed')

    state = _get_google_docs_auto_sync_state(db)
    now = datetime.now(timezone.utc)
    should_run = bool(force)

    if not should_run:
        if not state['enabled']:
            return {'ran': False, 'reason': 'disabled'}
        if not state['selected_targets']:
            return {'ran': False, 'reason': 'no-targets'}

        interval_seconds = GDOC_AUTO_SYNC_INTERVAL_SECONDS.get(state['interval_key'], GDOC_AUTO_SYNC_INTERVAL_SECONDS['daily'])
        if state['last_run_at'] is None:
            should_run = True
        else:
            elapsed_seconds = (now - state['last_run_at']).total_seconds()
            should_run = elapsed_seconds >= interval_seconds

    if not should_run:
        emit_progress(status='skipped', message='Sync is not due yet.', percent=100)
        return {'ran': False, 'reason': 'not-due'}

    selected_targets = list(state['selected_targets'])
    if not selected_targets and force:
        connected_docs = state.get('connected_docs') or []
        for doc in connected_docs:
            target_type, target_id = _parse_gdoc_target_key(doc.get('target_key'))
            if not target_type:
                continue
            selected_targets.append({
                'target_key': doc.get('target_key'),
                'target_type': target_type,
                'target_id': target_id,
                'mode': 'pull' if target_type == 'patient' else 'pull',
                'label': doc.get('label') or doc.get('target_key'),
            })

    if not selected_targets:
        emit_progress(status='failed', message='No connected Google Docs are selected for sync.', percent=100)
        return {'ran': False, 'reason': 'no-connected-targets'}

    targets_total = len(selected_targets)
    emit_progress(
        status='running',
        message='Starting Google Docs sync...',
        phase='prepare',
        targets_total=targets_total,
        targets_processed=0,
        percent=0,
    )

    dependency_error = _google_docs_dependency_error()
    if dependency_error:
        history_id = _record_gdocs_sync_history(
            db,
            trigger_source=trigger_source,
            status='failed',
            interval_key=state['interval_key'],
            selected_targets=selected_targets,
            processed_targets=[],
            synced_total=0,
            synced_patients=0,
            synced_groups=0,
            pushed_groups=0,
            errors=[dependency_error],
            details=[],
        )
        db.commit()
        emit_progress(
            status='failed',
            message='Google dependencies are unavailable.',
            phase='failed',
            targets_total=targets_total,
            targets_processed=0,
            percent=100,
        )
        return {'ran': False, 'reason': 'dependency', 'errors': [dependency_error], 'history_id': history_id}

    total_synced = 0
    synced_patients = 0
    synced_groups = 0
    pushed_groups = 0
    errors = []
    warnings = []
    processed = []
    details = []
    processed_steps = 0

    for target in selected_targets:
        target_key = target['target_key']
        target_type = target['target_type']
        target_id = target['target_id']
        target_mode = target.get('mode') or 'pull'
        target_label = target.get('label') or target_key

        emit_progress(
            status='running',
            phase='syncing',
            message=f'Syncing {target_label}...',
            current_target=target_label,
            current_target_key=target_key,
            targets_total=targets_total,
            targets_processed=processed_steps,
            percent=int((processed_steps / targets_total) * 100) if targets_total else 0,
        )

        if target_type == 'patient':
            patient = db.execute('SELECT * FROM patients WHERE id = ? AND COALESCE(gdoc_id, "") <> ""', (target_id,)).fetchone()
            if not patient:
                processed_steps += 1
                emit_progress(
                    status='running',
                    phase='syncing',
                    message=f'Skipped {target_label} (not connected).',
                    current_target=target_label,
                    current_target_key=target_key,
                    targets_total=targets_total,
                    targets_processed=processed_steps,
                    percent=int((processed_steps / targets_total) * 100) if targets_total else 100,
                )
                continue
            synced_count, sync_err, attempts = _run_sync_with_retry(lambda: _pull_gdoc_notes(db, patient))
            if sync_err:
                errors.append(f"patient:{target_id} -> {sync_err}")
                details.append({'target_key': target_key, 'mode': 'pull', 'action': 'pull', 'status': 'error', 'attempts': attempts, 'error': sync_err})
                processed_steps += 1
                emit_progress(
                    status='running',
                    phase='syncing',
                    message=f'Finished {target_label} with errors.',
                    current_target=target_label,
                    current_target_key=target_key,
                    targets_total=targets_total,
                    targets_processed=processed_steps,
                    percent=int((processed_steps / targets_total) * 100) if targets_total else 100,
                )
                continue

            total_synced += int(synced_count or 0)
            synced_patients += 1
            processed.append(target_key)
            details.append({'target_key': target_key, 'mode': 'pull', 'action': 'pull', 'status': 'ok', 'attempts': attempts, 'synced': int(synced_count or 0)})
            processed_steps += 1
            emit_progress(
                status='running',
                phase='syncing',
                message=f'Finished {target_label}.',
                current_target=target_label,
                current_target_key=target_key,
                targets_total=targets_total,
                targets_processed=processed_steps,
                percent=int((processed_steps / targets_total) * 100) if targets_total else 100,
            )
            continue

        group = db.execute('SELECT * FROM groups WHERE id = ? AND COALESCE(gdoc_id, "") <> ""', (target_id,)).fetchone()
        if not group:
            processed_steps += 1
            emit_progress(
                status='running',
                phase='syncing',
                message=f'Skipped {target_label} (not connected).',
                current_target=target_label,
                current_target_key=target_key,
                targets_total=targets_total,
                targets_processed=processed_steps,
                percent=int((processed_steps / targets_total) * 100) if targets_total else 100,
            )
            continue

        pulled_count, pull_err, pull_attempts = _run_sync_with_retry(lambda: _pull_group_gdoc_notes(db, group))
        if pull_err:
            errors.append(f"group:{target_id} pull -> {pull_err}")
            details.append({'target_key': target_key, 'mode': target_mode, 'action': 'pull', 'status': 'error', 'attempts': pull_attempts, 'error': pull_err})
            processed_steps += 1
            emit_progress(
                status='running',
                phase='syncing',
                message=f'Finished {target_label} with pull errors.',
                current_target=target_label,
                current_target_key=target_key,
                targets_total=targets_total,
                targets_processed=processed_steps,
                percent=int((processed_steps / targets_total) * 100) if targets_total else 100,
            )
            continue

        total_synced += int(pulled_count or 0)
        synced_groups += 1
        processed.append(target_key)
        details.append({'target_key': target_key, 'mode': target_mode, 'action': 'pull', 'status': 'ok', 'attempts': pull_attempts, 'synced': int(pulled_count or 0)})

        if target_mode == 'both':
            pushed_count, push_err, push_attempts = _run_sync_with_retry(lambda: _sync_group_gdoc_sessions(db, group))
            if push_err:
                warnings.append(f"group:{target_id} push skipped after errors: {push_err}")
                details.append({'target_key': target_key, 'mode': target_mode, 'action': 'push', 'status': 'error', 'attempts': push_attempts, 'error': push_err})
            else:
                total_synced += int(pushed_count or 0)
                pushed_groups += 1
                details.append({'target_key': target_key, 'mode': target_mode, 'action': 'push', 'status': 'ok', 'attempts': push_attempts, 'synced': int(pushed_count or 0)})

        processed_steps += 1
        emit_progress(
            status='running',
            phase='syncing',
            message=f'Finished {target_label}.',
            current_target=target_label,
            current_target_key=target_key,
            targets_total=targets_total,
            targets_processed=processed_steps,
            percent=int((processed_steps / targets_total) * 100) if targets_total else 100,
        )

    save_site_settings(db, {
        'gdocs_auto_sync_last_run_at': now.isoformat(),
    })

    if errors:
        run_status = 'partial' if processed else 'failed'
    elif warnings:
        run_status = 'partial'
    elif selected_targets and not processed:
        # All selected targets were skipped (e.g. gdoc_id removed between state load and sync)
        run_status = 'partial'
    else:
        run_status = 'success'

    history_id = _record_gdocs_sync_history(
        db,
        trigger_source=trigger_source,
        status=run_status,
        interval_key=state['interval_key'],
        selected_targets=selected_targets,
        processed_targets=processed,
        synced_total=total_synced,
        synced_patients=synced_patients,
        synced_groups=synced_groups,
        pushed_groups=pushed_groups,
        errors=errors,
        details=details,
    )
    db.commit()

    emit_progress(
        status=run_status,
        phase='done',
        message='Google Docs sync completed.',
        targets_total=targets_total,
        targets_processed=processed_steps,
        percent=100,
    )

    return {
        'ran': True,
        'status': run_status,
        'processed_targets': processed,
        'targets_total': targets_total,
        'targets_processed': processed_steps,
        'total_synced': total_synced,
        'synced_patients': synced_patients,
        'synced_groups': synced_groups,
        'pushed_groups': pushed_groups,
        'errors': errors,
        'warnings': warnings,
        'history_id': history_id,
    }


def _get_recent_gdocs_sync_history(db, limit=12):
    try:
        rows = db.execute('''
            SELECT *
            FROM gdocs_sync_history
            ORDER BY run_at DESC, id DESC
            LIMIT ?
        ''', (int(limit),)).fetchall()
    except sqlite3.OperationalError:
        return []

    history = []
    for row in rows:
        try:
            errors = json.loads(row['errors_json'] or '[]')
        except (TypeError, ValueError, json.JSONDecodeError):
            errors = []
        history.append({
            'id': int(row['id']),
            'run_at': row['run_at'],
            'trigger_source': row['trigger_source'] or 'auto',
            'status': row['status'] or 'unknown',
            'targets_total': int(row['targets_total'] or 0),
            'targets_processed': int(row['targets_processed'] or 0),
            'synced_total': int(row['synced_total'] or 0),
            'synced_patients': int(row['synced_patients'] or 0),
            'synced_groups': int(row['synced_groups'] or 0),
            'pushed_groups': int(row['pushed_groups'] or 0),
            'errors': errors,
        })
    return history


def _get_gdocs_auto_sync_health(db):
    state = _get_google_docs_auto_sync_state(db)
    interval_seconds = GDOC_AUTO_SYNC_INTERVAL_SECONDS.get(state['interval_key'], GDOC_AUTO_SYNC_INTERVAL_SECONDS['daily'])
    now = datetime.now(timezone.utc)
    last_run = state['last_run_at']
    next_run = (last_run + timedelta(seconds=interval_seconds)) if last_run else None
    overdue = bool(state['enabled'] and next_run and now > next_run)
    age_seconds = int((now - last_run).total_seconds()) if last_run else None

    recent_history = _get_recent_gdocs_sync_history(db, limit=1)
    last_record = recent_history[0] if recent_history else None
    last_status = last_record['status'] if last_record else None
    last_synced_total = last_record['synced_total'] if last_record else None

    return {
        'enabled': state['enabled'],
        'interval_key': state['interval_key'],
        'selected_count': len(state['selected_targets']),
        'last_run_at': last_run,
        'next_run_at': next_run,
        'last_run_age_seconds': age_seconds,
        'overdue': overdue,
        'last_status': last_status,
        'last_synced_total': last_synced_total,
    }


def _cleanup_manual_sync_jobs_locked():
    if len(_GDOC_MANUAL_SYNC_JOBS) <= _GDOC_MANUAL_SYNC_MAX_JOBS:
        return
    ordered_ids = sorted(
        _GDOC_MANUAL_SYNC_JOBS.keys(),
        key=lambda key: (_GDOC_MANUAL_SYNC_JOBS[key].get('started_at') or '', key)
    )
    removable = len(_GDOC_MANUAL_SYNC_JOBS) - _GDOC_MANUAL_SYNC_MAX_JOBS
    for job_id in ordered_ids:
        if removable <= 0:
            break
        if job_id == _GDOC_MANUAL_SYNC_ACTIVE_JOB_ID:
            continue
        _GDOC_MANUAL_SYNC_JOBS.pop(job_id, None)
        removable -= 1


def _create_manual_sync_job(admin_user_id):
    now_iso = datetime.now(timezone.utc).isoformat()
    job_id = secrets.token_hex(16)
    with _GDOC_MANUAL_SYNC_JOB_LOCK:
        global _GDOC_MANUAL_SYNC_ACTIVE_JOB_ID
        existing_id = _GDOC_MANUAL_SYNC_ACTIVE_JOB_ID
        existing_job = _GDOC_MANUAL_SYNC_JOBS.get(existing_id) if existing_id else None
        if existing_job and existing_job.get('status') == 'running':
            return None, existing_id

        _GDOC_MANUAL_SYNC_ACTIVE_JOB_ID = job_id
        _GDOC_MANUAL_SYNC_JOBS[job_id] = {
            'job_id': job_id,
            'status': 'running',
            'created_by': int(admin_user_id),
            'started_at': now_iso,
            'finished_at': None,
            'message': 'Starting Google Docs sync...',
            'phase': 'prepare',
            'targets_total': 0,
            'targets_processed': 0,
            'percent': 0,
            'current_target': None,
            'result': None,
            'error': None,
        }
        _cleanup_manual_sync_jobs_locked()
    return job_id, None


def _update_manual_sync_job(job_id, payload):
    with _GDOC_MANUAL_SYNC_JOB_LOCK:
        job = _GDOC_MANUAL_SYNC_JOBS.get(job_id)
        if not job:
            return
        for key in ('message', 'phase', 'current_target'):
            if key in payload:
                job[key] = payload.get(key)

        if 'targets_total' in payload:
            job['targets_total'] = int(payload.get('targets_total') or 0)
        if 'targets_processed' in payload:
            job['targets_processed'] = int(payload.get('targets_processed') or 0)

        total = max(0, int(job.get('targets_total') or 0))
        processed = max(0, int(job.get('targets_processed') or 0))
        if total > 0:
            computed_percent = int(round((processed / total) * 100))
        else:
            computed_percent = int(payload.get('percent') or 0)
        computed_percent = max(0, min(100, computed_percent))
        job['percent'] = int(payload.get('percent', computed_percent)) if payload.get('percent') is not None else computed_percent

        status = payload.get('status')
        if status in {'running', 'success', 'partial', 'failed', 'error', 'skipped'}:
            job['status'] = status


def _complete_manual_sync_job(job_id, status, result=None, error_message=None):
    now_iso = datetime.now(timezone.utc).isoformat()
    with _GDOC_MANUAL_SYNC_JOB_LOCK:
        global _GDOC_MANUAL_SYNC_ACTIVE_JOB_ID
        job = _GDOC_MANUAL_SYNC_JOBS.get(job_id)
        if not job:
            if _GDOC_MANUAL_SYNC_ACTIVE_JOB_ID == job_id:
                _GDOC_MANUAL_SYNC_ACTIVE_JOB_ID = None
            return

        job['status'] = status
        job['finished_at'] = now_iso
        job['percent'] = 100
        if result is not None:
            job['result'] = result
            job['run_status'] = result.get('run_status', status)
            job['result_message'] = result.get('message', job.get('message', ''))
            job['patients'] = int(result.get('patients', 0) or 0)
            job['groups'] = int(result.get('groups', 0) or 0)
            job['pushed_groups'] = int(result.get('pushed_groups', 0) or 0)
            job['warnings'] = result.get('warnings', [])
            job['errors'] = result.get('errors', [])
            if result.get('targets_total') is not None:
                job['targets_total'] = int(result.get('targets_total') or 0)
            if result.get('targets_processed') is not None:
                job['targets_processed'] = int(result.get('targets_processed') or 0)
            if status in {'success', 'partial', 'failed'}:
                job['message'] = result.get('message') or job.get('message')
        if error_message:
            job['error'] = str(error_message)
            job['message'] = str(error_message)
        if _GDOC_MANUAL_SYNC_ACTIVE_JOB_ID == job_id:
            _GDOC_MANUAL_SYNC_ACTIVE_JOB_ID = None
        _cleanup_manual_sync_jobs_locked()


def _snapshot_manual_sync_job(job_id):
    with _GDOC_MANUAL_SYNC_JOB_LOCK:
        job = _GDOC_MANUAL_SYNC_JOBS.get(job_id)
        if not job:
            return None
        return dict(job)


def _run_manual_google_docs_sync_job(job_id):
    try:
        with app.app_context():
            db = get_db()
            with _GDOC_AUTO_SYNC_LOCK:
                result = _run_google_docs_auto_sync(
                    db,
                    force=True,
                    trigger_source='manual',
                    progress_callback=lambda payload: _update_manual_sync_job(job_id, payload),
                )

            if not result.get('ran'):
                reason = result.get('reason')
                if reason == 'dependency':
                    status = 'failed'
                    message = '; '.join(result.get('errors') or ['Google dependencies unavailable'])
                elif reason in {'no-targets', 'no-connected-targets'}:
                    status = 'failed'
                    message = 'No connected Google Docs are selected for automatic sync.'
                elif reason == 'not-due':
                    status = 'skipped'
                    message = 'Google Docs auto-sync is not due yet.'
                else:
                    status = 'failed'
                    message = 'Google Docs auto-sync did not run.'

                result_payload = {
                    'status': status,
                    'run_status': status,
                    'synced': int(result.get('total_synced') or 0),
                    'patients': int(result.get('synced_patients') or 0),
                    'groups': int(result.get('synced_groups') or 0),
                    'pushed_groups': int(result.get('pushed_groups') or 0),
                    'targets_total': int(result.get('targets_total') or 0),
                    'targets_processed': int(result.get('targets_processed') or 0),
                    'errors': result.get('errors') or [message],
                    'warnings': result.get('warnings') or [],
                    'history_id': result.get('history_id'),
                    'message': message,
                }
                _complete_manual_sync_job(job_id, status=status, result=result_payload)
                return

            run_status = str(result.get('status') or 'success').lower()
            result_payload = {
                'status': 'ok',
                'run_status': run_status,
                'synced': int(result.get('total_synced') or 0),
                'patients': int(result.get('synced_patients') or 0),
                'groups': int(result.get('synced_groups') or 0),
                'pushed_groups': int(result.get('pushed_groups') or 0),
                'targets_total': int(result.get('targets_total') or 0),
                'targets_processed': int(result.get('targets_processed') or 0),
                'errors': result.get('errors') or [],
                'warnings': result.get('warnings') or [],
                'history_id': result.get('history_id'),
                'message': f"Synced {int(result.get('total_synced') or 0)} records from Google Docs.",
            }
            _complete_manual_sync_job(job_id, status=run_status, result=result_payload)
    except Exception as exc:
        app.logger.exception('Manual Google Docs sync job failed')
        _complete_manual_sync_job(job_id, status='error', error_message=f'Manual sync failed: {exc}')

@app.context_processor
def inject_translations():
    def t(text):
        if session.get('lang') == 'he':
            return HEBREW_TRANSLATIONS.get(text, text)
        return text
    ui_density = (session.get('ui_density') or 'balanced').strip().lower()
    if ui_density not in {'compact', 'balanced', 'large'}:
        ui_density = 'balanced'
    return dict(t=t, lang=session.get('lang', 'en'), ui_density=ui_density)

def _get_notification_unread_count(db, user):
    if not getattr(user, 'is_authenticated', False):
        return 0

    try:
        if getattr(user, 'role', None) == 'admin':
            row = db.execute('''
                SELECT COUNT(*) AS count
                FROM notifications
                WHERE COALESCE(is_read, 0) = 0
                  AND (COALESCE(audience, 'admin') = 'admin' OR recipient_user_id = ?)
            ''', (user.id,)).fetchone()
        else:
            row = db.execute('''
                SELECT COUNT(*) AS count
                FROM notifications
                WHERE COALESCE(is_read, 0) = 0
                  AND recipient_user_id = ?
            ''', (user.id,)).fetchone()
    except sqlite3.OperationalError:
        return 0
    return int((row['count'] if row else 0) or 0)


@app.context_processor
def inject_global_vars():
    unread_messages = 0
    notification_unread_count = 0
    db = None

    try:
        db = get_db()
    except Exception:
        db = None

    pending_cancel_count = 0
    if current_user.is_authenticated and db is not None:
        unread_messages = db.execute(
            'SELECT COUNT(*) as count FROM messages WHERE recipient_id = ? AND is_read = 0',
            (current_user.id,)
        ).fetchone()['count']
        notification_unread_count = _get_notification_unread_count(db, current_user)
        if current_user.role == 'admin':
            pending_cancel_count = db.execute(
                "SELECT COUNT(*) as count FROM cancel_requests WHERE status = 'pending'"
            ).fetchone()['count']

    return dict(
        unread_messages=unread_messages,
        notification_unread_count=notification_unread_count,
        pending_cancel_count=pending_cancel_count,
        site_settings=get_site_settings(db)
    )

@app.route('/set_lang/<lang>')
def set_lang(lang):
    if lang in ['en', 'he']:
        session['lang'] = lang
    return redirect(request.referrer or url_for('index'))


@app.route('/set_density/<density>')
def set_density(density):
    normalized = (density or '').strip().lower()
    if normalized in {'compact', 'balanced', 'large'}:
        session['ui_density'] = normalized
    return redirect(request.referrer or url_for('index'))


@app.before_request
def enforce_inactivity_timeout():
    if request.path.startswith('/static/'):
        return

    if not current_user.is_authenticated:
        session.pop('last_activity_at', None)
        return

    timeout_minutes = int(app.config.get('INACTIVITY_TIMEOUT_MINUTES', 5) or 5)
    timeout_seconds = max(timeout_minutes, 1) * 60
    now_ts = int(datetime.now(timezone.utc).timestamp())
    last_activity_at = session.get('last_activity_at')

    if last_activity_at is not None:
        try:
            idle_seconds = now_ts - int(last_activity_at)
        except (TypeError, ValueError):
            idle_seconds = 0

        if idle_seconds >= timeout_seconds:
            logout_user()
            session.pop('last_activity_at', None)
            flash('Session expired due to inactivity. Please log in again.')
            return redirect(url_for('login'))

    session['last_activity_at'] = now_ts


@app.before_request
def enforce_session_version_match():
    if request.path.startswith('/static/'):
        return

    if not current_user.is_authenticated:
        session.pop('session_version', None)
        return

    db = get_db()
    row = db.execute('SELECT session_version FROM users WHERE id = ?', (current_user.id,)).fetchone()
    if not row:
        logout_user()
        session.pop('session_version', None)
        flash('Session expired. Please log in again.')
        return redirect(url_for('login'))

    db_version = int((row['session_version'] if row['session_version'] is not None else 0) or 0)
    session_version = session.get('session_version')
    try:
        session_version = int(session_version if session_version is not None else -1)
    except (TypeError, ValueError):
        session_version = -1

    if session_version != db_version:
        logout_user()
        session.pop('session_version', None)
        session.pop('last_activity_at', None)
        flash('Your session was invalidated after a security change. Please sign in again.')
        return redirect(url_for('login'))


@app.before_request
def security_retention_guard():
    if request.path.startswith('/static/'):
        return

    if app.config.get('TESTING'):
        return

    try:
        db = get_db()
        _security_retention_cleanup(db)
    except Exception:
        app.logger.exception('Security retention guard failed')


@app.before_request
def routine_backup_guard():
    if request.path.startswith('/static/'):
        return
    if app.config.get('TESTING'):
        return

    db_path = app.config.get('DATABASE_PATH') or app.config.get('DATABASE', DATABASE)
    if not db_path or not os.path.exists(db_path):
        return

    last_modified = os.path.getmtime(db_path)
    now = _time.time()

    if now - last_modified > 86400: # 24 hours
        try:
            perform_routine_encrypted_backup(db_path)
        except Exception:
            app.logger.exception('Routine encrypted backup failed')


@app.before_request
def gdocs_auto_sync_guard():
    if request.path.startswith('/static/'):
        return
    if app.config.get('TESTING'):
        return

    global _GDOC_AUTO_SYNC_LAST_CHECK_TS

    import time

    now_ts = time.time()
    # Avoid checking on every request; due logic is handled inside the sync runner.
    if now_ts - float(_GDOC_AUTO_SYNC_LAST_CHECK_TS or 0.0) < 60:
        return
    if not _GDOC_AUTO_SYNC_LOCK.acquire(blocking=False):
        return

    try:
        _GDOC_AUTO_SYNC_LAST_CHECK_TS = now_ts
        db = get_db()
        _run_google_docs_auto_sync(db, force=False, trigger_source='request')
    except Exception:
        app.logger.exception('Google Docs auto-sync guard failed')
    finally:
        _GDOC_AUTO_SYNC_LOCK.release()


@app.after_request
def apply_security_headers(response):
    response.headers.setdefault('X-Content-Type-Options', 'nosniff')
    response.headers.setdefault('X-Frame-Options', 'SAMEORIGIN')
    response.headers.setdefault('Referrer-Policy', 'strict-origin-when-cross-origin')
    response.headers.setdefault('Permissions-Policy', 'camera=(), microphone=(), geolocation=()')
    csp_policy = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' 'unsafe-eval'; "
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com; "
        "font-src 'self' data: https://fonts.gstatic.com; "
        "img-src 'self' data: https://api.qrserver.com https://*.googleusercontent.com; "
        "connect-src 'self'; "
        "form-action 'self';"
    )
    response.headers.setdefault('Content-Security-Policy', csp_policy)
    return response


def _run_security_scan_logic():
    import subprocess
    import json
    import shutil
    
    # Locate virtual environment Python to run bandit/pip-audit if possible
    root_path = app.root_path
    venv_dir = os.path.abspath(os.path.join(root_path, '../venv'))
    if not os.path.isdir(venv_dir):
        venv_dir = os.path.abspath(os.path.join(root_path, '.venv'))
    
    bandit_bin = os.path.join(venv_dir, 'bin/bandit') if os.path.isdir(venv_dir) else 'bandit'
    pip_audit_bin = os.path.join(venv_dir, 'bin/pip-audit') if os.path.isdir(venv_dir) else 'pip-audit'
    
    bandit_available = shutil.which(bandit_bin) is not None
    pip_audit_available = shutil.which(pip_audit_bin) is not None
    
    # Run Bandit
    bandit_count = 0
    bandit_results = []
    if bandit_available:
        try:
            cmd = [bandit_bin, '-r', 'app.py', 'clinic_app/', '-f', 'json', '-ll']
            res = subprocess.run(cmd, capture_output=True, text=True, cwd=root_path)
            if res.stdout:
                data = json.loads(res.stdout)
                bandit_count = len(data.get('results', []))
                for item in data.get('results', []):
                    bandit_results.append({
                        'file': item.get('filename'),
                        'line': item.get('line_number'),
                        'issue_text': item.get('issue_text'),
                        'severity': item.get('issue_severity'),
                        'confidence': item.get('issue_confidence'),
                    })
        except Exception as exc:
            bandit_results.append({'error': f'Failed to run bandit: {exc}'})
    else:
        bandit_results.append({'error': 'bandit is not installed. Install with: pip install bandit'})
        
    # Run pip-audit
    pip_audit_count = 0
    pip_audit_results = []
    if pip_audit_available:
        try:
            cmd = [pip_audit_bin, '-r', 'requirements.txt', '-f', 'json']
            res = subprocess.run(cmd, capture_output=True, text=True, cwd=root_path)
            if res.stdout:
                data = json.loads(res.stdout)
                dependencies = data if isinstance(data, list) else data.get('dependencies', [])
                for dep in dependencies:
                    vulns = dep.get('vulns', [])
                    if vulns:
                        pip_audit_count += len(vulns)
                        for v in vulns:
                            pip_audit_results.append({
                                'package': dep.get('name'),
                                'version': dep.get('version'),
                                'id': v.get('id'),
                                'fix_versions': v.get('fix_versions', []),
                                'description': v.get('description'),
                            })
        except Exception as exc:
            pip_audit_results.append({'error': f'Failed to run pip-audit: {exc}'})
    else:
        pip_audit_results.append({'error': 'pip-audit is not installed. Install with: pip install pip-audit'})
        
    status = 'ok'
    if bandit_count > 0 or pip_audit_count > 0:
        status = 'warning'
        for b in bandit_results:
            if b.get('severity', '').upper() == 'HIGH':
                status = 'error'
                
    summary = {
        'status': status,
        'run_at': datetime.now(timezone.utc).isoformat(),
        'bandit': {
            'total_issues': bandit_count,
            'issues': bandit_results[:20],
        },
        'pip_audit': {
            'total_vulnerabilities': pip_audit_count,
            'vulnerabilities': pip_audit_results[:20],
        }
    }
    return summary


def _run_automated_security_scan(db, force=False):
    settings = get_site_settings(db)
    enabled = str(settings.get('security_scan_enabled') or '0') in {'1', 'true', 'yes', 'on'}
    interval_key = settings.get('security_scan_interval') or 'daily'
    
    if not force:
        if not enabled or interval_key == 'disabled':
            return {'ran': False, 'reason': 'disabled'}
        
        last_run_raw = settings.get('security_scan_last_run_at')
        if last_run_raw:
            try:
                last_run = datetime.fromisoformat(last_run_raw)
                now = datetime.now(timezone.utc)
                if last_run.tzinfo is None:
                    last_run = last_run.replace(tzinfo=timezone.utc)
                elapsed = (now - last_run).total_seconds()
                interval_seconds = 24 * 3600 if interval_key == 'daily' else 7 * 24 * 3600
                if elapsed < interval_seconds:
                    return {'ran': False, 'reason': 'not-due'}
            except Exception:
                pass
                
    results = _run_security_scan_logic()
    now_str = datetime.now(timezone.utc).isoformat()
    save_site_settings(db, {
        'security_scan_last_run_at': now_str,
        'security_scan_last_status': results['status'],
        'security_scan_last_results_json': json.dumps(results),
    })
    return {'ran': True, 'results': results}




def _request_expects_json_error():
    if request.path.startswith('/api/'):
        return True
    return request.accept_mimetypes.best == 'application/json'


@app.errorhandler(404)
def handle_not_found(error):
    if _request_expects_json_error():
        return jsonify({'status': 'error', 'message': 'Not found'}), 404
    return render_template('404.html'), 404


@app.errorhandler(400)
def handle_bad_request(error):
    if _request_expects_json_error():
        msg = str(error) if str(error) else 'Bad request'
        msg = re.sub(r'<[^>]+>', '', msg).strip()
        return jsonify({'status': 'error', 'message': msg}), 400
    return '<h1>400 Bad Request</h1>', 400


@app.errorhandler(500)
def handle_internal_error(error):
    app.logger.exception('Unhandled server error: %s', error)
    if _request_expects_json_error():
        return jsonify({'status': 'error', 'message': 'Internal server error'}), 500
    return render_template('500.html'), 500


def _gdocs_auto_sync_worker_loop():
    while not _GDOC_AUTO_SYNC_STOP_EVENT.is_set():
        if _GDOC_AUTO_SYNC_STOP_EVENT.wait(timeout=60):
            break
        if not _GDOC_AUTO_SYNC_LOCK.acquire(blocking=False):
            continue
        try:
            with app.app_context():
                db = get_db()
                _run_google_docs_auto_sync(db, force=False, trigger_source='worker')
        except Exception:
            app.logger.exception('Google Docs auto-sync worker iteration failed')
        finally:
            _GDOC_AUTO_SYNC_LOCK.release()


@app.before_request
def enforce_2fa_timeout():
    if request.path.startswith('/static/'):
        return
    if app.config.get('TESTING'):
        return
    if not current_user.is_authenticated:
        return
    if not session.get('totp_enabled'):
        return

    import time
    otp_verified_at = session.get('otp_verified_at', 0)
    if time.time() - float(otp_verified_at) > 1800:
        session.pop('otp_verified_at', None)
        logout_user()
        if request.is_json or request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.path.startswith('/api/'):
            return jsonify({'status': 'session_expired', 'message': 'Session expired due to inactivity.'}), 401
        flash('Session expired due to inactivity. Please log in again.', 'warning')
        return redirect(url_for('login'))
    session['otp_verified_at'] = time.time()



def ensure_gdocs_auto_sync_worker_started():
    if app.config.get('TESTING'):
        return
    with _GDOC_AUTO_SYNC_WORKER_STATE_LOCK:
        global _GDOC_AUTO_SYNC_WORKER_STARTED
        if _GDOC_AUTO_SYNC_WORKER_STARTED:
            return
        worker = threading.Thread(target=_gdocs_auto_sync_worker_loop, name='gdocs-auto-sync-worker', daemon=True)
        worker.start()
        _GDOC_AUTO_SYNC_WORKER_STARTED = True

class User(UserMixin):
    def __init__(self, id, username, role, patient_id=None, display_name=None, session_version=0):
        self.id = id
        self.username = username
        self.role = role
        self.patient_id = patient_id
        self.display_name = display_name or username
        self.session_version = int(session_version or 0)

@login_manager.user_loader
def load_user(user_id):
    db = get_db()
    user = db.execute('SELECT * FROM users WHERE id = ?', (user_id,)).fetchone()
    if user:
        return User(
            user['id'],
            user['username'],
            user['role'],
            user['patient_id'],
            user['display_name'],
            user['session_version'],
        )
    return None

def get_db():
    db = getattr(g, '_database', None)
    if db is None:
        database = app.config.get('DATABASE', DATABASE)
        db = g._database = sqlite3.connect(database, timeout=5)
        db.row_factory = sqlite3.Row
        db.execute("PRAGMA busy_timeout=5000")
    return db

@app.teardown_appcontext
def close_connection(exception):
    db = getattr(g, '_database', None)
    if db is not None:
        db.close()


def _verify_totp_code(secret, candidate_code):
    if not secret:
        return False
    normalized = re.sub(r'\s+', '', str(candidate_code or ''))
    if not normalized.isdigit():
        return False
    return pyotp.TOTP(secret).verify(normalized, valid_window=1)


def _admin_totp_uri(user_row, secret):
    issuer = 'Private Clinic CRM'
    account = user_row['username']
    return pyotp.TOTP(secret).provisioning_uri(name=account, issuer_name=issuer)


def _login_redirect_for_user(user_row):
    user_obj = User(
        user_row['id'],
        user_row['username'],
        user_row['role'],
        user_row['patient_id'],
        user_row['display_name'],
        user_row['session_version'],
    )
    login_user(user_obj)
    session.permanent = True
    session['session_version'] = int(user_row['session_version'] or 0)
    import time
    session['otp_verified_at'] = time.time()
    session['totp_enabled'] = bool(user_row['totp_enabled'] and user_row['totp_secret'])

    if user_row['role'] == 'admin':
        if not user_row['totp_enabled'] or not user_row['totp_secret']:
            flash('Set up two-factor authentication from the admin profile before continuing.')
            return redirect(url_for('admin_profile'))
        if user_row['force_password_change']:
            flash('Admin password must be changed before continuing.')
            return redirect(url_for('admin_profile'))
        return redirect(url_for('patients'))

    if user_row['role'] == 'patient' and user_row['force_password_change']:
        return redirect(url_for('patient_change_password'))

    return redirect(url_for('patient_home'))


register_auth_routes(
    app,
    get_db=get_db,
    verify_totp_code=_verify_totp_code,
    login_redirect_for_user=_login_redirect_for_user,
    dummy_password_hash=DUMMY_PASSWORD_HASH,
    send_smtp_email=_send_smtp_email,
    validate_password_strength=_validate_password_strength,
)


def _migrate_add_column(db, table, column, col_type, default=None):
    default_clause = f" DEFAULT {default}" if default is not None else ""
    try:
        db.execute(f'ALTER TABLE {table} ADD COLUMN {column} {col_type}{default_clause}')
    except sqlite3.OperationalError:
        pass


def _run_db_migrations(db):
    """Run all schema migrations and index creations."""
    db.execute('''
        CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            sender_id INTEGER NOT NULL,
            recipient_id INTEGER,
            content TEXT NOT NULL,
            timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            is_read BOOLEAN DEFAULT 0,
            FOREIGN KEY (sender_id) REFERENCES users (id),
            FOREIGN KEY (recipient_id) REFERENCES users (id)
        )
    ''')

    # Handle column migrations
    _migrate_add_column(db, 'appointments', 'duration_minutes', 'INTEGER', '60') # Column exists
    _migrate_add_column(db, 'appointments', 'is_recurring', 'BOOLEAN', '0')
    _migrate_add_column(db, 'appointments', 'recurrence_interval', 'INTEGER')
    _migrate_add_column(db, 'appointments', 'recurrence_days', 'TEXT')
    _migrate_add_column(db, 'appointments', 'meeting_type', 'TEXT', '"in-person"')
    _migrate_add_column(db, 'appointments', 'meeting_link', 'TEXT')
    _migrate_add_column(db, 'appointments', 'recurrence_end_date', 'DATE')
    _migrate_add_column(db, 'appointments', 'recurrence_count', 'INTEGER')
    _migrate_add_column(db, 'notes', 'content_hebrew', 'TEXT')
    _migrate_add_column(db, 'notes', 'note_date', 'DATE')
    _migrate_add_column(db, 'notes', 'patient_appearance', 'TEXT')
    _migrate_add_column(db, 'notes', 'key_topics', 'TEXT')
    _migrate_add_column(db, 'notes', 'updated_at', 'TIMESTAMP')
    _migrate_add_column(db, 'notes', 'behavior_checklist', 'TEXT')
    _migrate_add_column(db, 'notes', 'mood_summary', 'TEXT')
    _migrate_add_column(db, 'notes', 'behavior_notes', 'TEXT')
    _migrate_add_column(db, 'notes', 'is_missed_meeting', 'BOOLEAN', '0')
    _migrate_add_column(db, 'notes', 'missed_reason', 'TEXT')
    _migrate_add_column(db, 'files', 'treatment_id', 'INTEGER')
    _migrate_add_column(db, 'notes', 'appointment_id', 'INTEGER')
    _migrate_add_column(db, 'notes', 'session_number', 'INTEGER')
    _migrate_add_column(db, 'notes', 'needs_review', 'BOOLEAN', '0')
    try:
        db.execute('''CREATE TABLE IF NOT EXISTS resources (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            description TEXT,
            url TEXT,
            is_public BOOLEAN DEFAULT 0,
            allow_patient_view BOOLEAN DEFAULT 1,
            allow_patient_download BOOLEAN DEFAULT 1,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''')
    except sqlite3.OperationalError:
        pass
    try:
        db.execute('''CREATE TABLE IF NOT EXISTS patient_resources (
            patient_id INTEGER NOT NULL,
            resource_id INTEGER NOT NULL,
            PRIMARY KEY (patient_id, resource_id),
            FOREIGN KEY (patient_id) REFERENCES patients(id),
            FOREIGN KEY (resource_id) REFERENCES resources(id)
        )''')
    except sqlite3.OperationalError:
        pass
    _migrate_add_column(db, 'patient_resources', 'assigned_at', 'TIMESTAMP', 'CURRENT_TIMESTAMP')
    _migrate_add_column(db, 'resources', 'allow_patient_view', 'BOOLEAN', '1')
    _migrate_add_column(db, 'resources', 'allow_patient_download', 'BOOLEAN', '1')
    try:
        db.execute('UPDATE resources SET allow_patient_view = COALESCE(allow_patient_view, 1), allow_patient_download = COALESCE(allow_patient_download, 1)')
    except sqlite3.OperationalError:
        pass
    _migrate_add_column(db, 'files', 'treatment_id', 'INTEGER')
    _migrate_add_column(db, 'patients', 'background', 'TEXT')
    _migrate_add_column(db, 'patients', 'treatment_info', 'TEXT')
    _migrate_add_column(db, 'patients', 'profile_image', 'TEXT')
    try:
        db.execute('''CREATE TABLE IF NOT EXISTS site_settings (
            setting_key TEXT PRIMARY KEY,
            setting_value TEXT
        )''')
    except sqlite3.OperationalError:
        pass
    try:
        db.execute('''CREATE TABLE IF NOT EXISTS google_oauth_pending_states (
            state TEXT PRIMARY KEY,
            user_id INTEGER NOT NULL,
            redirect_uri TEXT NOT NULL,
            code_verifier TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )''')
        db.execute('CREATE INDEX IF NOT EXISTS idx_google_oauth_pending_created_at ON google_oauth_pending_states(created_at)')
    except sqlite3.OperationalError:
        pass
    try:
        db.execute('''CREATE TABLE IF NOT EXISTS gdocs_sync_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            run_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            trigger_source TEXT,
            status TEXT,
            interval_key TEXT,
            targets_total INTEGER DEFAULT 0,
            targets_processed INTEGER DEFAULT 0,
            synced_total INTEGER DEFAULT 0,
            synced_patients INTEGER DEFAULT 0,
            synced_groups INTEGER DEFAULT 0,
            pushed_groups INTEGER DEFAULT 0,
            errors_json TEXT,
            details_json TEXT
        )''')
    except sqlite3.OperationalError:
        pass
    try:
        db.execute('''CREATE TABLE IF NOT EXISTS availability (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            slot_date DATE,
            slot_time TIME NOT NULL,
            duration_minutes INTEGER DEFAULT 60,
            recurrence TEXT,
            weekday INTEGER CHECK(weekday >= 0 AND weekday <= 6),
            status TEXT NOT NULL DEFAULT 'available',
            booked_by_name TEXT,
            booked_by_phone TEXT,
            booked_notes TEXT,
            booked_at TIMESTAMP,
            share_token TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''')
    except sqlite3.OperationalError:
        pass
    _migrate_add_column(db, 'slots_override', 'duration_minutes', 'INTEGER', '60')
    _migrate_add_column(db, 'patients', 'can_self_schedule', 'BOOLEAN', '0')
    try:
        db.execute('''CREATE TABLE IF NOT EXISTS blocked_slots (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            blocked_date DATE NOT NULL,
            blocked_time TIME NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''')
    except sqlite3.OperationalError:
        pass
    _migrate_add_column(db, 'blocked_slots', 'duration_minutes', 'INTEGER', '60')
    _migrate_add_column(db, 'blocked_slots', 'title', 'TEXT')
    _migrate_add_column(db, 'blocked_slots', 'is_private', 'BOOLEAN', '0')
    try:
        db.execute("ALTER TABLE blocked_slots ADD COLUMN block_type TEXT DEFAULT 'blocked'")
    except sqlite3.OperationalError:
        pass
    _migrate_add_column(db, 'blocked_slots', 'created_by', 'INTEGER')
    try:
        db.execute("ALTER TABLE patients ADD COLUMN patient_type TEXT DEFAULT 'private'")
    except sqlite3.OperationalError:
        pass
    _migrate_add_column(db, 'patients', 'intake_assessment', 'TEXT')
    _migrate_add_column(db, 'patients', 'intake_questionnaire', 'TEXT')
    try:
        db.execute("ALTER TABLE patients ADD COLUMN is_deleted BOOLEAN DEFAULT 0")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE patients ADD COLUMN deleted_at TIMESTAMP")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE patients ADD COLUMN deleted_reason TEXT")
    except sqlite3.OperationalError:
        pass
    _migrate_add_column(db, 'patients', 'birth_date', 'DATE')
    _migrate_add_column(db, 'patients', 'id_number', 'TEXT')
    _migrate_add_column(db, 'patients', 'has_intake_tab', 'BOOLEAN', '0')
    _migrate_add_column(db, 'patients', 'has_questionnaire_tab', 'BOOLEAN', '0')
    try:
        db.execute('''CREATE TABLE IF NOT EXISTS patient_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_id INTEGER NOT NULL,
            encounter_date DATE,
            title TEXT,
            content TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (patient_id) REFERENCES patients (id)
        )''')
    except sqlite3.OperationalError:
        pass
    try:
        db.execute('CREATE INDEX IF NOT EXISTS idx_patient_logs_patient_date ON patient_logs(patient_id, encounter_date)')
    except sqlite3.OperationalError:
        pass
    db.execute("UPDATE patients SET status = 'candidate' WHERE status IN ('waiting', 'waiting for scheduling')")
    _migrate_add_column(db, 'appointments', 'meeting_platform', 'TEXT')
    _migrate_add_column(db, 'appointments', 'meeting_title', 'TEXT')
    _migrate_add_column(db, 'appointments', 'missed_reason', 'TEXT')
    _migrate_add_column(db, 'appointments', 'save_to_google', 'BOOLEAN', '0')
    _migrate_add_column(db, 'appointments', 'excluded_dates', 'TEXT')
    _migrate_add_column(db, 'appointments', 'cancelled_dates', 'TEXT')
    _migrate_add_column(db, 'appointments', 'recurrence_group_id', 'TEXT')
    _migrate_add_column(db, 'group_sessions', 'cancelled_dates', 'TEXT')
    _migrate_add_column(db, 'group_sessions', 'recurrence_group_id', 'TEXT')
    _migrate_add_column(db, 'users', 'display_name', 'TEXT')
    _migrate_add_column(db, 'users', 'email', 'TEXT')
    _migrate_add_column(db, 'users', 'phone', 'TEXT')
    _migrate_add_column(db, 'users', 'id_number', 'TEXT')
    _migrate_add_column(db, 'users', 'birth_date', 'DATE')
    _migrate_add_column(db, 'users', 'totp_secret', 'TEXT')
    _migrate_add_column(db, 'users', 'totp_enabled', 'BOOLEAN', '0')
    _migrate_add_column(db, 'users', 'force_password_change', 'BOOLEAN', '0')
    _migrate_add_column(db, 'users', 'session_version', 'INTEGER', '0')
    _migrate_add_column(db, 'users', 'totp_recovery_codes', 'TEXT')

    _migrate_add_column(db, 'slots_override', 'share_token', 'TEXT')
    _migrate_add_column(db, 'slots_override', 'booked_by_name', 'TEXT')
    _migrate_add_column(db, 'slots_override', 'booked_by_phone', 'TEXT')
    _migrate_add_column(db, 'slots_override', 'booked_notes', 'TEXT')
    _migrate_add_column(db, 'slots_override', 'booked_at', 'TIMESTAMP')
    try:
        db.execute('''CREATE TABLE IF NOT EXISTS vacancy_recurring (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            weekday INTEGER NOT NULL CHECK(weekday >= 0 AND weekday <= 6),
            slot_time TEXT NOT NULL,
            duration_minutes INTEGER DEFAULT 60,
            is_active BOOLEAN DEFAULT 1,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''')
    except sqlite3.OperationalError:
        pass
    _migrate_add_column(db, 'vacancy_recurring', 'duration_minutes', 'INTEGER', '60')
    try:
        db.execute('''CREATE TABLE IF NOT EXISTS audit_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_id INTEGER,
            action TEXT NOT NULL,
            details TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''')
    except sqlite3.OperationalError:
        pass

    try:
        db.execute('''CREATE TABLE IF NOT EXISTS password_reset_tokens (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            token_hash TEXT NOT NULL UNIQUE,
            expires_at TIMESTAMP NOT NULL,
            used_at TIMESTAMP,
            requested_ip TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )''')
    except sqlite3.OperationalError:
        pass
    try:
        db.execute('CREATE INDEX IF NOT EXISTS idx_password_reset_tokens_user_id ON password_reset_tokens(user_id)')
    except sqlite3.OperationalError:
        pass
    try:
        db.execute('CREATE INDEX IF NOT EXISTS idx_password_reset_tokens_expires_at ON password_reset_tokens(expires_at)')
    except sqlite3.OperationalError:
        pass

    try:
        db.execute('''CREATE TABLE IF NOT EXISTS notifications (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT,
            message TEXT NOT NULL,
            recipient_user_id INTEGER,
            sender_id INTEGER,
            audience TEXT DEFAULT 'admin',
            is_read BOOLEAN DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''')
    except sqlite3.OperationalError:
        pass
    _migrate_add_column(db, 'notifications', 'title', 'TEXT')
    _migrate_add_column(db, 'notifications', 'recipient_user_id', 'INTEGER')
    _migrate_add_column(db, 'notifications', 'sender_id', 'INTEGER')
    try:
        db.execute("ALTER TABLE notifications ADD COLUMN audience TEXT DEFAULT 'admin'")
    except sqlite3.OperationalError:
        pass

    try:
        db.execute('''CREATE TABLE IF NOT EXISTS public_booking_links (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            token TEXT UNIQUE NOT NULL,
            created_by INTEGER,
            is_active BOOLEAN DEFAULT 1,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''')
    except sqlite3.OperationalError:
        pass

    try:
        db.execute('''CREATE TABLE IF NOT EXISTS goals (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_id INTEGER NOT NULL,
            description TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'active',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (patient_id) REFERENCES patients (id)
        )''')
    except sqlite3.OperationalError:
        pass

    try:
        db.execute('''CREATE TABLE IF NOT EXISTS groups (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            group_type TEXT DEFAULT 'support',
            description TEXT,
            is_active BOOLEAN DEFAULT 1,
            gdoc_id TEXT,
            gdoc_watch_channel TEXT,
            gdoc_watch_expiry TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''')
    except sqlite3.OperationalError:
        pass
    _migrate_add_column(db, 'groups', 'gdoc_id', 'TEXT')
    _migrate_add_column(db, 'groups', 'gdoc_watch_channel', 'TEXT')
    _migrate_add_column(db, 'groups', 'gdoc_watch_expiry', 'TEXT')

    try:
        db.execute('''CREATE TABLE IF NOT EXISTS group_members (
            group_id INTEGER NOT NULL,
            patient_id INTEGER NOT NULL,
            joined_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            left_at TIMESTAMP,
            role TEXT DEFAULT 'member',
            PRIMARY KEY (group_id, patient_id),
            FOREIGN KEY (group_id) REFERENCES groups (id),
            FOREIGN KEY (patient_id) REFERENCES patients (id)
        )''')
    except sqlite3.OperationalError:
        pass

    try:
        db.execute('''CREATE TABLE IF NOT EXISTS group_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            group_id INTEGER NOT NULL,
            session_date DATE NOT NULL,
            session_time TIME NOT NULL,
            duration_minutes INTEGER DEFAULT 60,
            title TEXT,
            facilitator TEXT,
            meeting_type TEXT DEFAULT 'in-person',
            meeting_link TEXT,
            series_id INTEGER,
            occurrence_index INTEGER,
            session_summary TEXT,
            status TEXT DEFAULT 'scheduled',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (group_id) REFERENCES groups (id)
        )''')
    except sqlite3.OperationalError:
        pass

    _migrate_add_column(db, 'group_sessions', 'series_id', 'INTEGER')
    _migrate_add_column(db, 'group_sessions', 'occurrence_index', 'INTEGER')
    _migrate_add_column(db, 'group_sessions', 'session_summary', 'TEXT')
    _migrate_add_column(db, 'group_sessions', 'supervision_id', 'INTEGER')

    try:
        db.execute('''CREATE TABLE IF NOT EXISTS group_member_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            group_id INTEGER NOT NULL,
            patient_id INTEGER NOT NULL,
            joined_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            left_at TIMESTAMP,
            role TEXT DEFAULT 'member',
            FOREIGN KEY (group_id) REFERENCES groups (id),
            FOREIGN KEY (patient_id) REFERENCES patients (id)
        )''')
    except sqlite3.OperationalError:
        pass

    try:
        db.execute('''CREATE TABLE IF NOT EXISTS group_session_series (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            group_id INTEGER NOT NULL,
            start_date DATE NOT NULL,
            start_time TIME NOT NULL,
            duration_minutes INTEGER DEFAULT 60,
            recurrence_interval_weeks INTEGER DEFAULT 1,
            recurrence_end_date DATE,
            recurrence_count INTEGER,
            title TEXT,
            facilitator TEXT,
            meeting_type TEXT DEFAULT 'in-person',
            meeting_link TEXT,
            is_active BOOLEAN DEFAULT 1,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (group_id) REFERENCES groups (id)
        )''')
    except sqlite3.OperationalError:
        pass

    try:
        db.execute('''CREATE TABLE IF NOT EXISTS group_session_attendance (
            session_id INTEGER NOT NULL,
            patient_id INTEGER NOT NULL,
            attendance_status TEXT NOT NULL DEFAULT 'pending',
            absence_reason TEXT,
            notified_on_time BOOLEAN DEFAULT 0,
            attendance_note TEXT,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            PRIMARY KEY (session_id, patient_id),
            FOREIGN KEY (session_id) REFERENCES group_sessions (id),
            FOREIGN KEY (patient_id) REFERENCES patients (id)
        )''')
    except sqlite3.OperationalError:
        pass
    _migrate_add_column(db, 'group_session_attendance', 'notified_on_time', 'BOOLEAN', '0')

    # Performance indexes for common filters and sort paths.
    db.execute('CREATE INDEX IF NOT EXISTS idx_patients_status_deleted ON patients(status, is_deleted)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_patients_type_deleted ON patients(patient_type, is_deleted)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_patients_name ON patients(name COLLATE NOCASE)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_patients_email ON patients(email)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_patients_phone ON patients(phone)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_users_patient_id ON users(patient_id)')

    db.execute('CREATE INDEX IF NOT EXISTS idx_appointments_patient_date_time ON appointments(patient_id, appointment_date, appointment_time)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_appointments_patient_status_date ON appointments(patient_id, status, appointment_date)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_appointments_date_time_status ON appointments(appointment_date, appointment_time, status)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_appointments_recurrence_group ON appointments(recurrence_group_id)')

    db.execute('CREATE INDEX IF NOT EXISTS idx_notes_patient_created ON notes(patient_id, created_at)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_receipts_patient_created ON receipts(patient_id, created_at)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_files_patient_created ON files(patient_id, created_at)')

    db.execute('CREATE INDEX IF NOT EXISTS idx_messages_recipient_read_time ON messages(recipient_id, is_read, timestamp)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_messages_sender_recipient_time ON messages(sender_id, recipient_id, timestamp)')

    db.execute('CREATE INDEX IF NOT EXISTS idx_availability_date_time_status ON availability(slot_date, slot_time, status)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_availability_weekday_recurrence ON availability(weekday, recurrence)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_availability_share_token ON availability(share_token)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_blocked_slots_date_time ON blocked_slots(blocked_date, blocked_time)')

    db.execute('CREATE INDEX IF NOT EXISTS idx_group_members_patient_left ON group_members(patient_id, left_at)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_group_sessions_date_time_status ON group_sessions(session_date, session_time, status)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_group_member_history_group_patient ON group_member_history(group_id, patient_id, joined_at)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_group_series_group_start ON group_session_series(group_id, start_date)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_group_attendance_session_status ON group_session_attendance(session_id, attendance_status)')

    db.execute('CREATE INDEX IF NOT EXISTS idx_notifications_read_created ON notifications(is_read, created_at)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_goals_patient_status ON goals(patient_id, status)')

    try:
        db.execute('''CREATE TABLE IF NOT EXISTS supervisions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_id INTEGER,
            group_id INTEGER,
            supervision_date DATE NOT NULL,
            supervisor_name TEXT,
            content TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (patient_id) REFERENCES patients (id),
            FOREIGN KEY (group_id) REFERENCES groups (id)
        )''')
    except sqlite3.OperationalError:
        pass
    db.execute('CREATE INDEX IF NOT EXISTS idx_supervisions_patient ON supervisions(patient_id, supervision_date)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_supervisions_group ON supervisions(group_id, supervision_date)')

    try:
        db.execute('''CREATE TABLE IF NOT EXISTS diagnosis_documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_id INTEGER NOT NULL,
            category TEXT NOT NULL DEFAULT 'test_document',
            title TEXT,
            original_filename TEXT NOT NULL,
            stored_filename TEXT NOT NULL,
            notes TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (patient_id) REFERENCES patients (id)
        )''')
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE diagnosis_documents ADD COLUMN category TEXT NOT NULL DEFAULT 'test_document'")
    except sqlite3.OperationalError:
        pass
    _migrate_add_column(db, 'diagnosis_documents', 'title', 'TEXT')
    _migrate_add_column(db, 'diagnosis_documents', 'original_filename', 'TEXT')
    _migrate_add_column(db, 'diagnosis_documents', 'stored_filename', 'TEXT')
    _migrate_add_column(db, 'diagnosis_documents', 'notes', 'TEXT')
    db.execute('CREATE INDEX IF NOT EXISTS idx_diagnosis_documents_patient ON diagnosis_documents(patient_id, category, created_at)')

    # Google Calendar: add google_event_id to appointments and group_sessions
    _migrate_add_column(db, 'appointments', 'google_event_id', 'TEXT')
    _migrate_add_column(db, 'group_sessions', 'google_event_id', 'TEXT')
    # Ensure google_calendar_tokens table exists
    db.execute('''
        CREATE TABLE IF NOT EXISTS google_calendar_tokens (
            id INTEGER PRIMARY KEY,
            owner TEXT NOT NULL DEFAULT 'admin',
            token_json TEXT NOT NULL,
            calendar_id TEXT NOT NULL DEFAULT 'primary',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # Treatment method tag and manual sort order for patients
    _migrate_add_column(db, 'patients', 'treatment_method', 'TEXT')
    _migrate_add_column(db, 'patients', 'sort_order', 'INTEGER')
    db.execute('''CREATE TABLE IF NOT EXISTS treatment_method_options (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        label TEXT NOT NULL UNIQUE,
        display_order INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    # Seed default options (only inserts if they don't exist yet)
    for _label in ['Psychodynamic', 'CBT', 'EFT', 'Management', '15 sessions', '3 sessions']:
        db.execute('INSERT OR IGNORE INTO treatment_method_options (label) VALUES (?)', (_label,))

    # Google Docs integration columns
    _migrate_add_column(db, 'patients', 'gdoc_id', 'TEXT')
    _migrate_add_column(db, 'patients', 'gdoc_watch_channel', 'TEXT')
    _migrate_add_column(db, 'patients', 'gdoc_watch_expiry', 'TEXT')
    _migrate_add_column(db, 'patients', 'questionnaires_file_id', 'TEXT')
    _migrate_add_column(db, 'patients', 'questionnaires_file_url', 'TEXT')
    _migrate_add_column(db, 'patients', 'questionnaires_selected', 'TEXT')

    _migrate_add_column(db, 'notes', 'link_url', 'TEXT')
    _migrate_add_column(db, 'patient_logs', 'link_url', 'TEXT')

    # Public contact inquiries submitted from the About page by unauthenticated visitors
    db.execute('''CREATE TABLE IF NOT EXISTS contact_inquiries (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        email TEXT,
        phone TEXT,
        message TEXT NOT NULL,
        is_read BOOLEAN DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    # Shared rate-limit store (used across Gunicorn workers)
    db.execute('''CREATE TABLE IF NOT EXISTS rate_limits (
        bucket_key TEXT NOT NULL,
        scope TEXT NOT NULL,
        timestamp_real REAL NOT NULL,
        PRIMARY KEY (bucket_key, scope, timestamp_real)
    )''')
    db.execute('''CREATE INDEX IF NOT EXISTS idx_rate_limits_lookup
        ON rate_limits (bucket_key, scope, timestamp_real)''')

    # Appointment reminder columns
    try:
        db.execute("ALTER TABLE patients ADD COLUMN reminder_email_enabled BOOLEAN DEFAULT 1")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE appointments ADD COLUMN reminder_sent_at TIMESTAMP")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE appointments ADD COLUMN reminder_hours_before INTEGER")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE notes ADD COLUMN share_with_patient BOOLEAN DEFAULT 0")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE receipts ADD COLUMN receipt_number TEXT")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE patients ADD COLUMN treatment_plan TEXT")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE receipts ADD COLUMN status TEXT DEFAULT 'paid'")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE receipts ADD COLUMN morning_doc_id TEXT")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE receipts ADD COLUMN morning_sync_status TEXT DEFAULT 'pending'")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE receipts ADD COLUMN morning_synced_at TIMESTAMP")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE receipts ADD COLUMN vat_rate REAL DEFAULT 0")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE receipts ADD COLUMN vat_amount REAL DEFAULT 0")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE receipts ADD COLUMN net_amount REAL")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE receipts ADD COLUMN payment_method TEXT")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE receipts ADD COLUMN document_type TEXT DEFAULT 'receipt'")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE receipts ADD COLUMN client_email TEXT")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE assessments ADD COLUMN answers_json TEXT DEFAULT '{}'")
    except sqlite3.OperationalError:
        pass
    db.execute('''
        CREATE TABLE IF NOT EXISTS assessment_questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            assessment_type_id INTEGER NOT NULL REFERENCES assessment_types(id),
            question_order INTEGER NOT NULL,
            question_key TEXT NOT NULL,
            question_text_en TEXT NOT NULL,
            question_text_he TEXT NOT NULL,
            question_type TEXT NOT NULL DEFAULT 'radio',
            options_json TEXT DEFAULT '[]',
            required INTEGER DEFAULT 1,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    try:
        db.execute("ALTER TABLE patients ADD COLUMN street TEXT")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE patients ADD COLUMN city TEXT")
    except sqlite3.OperationalError:
        pass
    try:
        db.execute("ALTER TABLE patients ADD COLUMN zip_code TEXT")
    except sqlite3.OperationalError:
        pass

    db.execute('''
        CREATE TABLE IF NOT EXISTS service_types (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            description TEXT,
            default_price REAL NOT NULL,
            is_active INTEGER DEFAULT 1,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    db.execute('''
        CREATE TABLE IF NOT EXISTS receipt_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            receipt_id INTEGER NOT NULL,
            service_type_id INTEGER,
            quantity INTEGER DEFAULT 1,
            unit_price REAL NOT NULL,
            line_total REAL NOT NULL,
            description TEXT,
            FOREIGN KEY (receipt_id) REFERENCES receipts(id) ON DELETE CASCADE,
            FOREIGN KEY (service_type_id) REFERENCES service_types(id)
        )
    ''')
    db.execute('''
        CREATE TABLE IF NOT EXISTS cancel_requests (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            appointment_id INTEGER NOT NULL,
            patient_id INTEGER NOT NULL,
            reason TEXT NOT NULL,
            status TEXT DEFAULT 'pending',
            reviewed_by INTEGER,
            reviewed_at TIMESTAMP,
            admin_note TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (appointment_id) REFERENCES appointments(id),
            FOREIGN KEY (patient_id) REFERENCES patients(id),
            FOREIGN KEY (reviewed_by) REFERENCES users(id)
        )
    ''')

    db.execute('CREATE INDEX IF NOT EXISTS idx_cancel_requests_appointment ON cancel_requests(appointment_id)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_cancel_requests_patient_status ON cancel_requests(patient_id, status)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_cancel_requests_status_created ON cancel_requests(status, created_at)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_schedules_patient ON schedules(patient_id)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_slots_patient ON slots(patient_id)')

    db.execute('''CREATE INDEX IF NOT EXISTS idx_appointments_reminder_pending
        ON appointments (appointment_date, appointment_time)
        WHERE COALESCE(status, 'scheduled') = 'scheduled' AND reminder_sent_at IS NULL''')

    db.execute('''CREATE TABLE IF NOT EXISTS reminder_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        appointment_id INTEGER NOT NULL,
        patient_id INTEGER NOT NULL,
        recipient_email TEXT NOT NULL,
        status TEXT NOT NULL,
        error_message TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (appointment_id) REFERENCES appointments(id),
        FOREIGN KEY (patient_id) REFERENCES patients(id)
    )''')
    db.execute('''CREATE INDEX IF NOT EXISTS idx_reminder_log_appointment
        ON reminder_log (appointment_id)''')
    db.execute('''CREATE INDEX IF NOT EXISTS idx_reminder_log_patient_status
        ON reminder_log (patient_id, status)''')

    db.execute('''CREATE TABLE IF NOT EXISTS email_reminder_templates (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        event_type TEXT NOT NULL UNIQUE,
        hours_before REAL NOT NULL DEFAULT 24.0,
        subject_template TEXT NOT NULL DEFAULT 'Appointment Reminder',
        body_template TEXT NOT NULL DEFAULT '',
        enabled BOOLEAN NOT NULL DEFAULT 1,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    db.execute('''CREATE TABLE IF NOT EXISTS incoming_email (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        message_id TEXT,
        from_email TEXT NOT NULL,
        from_name TEXT,
        subject TEXT NOT NULL,
        body_text TEXT,
        body_html TEXT,
        related_type TEXT,
        related_id INTEGER,
        is_read BOOLEAN DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    db.execute('CREATE INDEX IF NOT EXISTS idx_incoming_email_read ON incoming_email(is_read)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_incoming_email_created ON incoming_email(created_at)')

    db.commit()

def _seed_admin_user(db):
    """Seed the default admin user and handle legacy migrations."""
    env_username = (os.environ.get('ADMIN_USERNAME') or '').strip()
    if not env_username:
        env_username = 'lioraloni'
    env_password = (os.environ.get('ADMIN_PASSWORD') or '').strip()

    # Check if admin exists
    admin = db.execute("SELECT * FROM users WHERE role = 'admin' ORDER BY id ASC").fetchone()
    if not admin:
        print("Creating default admin user...")
        if env_password:
            hashed_pw = generate_password_hash(env_password)
            force_change = 0
        else:
            # No password configured — generate a random one and force immediate change.
            generated = secrets.token_urlsafe(16)
            hashed_pw = generate_password_hash(generated)
            force_change = 1
            print(f"ADMIN_PASSWORD not set. Temporary password: {generated!r}  (you will be required to change it on first login)")
        db.execute(
            "INSERT OR IGNORE INTO users (username, password_hash, role, force_password_change) VALUES (?, ?, ?, ?)",
            (env_username, hashed_pw, 'admin', force_change)
        )
        db.commit()
        print(f"Admin user created (username: {env_username}).")
        admin = db.execute("SELECT * FROM users WHERE role = 'admin' ORDER BY id ASC").fetchone()

    # One-time migration from legacy default admin credentials.
    legacy_admin = db.execute("SELECT * FROM users WHERE username = 'admin' AND role = 'admin'").fetchone()
    if legacy_admin and legacy_admin['id'] != admin['id']:
        collision = db.execute("SELECT id FROM users WHERE username = ? AND id <> ?", (env_username, legacy_admin['id'])).fetchone()
        if not collision:
            if env_password:
                new_hash = generate_password_hash(env_password)
                force_change = 0
            else:
                generated = secrets.token_urlsafe(16)
                new_hash = generate_password_hash(generated)
                force_change = 1
                print(f"ADMIN_PASSWORD not set. Temporary password for migrated admin: {generated!r}")
            db.execute(
                '''
                UPDATE users
                SET username = ?, password_hash = ?, force_password_change = ?
                WHERE id = ?
                ''',
                (env_username, new_hash, force_change, legacy_admin['id'])
            )
            db.commit()
            admin = db.execute("SELECT * FROM users WHERE id = ?", (legacy_admin['id'],)).fetchone()
            print(f'Legacy admin account migrated to {env_username}.')

    # Never modify existing admin user — env vars only apply on FIRST creation above
    # This ensures password, TOTP, username changes persist across redeploys
    if admin and not admin['display_name']:
        db.execute('UPDATE users SET display_name = ? WHERE id = ?', ('Admin', admin['id']))
        db.commit()


_db_init_lock = threading.Lock()
_db_initialized = False

def init_db():
    global _db_initialized

    if _db_initialized:
        return

    with _db_init_lock:
        if _db_initialized:
            return

        database = app.config.get('DATABASE', DATABASE)
        with app.app_context():
            db = get_db()
            db.execute("PRAGMA journal_mode=WAL")
            db.execute("PRAGMA busy_timeout=5000")
            # Check if DB already has schema (e.g. alembic_version exists)
            already_initialized = False
            try:
                row = db.execute(
                    "SELECT name FROM sqlite_master WHERE type='table' AND name='alembic_version'"
                ).fetchone()
                if row:
                    already_initialized = True
            except Exception:
                pass

            if not already_initialized:
                with app.open_resource('clinic_app/schema.sql', mode='r') as f:
                    db.cursor().executescript(f.read())
                db.commit()

            _run_db_migrations(db)
            _seed_admin_user(db)

            if not already_initialized:
                print(f"Initialized the database at {database}.")

            if not app.config.get('TESTING'):
                try:
                    perform_routine_encrypted_backup(database)
                except Exception as backup_error:
                    print(f"Routine backup skipped: {backup_error}")
                ensure_gdocs_auto_sync_worker_started()
                ensure_appointment_reminder_worker_started()

        _db_initialized = True

    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
        print(f"Created upload folder: {app.config['UPLOAD_FOLDER']}")

@app.route('/')
def index():
    try:
        get_db()
    except sqlite3.OperationalError:
        init_db()

    if current_user.is_authenticated:
        if current_user.role == 'admin':
            return redirect(url_for('admin_dashboard'))
        elif current_user.role == 'patient':
            return redirect(url_for('patient_home'))
    return redirect(url_for('login'))


@app.route('/about')
def about_page():
    settings = get_site_settings()
    is_admin_preview = current_user.is_authenticated and current_user.role == 'admin'
    if settings.get('about_enabled') != '1' and not is_admin_preview:
        return 'Page not found', 404
    map_urls = _build_about_map_urls(settings.get('about_map_url'))
    form_data = session.pop('contact_form_data', {})
    return render_template(
        'about.html',
        site_settings=settings,
        is_admin_preview=is_admin_preview,
        about_map_open_url=map_urls['open_url'],
        about_map_embed_url=map_urls['embed_url'],
        form_data=form_data,
    )


@app.route('/accessibility-statement')
def accessibility_statement():
    settings = get_site_settings()
    return render_template('accessibility_statement.html', site_settings=settings)


@app.route('/privacy-policy')
def privacy_policy():
    settings = get_site_settings()
    return render_template('privacy_policy.html', site_settings=settings)



def normalize_summary_text(text):
    if not text:
        return ''
    return ' '.join(str(text).replace('\n', ' ').split())


def split_summary_segments(text):
    clean_text = normalize_summary_text(text)
    if not clean_text:
        return []

    segments = []
    for segment in re.split(r'[.!?\n\u05c3]+', clean_text):
        segment = segment.strip(' ,;:-')
        if len(segment) >= 18:
            segments.append(segment)
    return segments


def extract_background_sentence(text):
    segments = split_summary_segments(text)
    if segments:
        return segments[0][:180].strip()
    return normalize_summary_text(text)[:180].strip()


def trim_summary_segment(segment, limit=140):
    segment = normalize_summary_text(segment)
    if len(segment) <= limit:
        return segment.rstrip(' ,;:')

    trimmed = segment[:limit].rsplit(' ', 1)[0].rstrip(' ,;:')
    return f'{trimmed}...'


def find_best_summary_segment(texts, patient_name, keywords, prefer_earlier=True):
    best_segment = ''
    best_score = -1
    normalized_name = normalize_summary_text(patient_name)

    for index, text in enumerate(texts):
        for segment in split_summary_segments(text):
            score = 0
            for keyword in keywords:
                if keyword in segment:
                    score += 2
            if normalized_name and normalized_name in segment:
                score += 2
            if prefer_earlier:
                score += max(0, 4 - index)
            if score > best_score:
                best_score = score
                best_segment = segment

    return trim_summary_segment(best_segment) if best_score > 0 else ''


def pick_top_summary_topics(texts, topics, limit):
    counts = Counter()
    for text in texts:
        clean_text = normalize_summary_text(text)
        for label, keywords in topics.items():
            hits = sum(clean_text.count(keyword) for keyword in keywords)
            if hits:
                counts[label] += hits

    return [label for label, _ in counts.most_common(limit)]


def extract_age_fact(texts, patient_name):
    if not patient_name:
        return ''

    escaped_name = re.escape(patient_name)
    patterns = [
        rf'{escaped_name}[^.!?\n]{{0,40}}?\b(בן|בת)\s+(\d{{1,2}})\b',
        rf'\b(בן|בת)\s+(\d{{1,2}})\b[^.!?\n]{{0,40}}?{escaped_name}'
    ]

    for text in texts[:4]:
        clean_text = normalize_summary_text(text)
        for pattern in patterns:
            match = re.search(pattern, clean_text)
            if match:
                return f"גיל מתועד: {match.group(1)} {match.group(2)}"

    return 'גיל מדויק לא תועד במפורש'


def extract_occupation_fact(texts, patient_name):
    segment = find_best_summary_segment(
        texts,
        patient_name,
        ['עובד', 'עובדת', 'עבודה', 'לומד', 'לומדת', 'מכללה', 'מפעל', 'תפקיד', 'מנהל', 'מנהלת', 'צבא'],
        prefer_earlier=True
    )
    if not segment:
        return ''
    return f'בהיבט התפקודי/תעסוקתי עלה כי {segment}'


def extract_children_count(corpus):
    match = re.search(r'(?:יש\s+ל[וה]\s+|אם\s+ל|אב\s+ל)(\d+|אחד|אחת|שני|שניים|שתיים|שתי|שלושה|שלוש|ארבעה|ארבע|חמישה|חמש)\s+ילדים', corpus)
    if not match:
        return ''

    raw_count = match.group(1)
    return HEBREW_NUMBER_WORDS.get(raw_count, raw_count)


def extract_family_fact(texts):
    corpus = ' '.join(normalize_summary_text(text) for text in texts if text)
    facts = []

    if any(keyword in corpus for keyword in ['בעלה שנפטר', 'פטירת האב', 'פטירה של בעל', 'בן זוגה שנפטר', 'בעלה נפטר']):
        facts.append('מתמודד/ת עם אובדן בן או בת הזוג')

    children_count = extract_children_count(corpus)
    if children_count:
        facts.append(f'הורה ל-{children_count} ילדים')
    elif 'ילדים' in corpus or 'ילדיה' in corpus:
        facts.append('יחסיו/ה עם הילדים הם מוקד משמעותי')

    if any(keyword in corpus for keyword in ['אמא', 'אביה', 'אביו', 'אחים', 'אחיו', 'אחיה', 'משפחת המקור']):
        facts.append('עולה עיסוק משמעותי גם במשפחת המקור')

    return '; '.join(facts[:3])


def extract_recent_focus(notes):
    recent_texts = []
    for note in notes[-3:]:
        if note['mood_summary']:
            recent_texts.append(note['mood_summary'])
        if note['content']:
            recent_texts.append(note['content'])

    recent_topics = pick_top_summary_topics(recent_texts, BACKGROUND_THEME_TOPICS, 2)
    if recent_topics:
        return ', '.join(recent_topics)

    if notes:
        return extract_background_sentence(notes[-1]['mood_summary'] or notes[-1]['content'])
    return ''


def extract_key_summary_points(notes, limit=3):
    scored_segments = []
    keyword_sets = list(BACKGROUND_REASON_TOPICS.values()) + list(BACKGROUND_THEME_TOPICS.values())

    recent_notes = list(notes[-8:])
    for idx, note in enumerate(reversed(recent_notes)):
        note_text = ' '.join([
            normalize_summary_text(note['mood_summary'] or ''),
            normalize_summary_text(note['content'] or '')
        ]).strip()
        if not note_text:
            continue
        for segment in split_summary_segments(note_text):
            score = max(0, 6 - idx)
            for keywords in keyword_sets:
                if any(keyword in segment for keyword in keywords):
                    score += 2
            if len(segment) > 140:
                score += 1
            scored_segments.append((score, segment))

    unique_segments = []
    seen_prefix = set()
    for _, segment in sorted(scored_segments, key=lambda pair: pair[0], reverse=True):
        key = normalize_summary_text(segment)[:64]
        if not key or key in seen_prefix:
            continue
        seen_prefix.add(key)
        unique_segments.append(trim_summary_segment(segment, 170))
        if len(unique_segments) >= limit:
            break

    return unique_segments


def normalize_intake_payload(payload):
    if not isinstance(payload, dict):
        return {}
    allowed_fields = set(intake_form_fields())
    normalized = {}
    for key, value in payload.items():
        key_text = str(key or '').strip()
        if not key_text:
            continue
        if key_text.startswith('intake_'):
            key_text = key_text[7:]
        if key_text not in allowed_fields:
            continue
        if isinstance(value, list):
            clean_values = [str(item or '').strip() for item in value if str(item or '').strip()]
            normalized[key_text] = ', '.join(clean_values)
        else:
            normalized[key_text] = str(value or '').strip()
    return normalized


def parse_legacy_intake_text(raw_value):
    text = str(raw_value or '').strip()
    if not text:
        return {}

    label_map = {
        'main complaint': 'main_complaint',
        'problem history / current illness': 'problem_history',
        'problem history': 'problem_history',
        'early anamnesis': 'early_anamnesis',
    }

    parsed = {}
    current_key = None
    current_lines = []

    def flush_current():
        if current_key is None:
            return
        value = '\n'.join(current_lines).strip()
        if value:
            parsed[current_key] = value

    for line in text.splitlines():
        stripped = line.strip()
        lowered = stripped.lower()
        if lowered.endswith(':'):
            candidate = lowered[:-1].strip()
            mapped = label_map.get(candidate)
            if mapped:
                flush_current()
                current_key = mapped
                current_lines = []
                continue
        if current_key is not None:
            current_lines.append(stripped)

    flush_current()
    if parsed:
        return parsed

    # Fall back to using the entire legacy text as the main complaint.
    return {'main_complaint': text}


def parse_intake_questionnaire(raw_value, fallback_assessment=None):
    if raw_value:
        try:
            parsed = json.loads(raw_value)
            normalized = normalize_intake_payload(parsed)
            if normalized:
                return normalized
        except (json.JSONDecodeError, TypeError):
            pass

        # Some legacy records were stored as Python dict strings (True/False/None, single quotes).
        raw_text = str(raw_value).strip()
        if raw_text.startswith('{') and raw_text.endswith('}'):
            try:
                converted = raw_text.replace("'", '"')
                converted = converted.replace('True', 'true').replace('False', 'false').replace('None', 'null')
                parsed = json.loads(converted)
                normalized = normalize_intake_payload(parsed)
                if normalized:
                    return normalized
            except (ValueError, TypeError, json.JSONDecodeError):
                pass

        legacy_from_questionnaire = parse_legacy_intake_text(raw_value)
        if legacy_from_questionnaire:
            return legacy_from_questionnaire

    legacy_from_assessment = parse_legacy_intake_text(fallback_assessment)
    if legacy_from_assessment:
        return legacy_from_assessment

    return {}


def intake_form_fields():
    return [
        'meeting_location', 'meeting_location_specify', 'meeting_time', 'meeting_duration', 'meeting_conductor',
        'main_complaint', 'problem_history', 'early_anamnesis', 'referral_source', 'referral_date',
        'family_status', 'guardian_status', 'guardian_by_whom', 'living_with', 'living_with_other',
        'disability_status', 'disability_percent', 'self_harm_level', 'self_harm_recent', 'self_harm_count',
        'forced_treatment', 'substance_use', 'medical_cannabis', 'alcohol_use',
        'medical_conditions', 'psychiatric_conditions',
        'appearance_fit', 'appearance_fit_note', 'appearance_ordered', 'appearance_ordered_note',
        'cooperation', 'cooperation_note', 'eye_contact', 'eye_contact_note',
        'behavior_normal', 'behavior_note', 'speech_style', 'speech_note',
        'mood', 'mood_note', 'affect_match', 'affect_state', 'affect_note',
        'thinking_normal', 'thinking_rate', 'thinking_sequence', 'thinking_content',
        'perception_normal', 'perception_abnormal', 'reality_testing', 'judgment', 'self_insight',
        'orientation', 'memory',
        'referral_target', 'referral_details', 'patient_consent',
        'treatment_approach', 'treatment_frequency', 'treatment_estimated_duration',
        'diag_referral_question', 'diag_test_battery', 'diag_observations',
        'diag_differential', 'diag_impression', 'diag_recommendations',
        'diag_followup_plan', 'diag_final_summary'
    ]


def intake_multi_select_fields():
    return {
        'appearance_fit',
        'appearance_ordered',
        'behavior_normal',
        'speech_style',
        'mood',
        'affect_match',
        'affect_state',
        'thinking_normal',
        'thinking_rate',
        'thinking_sequence',
        'thinking_content',
        'referral_target',
    }


def intake_data_from_request(form, existing_data=None):
    if not any(key.startswith('intake_') for key in form.keys()):
        return None
    data = normalize_intake_payload(existing_data) if isinstance(existing_data, dict) else {}
    allowed = set(intake_form_fields())
    multi_fields = intake_multi_select_fields()
    for key in intake_form_fields():
        field_name = f'intake_{key}'
        if key not in allowed:
            continue
        if field_name not in form:
            continue
        if key in multi_fields:
            values = [value.strip() for value in form.getlist(field_name) if value and value.strip()]
            data[key] = ', '.join(values)
        else:
            raw = form.get(field_name, '')
            data[key] = (raw or '').strip()
    return data


def serialize_intake_assessment(data):
    main_complaint = data.get('main_complaint', '')
    problem_history = data.get('problem_history', '')
    early_anamnesis = data.get('early_anamnesis', '')
    parts = []
    if main_complaint:
        parts.append(f"Main complaint:\n{main_complaint}")
    if problem_history:
        parts.append(f"Problem history / current illness:\n{problem_history}")
    if early_anamnesis:
        parts.append(f"Early anamnesis:\n{early_anamnesis}")
    return '\n\n'.join(parts).strip()


def build_intake_docx(patient_name, intake_data, language='en'):
    document = Document()
    title = f"Intake Summary - {patient_name}" if language != 'he' else f"סיכום אינטייק - {patient_name}"
    document.add_heading(title, level=1)

    label_map = {
        'main_complaint': ('Main complaint', 'תלונה עיקרית'),
        'problem_history': ('Problem history / current illness', 'היסטוריה של הבעיה / מחלה נוכחית'),
        'early_anamnesis': ('Early anamnesis', 'אנמנזה מוקדמת'),
        'referral_source': ('Referral source', 'גורם מפנה'),
        'referral_date': ('Referral date', 'תאריך הפניה'),
        'meeting_location': ('Meeting location', 'מיקום הפגישה'),
        'summary': ('Summary', 'סיכום'),
    }

    for field in intake_form_fields():
        value = (intake_data.get(field) or '').strip()
        if not value:
            continue
        label_en, label_he = label_map.get(field, (field.replace('_', ' ').title(), field.replace('_', ' ')))
        document.add_heading(label_he if language == 'he' else label_en, level=2)
        document.add_paragraph(value)

    if len(document.paragraphs) == 1:
        empty_text = 'No intake data available.' if language != 'he' else 'אין נתוני אינטייק זמינים.'
        document.add_paragraph(empty_text)

    return document



def _get_intake_data(patient_row):
    if not patient_row:
        return {}, ''
    intake_questionnaire = parse_intake_questionnaire(
        patient_row['intake_questionnaire'],
        patient_row['intake_assessment']
    )
    return intake_questionnaire, patient_row['intake_assessment'] or ''


def _extract_main_problem_no_notes(patient_row):
    intake_questionnaire, intake_assessment = _get_intake_data(patient_row)
    main_complaint = (intake_questionnaire.get('main_complaint') or '').strip()
    problem_history = (intake_questionnaire.get('problem_history') or '').strip()

    main_problem = main_complaint or problem_history
    if not main_problem and intake_assessment.strip():
        main_problem = intake_assessment.strip().splitlines()[0]
    return main_problem


def _extract_main_problem_with_notes(patient_row, notes, reason_topics, theme_topics, recent_focus):
    intake_questionnaire, intake_assessment = _get_intake_data(patient_row)
    main_complaint = (intake_questionnaire.get('main_complaint') or '').strip()
    problem_history = (intake_questionnaire.get('problem_history') or '').strip()

    intake_problem_line = ''
    if intake_assessment:
        for line in intake_assessment.splitlines():
            cleaned = line.strip()
            if cleaned and not cleaned.lower().startswith('main complaint:'):
                intake_problem_line = cleaned
                break

    if main_complaint:
        return main_complaint
    elif problem_history:
        return problem_history
    elif intake_problem_line:
        return intake_problem_line
    elif reason_topics:
        return f"קושי מרכזי סביב {', '.join(reason_topics)}"
    elif theme_topics:
        return f"מוקד קושי חוזר סביב {', '.join(theme_topics)}"
    elif recent_focus:
        return recent_focus
    else:
        return extract_background_sentence(notes[-1]['mood_summary'] or notes[-1]['content'])


def _get_notes_timeframe(notes):
    first_date = notes[0]['note_date']
    last_date = notes[-1]['note_date']

    if first_date and last_date:
        return f" בין {first_date} ל-{last_date}"
    if last_date:
        return f" עד המפגש האחרון המתועד ב-{last_date}"
    return ''


def _format_patient_summary(patient_name, notes, timeframe, age_fact, occupation_fact, family_fact, reason_topics, theme_topics, key_points, recent_focus, main_problem):
    parts = [f"סיכום מטופל: {patient_name}."]
    parts.append(f"תמונת זמן: {len(notes)} מפגשים מתועדים{timeframe}.")

    identity_parts = [age_fact]
    if occupation_fact:
        identity_parts.append(occupation_fact)
    if family_fact:
        identity_parts.append(family_fact)
    parts.append(f"פרופיל רקע: {'; '.join(identity_parts)}.")

    if reason_topics:
        parts.append(f"סיבות ופניות מרכזיות: {', '.join(reason_topics)}.")
    if theme_topics:
        parts.append(f"דפוסים חוזרים לאורך המפגשים: {', '.join(theme_topics)}.")
    if key_points:
        parts.append(f"תובנות מפתח מהתיעוד: {' | '.join(key_points)}.")
    if recent_focus:
        parts.append(f"מיקוד עדכני לטווח הקרוב: {recent_focus}.")

    parts.append(f"תמצית קלינית נוכחית: {main_problem}.")

    return ' '.join(part.strip() for part in parts if part).strip()


def build_patient_background_from_notes(db, patient_id, patient_name=None):
    if patient_name is None:
        patient_row = db.execute('SELECT name FROM patients WHERE id = ?', (patient_id,)).fetchone()
        patient_name = patient_row['name'] if patient_row else 'המטופל/ת'

    patient_row = db.execute('''
        SELECT intake_assessment, intake_questionnaire
        FROM patients
        WHERE id = ?
    ''', (patient_id,)).fetchone()

    notes = db.execute('''
        SELECT note_date, content, mood_summary, created_at
        FROM notes
        WHERE patient_id = ?
        ORDER BY COALESCE(note_date, date(created_at)) ASC, created_at ASC
    ''', (patient_id,)).fetchall()

    if not notes:
        main_problem = _extract_main_problem_no_notes(patient_row)
        if main_problem:
            return (
                f"סיכום מטופל: {patient_name}. "
                "סטטוס תיעוד: מידע ראשוני מאינטייק בלבד (ללא מפגשים שוטפים מתועדים). "
                f"מוקד עיקרי נוכחי: {main_problem}."
            )
        return 'לא נמצאה היסטוריה טיפולית מתועדת במערכת.'

    note_texts = []
    for note in notes:
        note_texts.append(note['content'])
        if note['mood_summary']:
            note_texts.append(note['mood_summary'])

    timeframe = _get_notes_timeframe(notes)

    age_fact = extract_age_fact(note_texts, patient_name)
    occupation_fact = extract_occupation_fact(note_texts, patient_name)
    family_fact = extract_family_fact(note_texts)
    reason_topics = pick_top_summary_topics(note_texts[:8], BACKGROUND_REASON_TOPICS, 2)
    theme_topics = pick_top_summary_topics(note_texts, BACKGROUND_THEME_TOPICS, 2)
    recent_focus = extract_recent_focus(notes)
    key_points = extract_key_summary_points(notes, limit=3)

    main_problem = _extract_main_problem_with_notes(patient_row, notes, reason_topics, theme_topics, recent_focus)

    return _format_patient_summary(
        patient_name, notes, timeframe, age_fact, occupation_fact, family_fact,
        reason_topics, theme_topics, key_points, recent_focus, main_problem
    )


def _normalize_session_number(value):
    if value is None:
        return None
    raw = str(value).strip()
    if not raw:
        return None
    try:
        return str(int(raw))
    except (TypeError, ValueError):
        return raw


def _parse_note_fields(item):
    meeting_number = _normalize_session_number(item.get('meeting_number') or item.get('session_number'))
    date_str = (item.get('date') or item.get('note_date') or '').strip() or None
    content_text = (item.get('content') or '').strip()
    appearance_text = (item.get('patient_appearance') or '').strip()
    checklist_text = item.get('behavior_checklist')
    if isinstance(checklist_text, list):
        checklist_text = ','.join([str(i).strip() for i in checklist_text if str(i).strip()])
    checklist_text = (checklist_text or '').strip()
    mood_summary = (item.get('mood_summary') or '').strip()
    behavior_notes = (item.get('behavior_notes') or '').strip()

    if not meeting_number and not _has_meaningful_note_information(
        content_text,
        mood_summary,
        behavior_notes,
        appearance_text,
        checklist_text,
    ):
        return None

    if not content_text:
        content_text = mood_summary or behavior_notes or appearance_text
    if not content_text:
        return None

    return {
        'meeting_number': meeting_number,
        'date_str': date_str,
        'content_text': content_text,
        'appearance_text': appearance_text,
        'checklist_text': checklist_text,
        'mood_summary': mood_summary,
        'behavior_notes': behavior_notes
    }


def _import_flat_patient_history(db, patient_id, data):
    appointments_added = 0
    notes_added = 0

    def _sort_key(item):
        raw_date = (item.get('date') or item.get('note_date') or '').strip()
        meeting_raw = item.get('meeting_number') or item.get('session_number') or 0
        try:
            meeting_num = int(meeting_raw)
        except (TypeError, ValueError):
            meeting_num = 0
        return (raw_date, meeting_num)

    for item in sorted(data, key=_sort_key):
        parsed = _parse_note_fields(item)
        if not parsed:
            continue

        appt_id = None
        if parsed['date_str']:
            existing = db.execute('SELECT id FROM appointments WHERE patient_id = ? AND appointment_date = ?', (patient_id, parsed['date_str'])).fetchone()
            if not existing:
                cursor = db.execute('INSERT INTO appointments (patient_id, appointment_date, appointment_time, status) VALUES (?, ?, ?, ?)', (patient_id, parsed['date_str'], '00:00', 'completed'))
                appt_id = cursor.lastrowid
                appointments_added += 1
            else:
                appt_id = existing['id']

        db.execute(
            '''INSERT INTO notes (patient_id, appointment_id, session_number, note_date, content,
                                  patient_appearance, behavior_checklist, mood_summary, behavior_notes)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''',
            (
                patient_id,
                appt_id,
                parsed['meeting_number'],
                parsed['date_str'],
                parsed['content_text'],
                parsed['appearance_text'],
                parsed['checklist_text'],
                parsed['mood_summary'],
                parsed['behavior_notes']
            )
        )
        notes_added += 1

    return appointments_added, notes_added, 0

def _import_structured_patient_history(db, patient_id, data):
    appointments_added = 0
    notes_added = 0
    receipts_added = 0

    # Import appointments
    appt_id_map = {}
    sorted_appts = sorted(data.get('appointments', []), key=lambda x: (x.get('appointment_date', ''), x.get('appointment_time', '')))
    for appt in sorted_appts:
        existing = db.execute('SELECT id FROM appointments WHERE patient_id = ? AND appointment_date = ? AND appointment_time = ?',
            (patient_id, appt.get('appointment_date'), appt.get('appointment_time'))).fetchone()
        if not existing:
            cursor = db.execute('''INSERT INTO appointments
                (patient_id, appointment_date, appointment_time, cost, duration_minutes, is_recurring, recurrence_interval, recurrence_days, meeting_type, meeting_link, status, recurrence_end_date, recurrence_count)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                (patient_id, appt.get('appointment_date'), appt.get('appointment_time'), appt.get('cost'), appt.get('duration_minutes'),
                 appt.get('is_recurring'), appt.get('recurrence_interval'), appt.get('recurrence_days'), appt.get('meeting_type'),
                 appt.get('meeting_link'), appt.get('status'), appt.get('recurrence_end_date'), appt.get('recurrence_count')))
            appt_id_map[appt.get('id')] = cursor.lastrowid
            appointments_added += 1
        else:
            appt_id_map[appt.get('id')] = existing['id']

    # Import notes
    sorted_notes = sorted(
        data.get('notes', []),
        key=lambda x: (
            x.get('note_date') or x.get('date') or x.get('created_at', ''),
            str(x.get('session_number') or x.get('meeting_number') or '')
        )
    )
    for note in sorted_notes:
        parsed = _parse_note_fields(note)
        if not parsed:
            continue

        new_appt_id = appt_id_map.get(note.get('appointment_id')) if note.get('appointment_id') else None

        db.execute('''INSERT INTO notes
            (patient_id, appointment_id, session_number, note_date, content, patient_appearance,
             behavior_checklist, mood_summary, behavior_notes, needs_review, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
            (
                patient_id,
                new_appt_id,
                parsed['meeting_number'],
                parsed['date_str'],
                parsed['content_text'],
                parsed['appearance_text'],
                parsed['checklist_text'],
                parsed['mood_summary'],
                parsed['behavior_notes'],
                note.get('needs_review'),
                note.get('created_at')
            ))
        notes_added += 1

    # Import receipts
    for receipt in data.get('receipts', []):
        db.execute('''INSERT INTO receipts
            (patient_id, amount, description, created_at)
            VALUES (?, ?, ?, ?)''',
            (patient_id, receipt.get('amount'), receipt.get('description'), receipt.get('created_at')))
        receipts_added += 1

    return appointments_added, notes_added, receipts_added


def _has_meaningful_note_information(content_text, mood_summary, behavior_notes, appearance_text, checklist_text):
    placeholders = {
        '',
        'n/a',
        'na',
        'none',
        'unknown',
        'yyyy-mm-dd',
        'brief mood summary.',
        'short behavior notes.',
        'general appearance observations.',
    }

    values = [content_text, mood_summary, behavior_notes, appearance_text, checklist_text]
    for value in values:
        cleaned = (value or '').strip().lower()
        if cleaned and cleaned not in placeholders:
            return True
    return False


def _get_patients_select_clause(admin_user_id):
    unread_case = '0'
    if admin_user_id is not None:
        unread_case = f'''(
            SELECT COUNT(*)
            FROM messages m
            JOIN users pu ON pu.patient_id = p.id AND pu.role = 'patient'
            WHERE m.sender_id = pu.id
              AND m.recipient_id = {int(admin_user_id)}
              AND COALESCE(m.is_read, 0) = 0
        )'''

    return f'''
        SELECT p.*,
        (SELECT COUNT(*) FROM appointments a WHERE a.patient_id = p.id AND a.is_recurring = 1 AND COALESCE(a.status, 'scheduled') = 'scheduled') as has_recurring,
        (
            SELECT MIN(event_date) FROM (
                SELECT MIN(a0.appointment_date) as event_date
                FROM appointments a0
                WHERE a0.patient_id = p.id
                  AND COALESCE(a0.status, 'scheduled') = 'scheduled'
                  AND a0.appointment_date >= DATE('now')
                UNION ALL
                SELECT MIN(gs.session_date)
                FROM group_sessions gs
                JOIN group_members gm ON gm.group_id = gs.group_id AND gm.patient_id = p.id
                WHERE COALESCE(gs.status, 'scheduled') = 'scheduled'
                  AND gs.session_date >= DATE('now')
                  AND date(COALESCE(gm.joined_at, gs.session_date)) <= date(gs.session_date)
                  AND (gm.left_at IS NULL OR date(gm.left_at) >= date(gs.session_date))
            )
        ) AS next_appointment_date,
        (
            SELECT event_time FROM (
                SELECT a1.appointment_date as event_date, a1.appointment_time as event_time
                FROM appointments a1
                WHERE a1.patient_id = p.id
                  AND COALESCE(a1.status, 'scheduled') = 'scheduled'
                  AND a1.appointment_date >= DATE('now')
                UNION ALL
                SELECT gs2.session_date, gs2.session_time
                FROM group_sessions gs2
                JOIN group_members gm2 ON gm2.group_id = gs2.group_id AND gm2.patient_id = p.id
                WHERE COALESCE(gs2.status, 'scheduled') = 'scheduled'
                  AND gs2.session_date >= DATE('now')
                  AND date(COALESCE(gm2.joined_at, gs2.session_date)) <= date(gs2.session_date)
                  AND (gm2.left_at IS NULL OR date(gm2.left_at) >= date(gs2.session_date))
                ORDER BY 1 ASC, 2 ASC
                LIMIT 1
            )
        ) AS next_appointment_time,
        {unread_case} AS unread_messages,
        (
            CASE
                WHEN p.status = 'candidate' AND EXISTS (
                    SELECT 1 FROM appointments a1
                    WHERE a1.patient_id = p.id
                      AND a1.is_recurring = 0
                      AND COALESCE(a1.status, 'scheduled') = 'scheduled'
                      AND a1.appointment_date < DATE('now')
                ) AND NOT EXISTS (
                    SELECT 1 FROM appointments a2
                    WHERE a2.patient_id = p.id
                      AND COALESCE(a2.status, 'scheduled') = 'scheduled'
                      AND a2.appointment_date >= DATE('now')
                )
                THEN 1
                ELSE 0
            END
        ) AS needs_followup_decision,
        (
            SELECT GROUP_CONCAT(g.name, ', ')
            FROM group_members gm
            JOIN groups g ON g.id = gm.group_id
            WHERE gm.patient_id = p.id
              AND gm.left_at IS NULL
              AND COALESCE(g.is_active, 1) = 1
        ) AS group_names,
        (
            SELECT GROUP_CONCAT(COALESCE(gm.role, 'member'), ', ')
            FROM group_members gm
            JOIN groups g ON g.id = gm.group_id
            WHERE gm.patient_id = p.id
              AND gm.left_at IS NULL
              AND COALESCE(g.is_active, 1) = 1
        ) AS group_roles
        FROM patients p
        WHERE COALESCE(p.is_deleted, 0) = 0
    '''

def _get_patients_where_clause(status, patient_type, search_query, include_group, treatment_method, show_archived=False):
    where_query = ""
    params = []

    if status in LEGACY_WAITING_STATUSES:
        where_query += " AND p.status IN ('candidate', 'waiting for scheduling', 'waiting')"
    elif status == 'all':
        if not show_archived:
            where_query += " AND COALESCE(p.status, 'candidate') != 'archived'"
    else:
        where_query += ' AND p.status = ?'
        params.append(status)

    if patient_type in ('private', 'residency', 'initial-intake', 'diagnosee', 'group'):
        where_query += ' AND COALESCE(p.patient_type, "private") = ?'
        params.append(patient_type)
    elif not include_group:
        where_query += ' AND COALESCE(p.patient_type, "private") != "group"'

    if search_query:
        where_query += ' AND (LOWER(p.name) LIKE ? OR LOWER(COALESCE(p.email, "")) LIKE ? OR LOWER(COALESCE(p.phone, "")) LIKE ?)'
        like_value = f"%{search_query.lower()}%"
        params.extend([like_value, like_value, like_value])

    if treatment_method and treatment_method != 'all':
        where_query += ' AND COALESCE(p.treatment_method, "") = ?'
        params.append(treatment_method)

    return where_query, params

def _get_patients_order_clause(sort_by):
    order_map = {
        'name_asc': 'p.name ASC',
        'name_desc': 'p.name DESC',
        'newest': 'p.created_at DESC',
        'oldest': 'p.created_at ASC',
        'manual_order': 'COALESCE(p.sort_order, 999999) ASC, p.created_at DESC',
        'status_priority': '''
            CASE
                WHEN p.status = 'ongoing' THEN 0
                WHEN p.status IN ('candidate', 'waiting for scheduling', 'waiting') THEN 1
                WHEN p.status = 'archived' THEN 2
                ELSE 3
            END ASC,
            p.created_at DESC
        '''
    }
    return " ORDER BY " + order_map.get(sort_by, order_map['status_priority'])


def _normalize_patient_status(status):
    normalized = (status or 'candidate').strip().lower()
    if normalized in LEGACY_WAITING_STATUSES:
        return 'candidate'
    if normalized in {'ongoing', 'archived'}:
        return normalized
    return 'candidate'


def fetch_patients_by_status(db, status, patient_type='all', search_query='', sort_by='status_priority', admin_user_id=None, include_group=True, treatment_method='all', show_archived=False):
    select_clause = _get_patients_select_clause(admin_user_id)
    where_clause, params = _get_patients_where_clause(status, patient_type, search_query, include_group, treatment_method, show_archived)
    order_clause = _get_patients_order_clause(sort_by)

    final_query = f"{select_clause}{where_clause}{order_clause}"
    return db.execute(final_query, tuple(params)).fetchall()


@app.route('/crm')
@login_required
def crm_dashboard():
    if current_user.role != 'admin':
        return redirect(url_for('patient_home'))

    db = get_db()
    saved_filters = session.get('crm_filters', {})
    status = request.args.get('status', saved_filters.get('status', 'all')).strip()
    clinic_type = request.args.get('clinic_type', saved_filters.get('clinic_type', 'all')).strip()
    search_query = request.args.get('q', saved_filters.get('q', '')).strip()
    sort_by = request.args.get('sort', saved_filters.get('sort', 'status_priority')).strip()
    include_group_raw = request.args.get('include_group', saved_filters.get('include_group', 'false'))
    include_group = include_group_raw == 'true'
    treatment_method = request.args.get('treatment_method', saved_filters.get('treatment_method', 'all')).strip()
    show_deleted_raw = request.args.get('show_deleted', saved_filters.get('show_deleted', 'false'))
    show_deleted = show_deleted_raw == 'true'
    show_archived_raw = request.args.get('show_archived', saved_filters.get('show_archived', 'false'))
    show_archived = show_archived_raw == 'true'

    status = _normalize_patient_status(status) if status != 'all' else 'all'
    if status not in {'all', 'ongoing', 'candidate', 'archived'}:
        status = 'all'
    if clinic_type not in {'all', 'private', 'residency', 'group'}:
        clinic_type = 'all'
    if sort_by not in {'status_priority', 'name_asc', 'name_desc', 'newest', 'oldest', 'manual_order'}:
        sort_by = 'status_priority'

    session['crm_filters'] = {
        'status': status,
        'clinic_type': clinic_type,
        'q': search_query,
        'sort': sort_by,
        'include_group': 'true' if include_group else 'false',
        'treatment_method': treatment_method,
        'show_deleted': 'true' if show_deleted else 'false',
        'show_archived': 'true' if show_archived else 'false'
    }
    
    patient_type = clinic_type

    treatment_method_options = db.execute(
        'SELECT label FROM treatment_method_options ORDER BY display_order ASC, label ASC'
    ).fetchall()
    treatment_method_labels = [row['label'] for row in treatment_method_options]

    # Get patients - if showing deleted, only show deleted; otherwise show active patients
    if show_deleted:
        # Show only deleted patients with optional search and filtering
        select_clause = '''
            SELECT p.*,
            (SELECT COUNT(*) FROM appointments a WHERE a.patient_id = p.id AND a.is_recurring = 1 AND COALESCE(a.status, 'scheduled') = 'scheduled') as has_recurring,
            0 AS unread_messages,
            0 AS needs_followup_decision,
            (
                SELECT GROUP_CONCAT(g.name, ', ')
                FROM group_members gm
                JOIN groups g ON g.id = gm.group_id
                WHERE gm.patient_id = p.id
                  AND gm.left_at IS NULL
                  AND COALESCE(g.is_active, 1) = 1
            ) AS group_names,
            (
                SELECT GROUP_CONCAT(COALESCE(gm.role, 'member'), ', ')
                FROM group_members gm
                JOIN groups g ON g.id = gm.group_id
                WHERE gm.patient_id = p.id
                  AND gm.left_at IS NULL
                  AND COALESCE(g.is_active, 1) = 1
            ) AS group_roles
            FROM patients p
            WHERE COALESCE(p.is_deleted, 0) = 1
        '''
        where_clause = ""
        params = []
        
        if search_query:
            where_clause += ' AND (LOWER(p.name) LIKE ? OR LOWER(COALESCE(p.email, "")) LIKE ? OR LOWER(COALESCE(p.phone, "")) LIKE ?)'
            like_value = f"%{search_query.lower()}%"
            params.extend([like_value, like_value, like_value])

        if clinic_type != 'all':
            where_clause += ' AND p.patient_type = ?'
            params.append(clinic_type)
        elif not include_group:
            where_clause += ' AND COALESCE(p.patient_type, "private") <> ?'
            params.append('group')

        if treatment_method and treatment_method != 'all':
            where_clause += ' AND COALESCE(p.treatment_method, "") = ?'
            params.append(treatment_method)
        
        order_clause = ' ORDER BY p.deleted_at DESC, p.name ASC'
        final_query = f"{select_clause}{where_clause}{order_clause}"
        patients = db.execute(final_query, tuple(params)).fetchall()
    else:
        patients = fetch_patients_by_status(db, status, patient_type=patient_type, search_query=search_query, sort_by=sort_by, admin_user_id=current_user.id, include_group=include_group, treatment_method=treatment_method, show_archived=show_archived)
    
    count_where = ['COALESCE(is_deleted, 0) = 0']
    count_params = []
    if clinic_type != 'all':
        count_where.append('patient_type = ?')
        count_params.append(clinic_type)
    elif not include_group:
        count_where.append('COALESCE(patient_type, "private") <> ?')
        count_params.append('group')
    if treatment_method and treatment_method != 'all':
        count_where.append('COALESCE(treatment_method, "") = ?')
        count_params.append(treatment_method)

    counts_row = db.execute(f'''
        SELECT
            COUNT(*) AS all_count,
            SUM(CASE WHEN status = 'ongoing' THEN 1 ELSE 0 END) AS ongoing_count,
            SUM(CASE WHEN status IN ('candidate', 'waiting for scheduling', 'waiting') THEN 1 ELSE 0 END) AS candidate_waiting_count,
            SUM(CASE WHEN status = 'archived' THEN 1 ELSE 0 END) AS archived_count
        FROM patients
        WHERE {' AND '.join(count_where)}
    ''', tuple(count_params)).fetchone()
    
    # Get deleted patients count
    deleted_count_where = ['COALESCE(is_deleted, 0) = 1']
    deleted_count_params = []
    if clinic_type != 'all':
        deleted_count_where.append('patient_type = ?')
        deleted_count_params.append(clinic_type)
    elif not include_group:
        deleted_count_where.append('COALESCE(patient_type, "private") <> ?')
        deleted_count_params.append('group')
    if treatment_method and treatment_method != 'all':
        deleted_count_where.append('COALESCE(treatment_method, "") = ?')
        deleted_count_params.append(treatment_method)

    deleted_count_row = db.execute(
        f'''SELECT COUNT(*) AS deleted_count FROM patients WHERE {' AND '.join(deleted_count_where)}''',
        tuple(deleted_count_params)
    ).fetchone()
    
    waiting_count = counts_row['candidate_waiting_count'] or 0
    counts = {
        'all': counts_row['all_count'] or 0,
        'ongoing': counts_row['ongoing_count'] or 0,
        'waiting': waiting_count,
        'candidate': waiting_count,
        'candidate_waiting': waiting_count,
        'archived': counts_row['archived_count'] or 0,
        'deleted': deleted_count_row['deleted_count'] or 0
    }
    
    # Get today's appointments (include tomorrow for UTC timezone offset)
    today = datetime.now().date()
    tomorrow = today + timedelta(days=1)
    today_appointments = db.execute('''
        SELECT a.id FROM appointments a
        JOIN patients p ON p.id = a.patient_id
        WHERE COALESCE(a.status, 'scheduled') = 'scheduled'
          AND COALESCE(p.is_deleted, 0) = 0
          AND a.appointment_date IN (?, ?)
    ''', (today.isoformat(), tomorrow.isoformat())).fetchall()
    
    # Get new admissions this week
    week_ago = today - timedelta(days=7)
    new_this_week = db.execute('''
        SELECT COUNT(*) as count FROM patients
        WHERE COALESCE(is_deleted, 0) = 0
          AND created_at >= ?
    ''', (week_ago.isoformat(),)).fetchone()['count']
    counts['new_this_week'] = new_this_week or 0
    
    # Calculate average wait time (placeholder for now)
    avg_wait_time = '18 min'

    # For patients with recurring appointments whose base date is in the past,
    # compute the actual next occurrence date dynamically.
    if not show_deleted:
        recurrence_range_end = today + timedelta(days=365)
        patients_needing_check = [p for p in patients
                                   if dict(p).get('has_recurring') and (not dict(p).get('next_appointment_date') or not dict(p).get('next_appointment_time'))]
        if patients_needing_check:
            pid_list = [p['id'] for p in patients_needing_check]
            placeholders = ','.join('?' * len(pid_list))
            recurring_appts = db.execute(
                f'''SELECT * FROM appointments
                    WHERE patient_id IN ({placeholders})
                      AND COALESCE(is_recurring, 0) = 1
                      AND COALESCE(status, 'scheduled') = 'scheduled'
                    ORDER BY appointment_date ASC''',
                pid_list
            ).fetchall()
            next_occurrences = {}
            for appt in recurring_appts:
                pid = appt['patient_id']
                occurrences = recurring_occurrences_between(appt, today, recurrence_range_end)
                if occurrences:
                    next_occ = occurrences[0]
                    if pid not in next_occurrences or next_occ < next_occurrences[pid][0]:
                        next_occurrences[pid] = (next_occ, appt['appointment_time'])
            patients = [
                dict(p, next_appointment_date=next_occurrences[p['id']][0].isoformat(),
                        next_appointment_time=next_occurrences[p['id']][1])
                if (p['has_recurring'] and (not p['next_appointment_date'] or not p['next_appointment_time']) and p['id'] in next_occurrences)
                else dict(p)
                for p in patients
            ]

        # For patients who still have no next_appointment after checking recurring,
        # check if they have upcoming group sessions.
        patients_missing_next = [p for p in patients
                                  if (not dict(p).get('next_appointment_date') or not dict(p).get('next_appointment_time'))]
        if patients_missing_next:
            pid_list = [p['id'] for p in patients_missing_next]
            placeholders = ','.join('?' * len(pid_list))
            group_next = db.execute(
                f'''SELECT gm.patient_id, MIN(gs.session_date) as next_date, gs.session_time
                    FROM group_sessions gs
                    JOIN group_members gm ON gm.group_id = gs.group_id
                    WHERE gm.patient_id IN ({placeholders})
                      AND COALESCE(gs.status, 'scheduled') = 'scheduled'
                      AND gs.session_date >= DATE('now')
                      AND date(COALESCE(gm.joined_at, gs.session_date)) <= date(gs.session_date)
                      AND (gm.left_at IS NULL OR date(gm.left_at) >= date(gs.session_date))
                    GROUP BY gm.patient_id
                    ORDER BY gs.session_date ASC
                ''',
                pid_list
            ).fetchall()
            group_map = {}
            for row in group_next:
                pid = row['patient_id']
                if pid not in group_map or row['next_date'] < group_map[pid][0]:
                    group_map[pid] = (row['next_date'], row['session_time'])
            patients = [
                dict(p, next_appointment_date=group_map[p['id']][0],
                        next_appointment_time=group_map[p['id']][1])
                if p['id'] in group_map
                else dict(p)
                for p in patients
            ]

    reminders = send_appointment_reminders(db)
    return render_template('crm.html', patients=patients, status=status, counts=counts,
                           clinic_type=clinic_type, search_query=search_query, sort_by=sort_by,
                           include_group=include_group, reminders=reminders,
                           treatment_method=treatment_method,
                           treatment_method_options=treatment_method_labels,
                           today_appointments=today_appointments,
                           avg_wait_time=avg_wait_time,
                           show_deleted=show_deleted,
                           show_archived=show_archived)



@app.route('/api/patients/reorder', methods=['POST'])
@login_required
def api_patients_reorder():
    data = request.json
    if not data or 'order' not in data:
        return jsonify({'error': 'No order provided'}), 400
    db = get_db()
    update_data = []
    for idx, patient_id in enumerate(data['order']):
        if not isinstance(patient_id, int):
            return jsonify({'error': 'Invalid patient id'}), 400
        update_data.append((idx, patient_id))
    db.executemany('UPDATE patients SET sort_order = ? WHERE id = ? AND COALESCE(is_deleted,0) = 0', update_data)

    db.commit()
    return jsonify({'success': True})


@app.route('/api/treatment_method_options', methods=['GET'])
@login_required
def api_treatment_method_options_get():
    if current_user.role != 'admin':
        return jsonify({'error': 'Unauthorized'}), 403
    db = get_db()
    rows = db.execute('SELECT id, label FROM treatment_method_options ORDER BY display_order ASC, label ASC').fetchall()
    return jsonify([{'id': r['id'], 'label': r['label']} for r in rows])


@app.route('/api/treatment_method_options', methods=['POST'])
@login_required
def api_treatment_method_options_add():
    if current_user.role != 'admin':
        return jsonify({'error': 'Unauthorized'}), 403
    data = request.get_json(silent=True)
    label = (data or {}).get('label', '').strip()
    if not label:
        return jsonify({'error': 'Label is required'}), 400
    if len(label) > 80:
        return jsonify({'error': 'Label too long'}), 400
    db = get_db()
    try:
        db.execute('INSERT INTO treatment_method_options (label) VALUES (?)', (label,))
        db.commit()
    except sqlite3.IntegrityError:
        return jsonify({'error': 'Option already exists'}), 409
    return jsonify({'ok': True, 'label': label}), 201


@app.route('/api/treatment_method_options/<int:option_id>', methods=['DELETE'])
@login_required
def api_treatment_method_options_delete(option_id):
    if current_user.role != 'admin':
        return jsonify({'error': 'Unauthorized'}), 403
    db = get_db()
    db.execute('DELETE FROM treatment_method_options WHERE id = ?', (option_id,))
    db.commit()
    return jsonify({'ok': True})


@app.route('/patient/change-password', methods=['GET', 'POST'])
@login_required
def patient_change_password():
    if current_user.role != 'patient':
        return redirect(url_for('patient_home'))

    db = get_db()
    user = db.execute('SELECT * FROM users WHERE id = ?', (current_user.id,)).fetchone()

    if not user or not user['force_password_change']:
        return redirect(url_for('patient_home'))

    if request.method == 'POST':
        new_password = (request.form.get('new_password') or '').strip()
        confirm_password = (request.form.get('confirm_password') or '').strip()

        ok, err = _validate_password_strength(new_password, username=user['username'])
        if not ok:
            flash(err)
            return render_template('patient_change_password.html')

        if new_password != confirm_password:
            flash('New password confirmation does not match.')
            return render_template('patient_change_password.html')

        db.execute(
            'UPDATE users SET password_hash = ?, force_password_change = 0, session_version = COALESCE(session_version, 0) + 1 WHERE id = ?',
            (generate_password_hash(new_password), current_user.id)
        )
        db.commit()
        flash('Password updated successfully.')
        return redirect(url_for('patient_home'))

    return render_template('patient_change_password.html')


@app.route('/patient/settings')
@login_required
def patient_settings():
    if current_user.role != 'patient':
        return redirect(url_for('patient_home'))

    db = get_db()
    user = db.execute('SELECT * FROM users WHERE id = ?', (current_user.id,)).fetchone()
    if not user:
        return "User not found", 404

    from urllib.parse import quote
    import pyotp
    pending_secret = session.get('pending_totp_secret')
    totp_uri = None
    if pending_secret:
        totp_uri = pyotp.totp.TOTP(pending_secret).provisioning_uri(name=user['username'], issuer_name='Private Clinic')

    recovery_codes = session.pop('mfa_recovery_codes', None)

    return render_template(
        'patient_settings.html',
        user=user,
        pending_totp_secret=pending_secret,
        totp_qr_url=f"https://api.qrserver.com/v1/create-qr-code/?size=220x220&data={quote(totp_uri)}" if totp_uri else None,
        recovery_codes=recovery_codes
    )



@app.route('/patient/home')
@login_required
def patient_home():
    if current_user.role != 'patient':
        return redirect(url_for('patients'))

    db = get_db()
    patient_id = current_user.patient_id
    patient = db.execute('SELECT * FROM patients WHERE id = ?', (patient_id,)).fetchone()

    upcoming = build_patient_upcoming_events(db, patient_id, days_ahead=120, limit=10)

    messages = db.execute('''
        SELECT m.*, u.username as sender_name
        FROM messages m
        LEFT JOIN users u ON m.sender_id = u.id
        WHERE m.sender_id = ? OR m.recipient_id = ?
        ORDER BY m.timestamp ASC
        LIMIT 20
    ''', (current_user.id, current_user.id)).fetchall()

    assigned_resources = db.execute('''
        SELECT r.*
        FROM resources r
        JOIN patient_resources pr ON r.id = pr.resource_id
        WHERE pr.patient_id = ?
          AND COALESCE(r.allow_patient_view, 1) = 1
        ORDER BY pr.assigned_at DESC
    ''', (patient_id,)).fetchall()

    receipts = db.execute('''
        SELECT *
        FROM receipts
        WHERE patient_id = ?
        ORDER BY created_at DESC
    ''', (patient_id,)).fetchall()

    past_appointments = db.execute('''
        SELECT a.id, a.appointment_date, a.appointment_time, a.duration_minutes,
               a.meeting_type, a.meeting_link, a.meeting_title, a.status, a.is_recurring,
               a.missed_reason
        FROM appointments a
        WHERE a.patient_id = ?
          AND a.appointment_date < ?
        ORDER BY a.appointment_date DESC, a.appointment_time DESC
        LIMIT 20
    ''', (patient_id, datetime.now().date().isoformat())).fetchall()

    shared_notes = db.execute('''
        SELECT id, note_date, content, session_number, key_topics, created_at
        FROM notes
        WHERE patient_id = ?
          AND share_with_patient = 1
        ORDER BY created_at DESC
        LIMIT 10
    ''', (patient_id,)).fetchall()

    goals = db.execute(
        'SELECT * FROM goals WHERE patient_id = ? ORDER BY status ASC, created_at DESC',
        (patient_id,)
    ).fetchall()

    notes_for_chart = db.execute(
        'SELECT note_date, key_topics, created_at FROM notes WHERE patient_id = ? ORDER BY created_at DESC LIMIT 200',
        (patient_id,)
    ).fetchall()
    chart_data = _build_patient_chart_data(notes_for_chart)
    chart_data['session_labels'] = json.loads(chart_data['session_labels'])
    chart_data['session_data'] = json.loads(chart_data['session_data'])

    db.execute('UPDATE messages SET is_read = 1 WHERE recipient_id = ?', (current_user.id,))
    db.commit()

    return render_template('patient_home.html', patient=patient,
                           upcoming=upcoming, messages=messages,
                           assigned_resources=assigned_resources,
                           receipts=receipts,
                           past_appointments=past_appointments,
                           shared_notes=shared_notes,
                           goals=goals,
                           chart_data=chart_data)


@app.route('/dashboard')
@login_required
def patient_dashboard():
    """Enhanced patient engagement dashboard with stats and insights"""
    db = get_db()
    
    if current_user.role == 'admin':
        return redirect(url_for('patients'))
    
    patient_id = current_user.patient_id
    patient = db.execute(
        'SELECT * FROM patients WHERE id = ?', 
        (patient_id,)
    ).fetchone()
    
    if not patient:
        return redirect(url_for('patient_home'))
    
    # Get upcoming appointments
    today = datetime.now().date()
    upcoming_appointments = db.execute('''
        SELECT * FROM appointments
        WHERE patient_id = ?
        AND appointment_date >= ?
        AND COALESCE(status, 'scheduled') = 'scheduled'
        ORDER BY appointment_date ASC, appointment_time ASC
        LIMIT 5
    ''', (patient_id, today.isoformat())).fetchall()
    
    # Get total appointment count
    total_appointments = db.execute(
        'SELECT COUNT(*) as count FROM appointments WHERE patient_id = ?',
        (patient_id,)
    ).fetchone()['count']
    
    # Get notes/progress
    recent_notes = db.execute('''
        SELECT * FROM notes
        WHERE patient_id = ?
        ORDER BY created_at DESC
        LIMIT 3
    ''', (patient_id,)).fetchall()
    
    # Get therapy goals
    goals = db.execute('''
        SELECT * FROM goals
        WHERE patient_id = ?
        AND status = 'active'
        ORDER BY created_at DESC
    ''', (patient_id,)).fetchall()
    
    # Calculate engagement metrics
    days_since_last_session = None
    if total_appointments > 0:
        last_appointment = db.execute('''
            SELECT appointment_date FROM appointments
            WHERE patient_id = ?
            ORDER BY appointment_date DESC
            LIMIT 1
        ''', (patient_id,)).fetchone()
        
        if last_appointment:
            last_date = datetime.fromisoformat(last_appointment['appointment_date']).date()
            days_since_last_session = (today - last_date).days
    
    # Get zoom/online meetings count
    zoom_meetings = db.execute('''
        SELECT COUNT(*) as count FROM appointments
        WHERE patient_id = ?
        AND meeting_type IN ('zoom', 'google-meet', 'online')
    ''', (patient_id,)).fetchone()['count']
    
    engagement_data = {
        'total_appointments': total_appointments,
        'upcoming_appointments': len(upcoming_appointments),
        'days_since_last': days_since_last_session,
        'zoom_meetings': zoom_meetings,
        'active_goals': len(goals),
        'recent_notes': len(recent_notes)
    }
    
    return render_template('patient_dashboard.html',
        patient=patient,
        upcoming_appointments=upcoming_appointments,
        recent_notes=recent_notes,
        goals=goals,
        engagement=engagement_data,
        now=datetime.now()
    )


@app.route('/api/engagement/stats')
@login_required
def api_engagement_stats():
    """Get engagement statistics for the patient"""
    db = get_db()
    
    total_appts = db.execute(
        'SELECT COUNT(*) as count FROM appointments WHERE patient_id = ?',
        (current_user.patient_id,)
    ).fetchone()['count']
    
    completed_appts = db.execute('''
        SELECT COUNT(*) as count FROM appointments
        WHERE patient_id = ?
        AND COALESCE(status, 'scheduled') = 'completed'
    ''', (current_user.patient_id,)).fetchone()['count']
    
    this_month_appts = db.execute('''
        SELECT COUNT(*) as count FROM appointments
        WHERE patient_id = ?
        AND strftime('%Y-%m', appointment_date) = strftime('%Y-%m', 'now')
    ''', (current_user.patient_id,)).fetchone()['count']
    
    online_appts = db.execute('''
        SELECT COUNT(*) as count FROM appointments
        WHERE patient_id = ?
        AND meeting_type IN ('zoom', 'google-meet', 'online')
    ''', (current_user.patient_id,)).fetchone()['count']
    
    return jsonify({
        'total_appointments': total_appts,
        'completed_appointments': completed_appts,
        'appointments_this_month': this_month_appts,
        'online_appointments': online_appts,
        'completion_rate': round((completed_appts / max(total_appts, 1)) * 100) if total_appts > 0 else 0
    })


@app.route('/patient/appointment/<int:appointment_id>/request_cancel', methods=['POST'])
@login_required
def request_cancel_appointment(appointment_id):
    if current_user.role != 'patient':
        return 'Unauthorized', 403

    db = get_db()
    from clinic_app.utils import _check_db_rate_limit, _record_db_rate_limit
    bucket_key = f"cancel-appt-{current_user.id}"
    retry_after = _check_db_rate_limit(db, bucket_key, 'cancel', 5, 3600) # Max 5 cancel requests per hour
    if retry_after:
        flash(f'Too many cancellation requests. Please wait {retry_after} seconds.')
        return redirect(url_for('patient_home'))
    _record_db_rate_limit(db, bucket_key, 'cancel')

    reason = (request.form.get('reason') or '').strip()
    if not reason:
        flash('Please explain why you want to cancel.')
        return redirect(url_for('patient_home'))

    appointment = db.execute('''
        SELECT a.id, a.patient_id, a.appointment_date, a.appointment_time, a.meeting_type, p.name AS patient_name
        FROM appointments a
        JOIN patients p ON p.id = a.patient_id
        WHERE a.id = ? AND a.patient_id = ?
    ''', (appointment_id, current_user.patient_id)).fetchone()
    if not appointment:
        return 'Appointment not found', 404

    appointment_dt = datetime.combine(
        parse_date_safe(appointment['appointment_date']),
        parse_time_safe(appointment['appointment_time'])
    )
    lead_time = format_lead_time_for_notice(appointment_dt)
    admin_message = (
        f"System cancellation request from {appointment['patient_name']}: "
        f"appointment on {appointment['appointment_date']} at {appointment['appointment_time']}. "
        f"Time before meeting: {lead_time}. Notes: {reason}"
    )
    patient_ack = (
        f"System: Your cancellation request for {appointment['appointment_date']} at {appointment['appointment_time']} was sent. "
        f"Time before meeting: {lead_time}. Notes: {reason}"
    )

    db.execute(
        '''INSERT INTO cancel_requests (appointment_id, patient_id, reason)
           VALUES (?, ?, ?)''',
        (appointment_id, current_user.patient_id, reason)
    )

    add_patient_chat_request(
        db,
        current_user.id,
        current_user.patient_id,
        admin_message,
        patient_ack,
        audit_action='cancel_request',
        audit_details=admin_message
    )

    try:
        admin_users = db.execute("SELECT email FROM users WHERE role = 'admin' AND email IS NOT NULL AND email != ''").fetchall()
        subject = f'Cancel Request: {appointment["patient_name"]} — {appointment["appointment_date"]} {appointment["appointment_time"]}'
        body = f'A patient has requested to cancel an appointment.\n\nPatient: {appointment["patient_name"]}\nDate: {appointment["appointment_date"]}\nTime: {appointment["appointment_time"]}\nReason: {reason}\n\nLog in to review: {request.host_url}cancel_requests'
        for admin in admin_users:
            _send_smtp_email(admin['email'], subject, body)
    except Exception:
        app.logger.exception('Failed to send cancel request email notification')

    db.commit()
    flash('Cancellation request sent.')
    return redirect(url_for('patient_home'))


@app.route('/patient/request_booking_access', methods=['POST'])
@login_required
def request_booking_access():
    if current_user.role != 'patient':
        return 'Unauthorized', 403

    db = get_db()
    from clinic_app.utils import _check_db_rate_limit, _record_db_rate_limit
    bucket_key = f"booking-access-{current_user.id}"
    retry_after = _check_db_rate_limit(db, bucket_key, 'booking', 5, 3600) # Max 5 booking requests per hour
    if retry_after:
        flash(f'Too many booking requests. Please wait {retry_after} seconds.')
        return redirect(url_for('patient_home'))
    _record_db_rate_limit(db, bucket_key, 'booking')

    notes = (request.form.get('notes') or '').strip()
    if not notes:
        flash('Please add a note for your booking request.')
        return redirect(url_for('patient_home'))

    patient = db.execute('SELECT id, name, can_self_schedule FROM patients WHERE id = ?', (current_user.patient_id,)).fetchone()
    if not patient:
        return 'Patient not found', 404

    next_appointment = db.execute('''
        SELECT appointment_date, appointment_time
        FROM appointments
        WHERE patient_id = ? AND COALESCE(status, 'scheduled') = 'scheduled' AND datetime(appointment_date || ' ' || appointment_time) >= datetime('now')
        ORDER BY appointment_date ASC, appointment_time ASC
        LIMIT 1
    ''', (current_user.patient_id,)).fetchone()
    next_fragment = ''
    if next_appointment:
        next_fragment = f" Current scheduled meeting: {next_appointment['appointment_date']} at {next_appointment['appointment_time']}."

    admin_message = (
        f"System booking request from {patient['name']}: patient asked to open self-booking for another meeting from available slots."
        f"{next_fragment} Notes: {notes}"
    )
    patient_ack = (
        'System: Your request for another meeting was sent to the clinic. '
        'If approved, self-booking can be opened for you from the available slots. '
        f'Notes: {notes}'
    )

    add_patient_chat_request(
        db,
        current_user.id,
        current_user.patient_id,
        admin_message,
        patient_ack,
        audit_action='booking_access_request',
        audit_details=admin_message
    )
    db.commit()
    flash('Booking request sent.')
    return redirect(url_for('patient_home'))


@app.route('/resources')
def public_resources():
    db = get_db()
    resources = db.execute('''
        SELECT *
        FROM resources
        WHERE is_public = 1 AND COALESCE(allow_patient_view, 1) = 1
        ORDER BY created_at DESC
    ''').fetchall()
    return render_template('resources.html', resources=resources, is_admin=False)


def _can_access_resource(db, resource, action='view'):
    if resource is None:
        return False

    if current_user.is_authenticated and current_user.role == 'admin':
        return True

    allow_view = int(resource['allow_patient_view'] or 0) == 1
    allow_download = int(resource['allow_patient_download'] or 0) == 1

    if action == 'download' and not allow_download:
        return False
    if not allow_view:
        return False

    if int(resource['is_public'] or 0) == 1:
        return True

    if current_user.is_authenticated and current_user.role == 'patient' and current_user.patient_id:
        assigned = db.execute(
            'SELECT 1 FROM patient_resources WHERE patient_id = ? AND resource_id = ?',
            (current_user.patient_id, resource['id'])
        ).fetchone()
        return assigned is not None

    return False


@app.route('/resource/<int:resource_id>/open')
def open_resource_link(resource_id):
    db = get_db()
    resource = db.execute('SELECT * FROM resources WHERE id = ?', (resource_id,)).fetchone()
    if resource is None:
        return 'Resource not found', 404
    if not _can_access_resource(db, resource, action='view'):
        return 'Access denied', 403
    if not resource['url']:
        return 'Resource URL not found', 404
    return redirect(resource['url'])


@app.route('/resource/<int:resource_id>/download')
def download_resource_link(resource_id):
    db = get_db()
    resource = db.execute('SELECT * FROM resources WHERE id = ?', (resource_id,)).fetchone()
    if resource is None:
        return 'Resource not found', 404
    if not _can_access_resource(db, resource, action='download'):
        return 'Download not allowed', 403
    if not resource['url']:
        return 'Resource URL not found', 404
    return redirect(resource['url'])

@app.route('/patients')
@login_required
def patients():
    if current_user.role != 'admin':
        return redirect(url_for('patient_home'))
    status = request.args.get('status', 'all')
    show_deleted = request.args.get('show_deleted', '0') == '1'
    return redirect(url_for('crm_dashboard', status=status, show_deleted='1' if show_deleted else None))

@app.route('/add_patient', methods=('GET', 'POST'))
@login_required
def add_patient():
    if current_user.role != 'admin':
        flash('Access denied.')
        return redirect(url_for('patient_home'))

    if request.method == 'POST':
        name = request.form['name']
        status = _normalize_patient_status(request.form['status'])
        email = request.form.get('email')
        phone = request.form.get('phone')
        birth_date = request.form.get('birth_date') or None
        id_number = (request.form.get('id_number') or '').strip() or None
        patient_type = request.form.get('patient_type', 'private')
        if patient_type not in ('private', 'residency', 'initial-intake', 'diagnosee', 'group'):
            patient_type = 'private'
        selected_questionnaires = [
            item.strip() for item in request.form.getlist('diagnosee_questionnaires') if item and item.strip()
        ] if patient_type == 'diagnosee' else []
        has_intake_tab = 1 if patient_type in ('initial-intake', 'diagnosee') else 0
        has_questionnaire_tab = 1 if patient_type == 'diagnosee' else 0
        intake_assessment = request.form.get('intake_assessment', '').strip() if patient_type in ('initial-intake', 'diagnosee') else ''
        intake_questionnaire = request.form.get('intake_questionnaire', '').strip() if patient_type in ('initial-intake', 'diagnosee') else ''
        treatment_method = request.form.get('treatment_method', '').strip() or None

        if not name:
            flash('Name is required!')
        else:
            field_errors = _validate_patient_fields(name, phone=phone, birth_date=birth_date, email=email)
            for err in field_errors:
                flash(err)
            if not field_errors:
                db = get_db()
                cursor = db.execute('''INSERT INTO patients
                                      (name, status, email, phone, birth_date, id_number, patient_type, has_intake_tab, has_questionnaire_tab, intake_assessment, intake_questionnaire, treatment_method)
                                      VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                                  (name, status, email, phone, birth_date, id_number, patient_type, has_intake_tab,
                                   has_questionnaire_tab, intake_assessment or None, intake_questionnaire or None, treatment_method))

                created_patient_id = int(cursor.lastrowid)
                if patient_type == 'diagnosee' and selected_questionnaires:
                    result, create_err = _create_diagnosee_questionnaires_sheet(db, name, selected_questionnaires)
                    if create_err:
                        db.rollback()
                        flash(f'Failed to create diagnosee questionnaires file: {create_err}')
                        tabs, _ = _list_questionnaire_tabs(db)
                        treatment_method_options = [r['label'] for r in db.execute(
                            'SELECT label FROM treatment_method_options ORDER BY display_order ASC, label ASC'
                        ).fetchall()]
                        return render_template(
                            'add_patient.html',
                            treatment_method_options=treatment_method_options,
                            questionnaire_options=[item['title'] for item in tabs],
                        )

                    db.execute('''
                        UPDATE patients
                        SET questionnaires_file_id = ?, questionnaires_file_url = ?, questionnaires_selected = ?
                        WHERE id = ?
                    ''', (
                        result['spreadsheet_id'],
                        result['spreadsheet_url'],
                        json.dumps(result['selected_titles'], ensure_ascii=False),
                        created_patient_id,
                    ))

                db.commit()
                return redirect(url_for('patients', status=status))

    db = get_db()
    treatment_method_options = [r['label'] for r in db.execute(
        'SELECT label FROM treatment_method_options ORDER BY display_order ASC, label ASC'
    ).fetchall()]
    questionnaire_tabs, _ = _list_questionnaire_tabs(db)
    questionnaire_options = [item['title'] for item in questionnaire_tabs]
    return render_template(
        'add_patient.html',
        treatment_method_options=treatment_method_options,
        questionnaire_options=questionnaire_options,
    )

def _get_patient_notes(db, patient_id):
    notes = db.execute('''
        SELECT * FROM notes
        WHERE patient_id = ?
        ORDER BY COALESCE(note_date, date(created_at)) DESC,
                 datetime(created_at) DESC,
                 id DESC
        LIMIT 200
    ''', (patient_id,)).fetchall()
    return notes

def _get_patient_group_data(db, patient_id):
    group_attendance_rows = db.execute('''
        SELECT gsa.session_id,
               gsa.attendance_status,
               gsa.absence_reason,
               gsa.notified_on_time,
               gsa.attendance_note,
               gsa.updated_at,
               gs.group_id,
               gs.session_date,
               gs.session_time,
               gs.title AS session_title,
               gs.session_summary,
               g.name AS group_name
        FROM group_session_attendance gsa
        JOIN group_sessions gs ON gs.id = gsa.session_id
        JOIN groups g ON g.id = gs.group_id
        WHERE gsa.patient_id = ?
        ORDER BY gs.session_date DESC, gs.session_time DESC
    ''', (patient_id,)).fetchall()

    group_membership_rows = db.execute('''
        SELECT h.id,
               h.group_id,
               g.name AS group_name,
               h.joined_at,
               h.left_at,
               COALESCE(h.role, 'member') AS role
        FROM group_member_history h
        JOIN groups g ON g.id = h.group_id
        WHERE h.patient_id = ?
        ORDER BY h.joined_at DESC
    ''', (patient_id,)).fetchall()

    group_arrived_count = sum(1 for row in group_attendance_rows if (row['attendance_status'] or '') == 'present')
    return group_attendance_rows, group_membership_rows, group_arrived_count

def _get_patient_messages(db, user, current_user_id):
    messages = []
    unread_messages_count = 0
    if user:
        unread_messages_count = db.execute('''
            SELECT COUNT(*) AS c
            FROM messages
            WHERE sender_id = ? AND recipient_id = ? AND COALESCE(is_read, 0) = 0
        ''', (user['id'], current_user_id)).fetchone()['c']
        messages = db.execute('''
            SELECT m.*, u.username as sender_name
            FROM messages m
            LEFT JOIN users u ON m.sender_id = u.id
            WHERE (m.sender_id = ? AND m.recipient_id = ?)
               OR (m.sender_id = ? AND m.recipient_id = ?)
            ORDER BY timestamp ASC
        ''', (current_user_id, user['id'], user['id'], current_user_id)).fetchall()
    return messages, unread_messages_count

def _get_patient_behavior_info(notes):
    behavior_options = [
        'Calm', 'Anxious', 'Restless', 'Withdrawn', 'Cooperative', 'Engaged', 'Low Energy', 'Irritable'
    ]
    latest_behavior = {
        'patient_appearance': '',
        'behavior_checklist': set(),
        'mood_summary': '',
        'behavior_notes': ''
    }
    if notes:
        latest_behavior['patient_appearance'] = notes[0]['patient_appearance'] or ''
        latest_behavior['mood_summary'] = notes[0]['mood_summary'] or ''
        latest_behavior['behavior_notes'] = notes[0]['behavior_notes'] or ''
        checklist_raw = notes[0]['behavior_checklist'] or ''
        latest_behavior['behavior_checklist'] = {
            item.strip() for item in checklist_raw.split(',') if item.strip()
        }
    return behavior_options, latest_behavior


def _build_patient_chart_data(notes):
    month_counts = defaultdict(int)
    topic_counts = Counter()

    today = datetime.now().date()
    current_month = today.replace(day=1)
    month_starts = []
    for offset in range(5, -1, -1):
        year = current_month.year
        month = current_month.month - offset
        while month <= 0:
            month += 12
            year -= 1
        month_starts.append(datetime(year, month, 1).date())

    month_labels = [month_start.strftime('%b %Y') for month_start in month_starts]
    month_index = {month_start.strftime('%Y-%m'): idx for idx, month_start in enumerate(month_starts)}
    session_data = [0] * len(month_starts)

    for note in notes:
        note_date_raw = note['note_date'] or note['created_at']
        try:
            if not note_date_raw:
                continue
            note_day = datetime.fromisoformat(str(note_date_raw).replace('Z', '+00:00')).date()
        except (ValueError, TypeError, AttributeError):
            continue

        month_key = note_day.strftime('%Y-%m')
        if month_key in month_index:
            month_counts[month_key] += 1

        raw_topics = (note['key_topics'] or '').strip()
        if raw_topics:
            for topic in raw_topics.split(','):
                normalized = topic.strip()
                if normalized:
                    topic_counts[normalized] += 1

    for month_key, count in month_counts.items():
        if month_key in month_index:
            session_data[month_index[month_key]] = count

    top_topics = topic_counts.most_common(6)
    topic_labels = [topic for topic, _count in top_topics] or ['No topics']
    topic_data = [count for _topic, count in top_topics] or [1]

    return {
        'session_labels': json.dumps(month_labels, ensure_ascii=True),
        'session_data': json.dumps(session_data, ensure_ascii=True),
        'topic_labels': json.dumps(topic_labels, ensure_ascii=True),
        'topic_data': json.dumps(topic_data, ensure_ascii=True),
    }


def _get_patient_followup_status(db, patient_id, next_appointment):
    row = db.execute('''
        SELECT MAX(COALESCE(NULLIF(note_date, ''), substr(created_at, 1, 10))) AS last_note_date
        FROM notes
        WHERE patient_id = ?
    ''', (patient_id,)).fetchone()

    last_note_raw = row['last_note_date'] if row else None
    last_note_date = parse_date_safe(last_note_raw)
    if not last_note_date:
        return {
            'needs_followup': False,
            'has_upcoming': bool(next_appointment),
            'last_note_date': None,
            'days_since_last_note': None,
            'severity': 'warning'
        }

    today = datetime.now().date()
    days_since_last_note = (today - last_note_date).days
    has_upcoming = bool(next_appointment)
    needs_followup = (not has_upcoming) and days_since_last_note >= 30

    return {
        'needs_followup': needs_followup,
        'has_upcoming': has_upcoming,
        'last_note_date': last_note_date.isoformat(),
        'days_since_last_note': days_since_last_note,
        'severity': 'danger' if days_since_last_note >= 60 else 'warning'
    }


@app.route('/patient/<int:patient_id>')
@login_required
def patient_detail(patient_id):
    if current_user.role != 'admin':
        flash('Access denied.')
        return redirect(url_for('patient_home'))

    db = get_db()
    show_deleted = request.args.get('show_deleted', '0') == '1'
    if show_deleted:
        patient = db.execute('SELECT * FROM patients WHERE id = ?', (patient_id,)).fetchone()
    else:
        patient = db.execute('SELECT * FROM patients WHERE id = ? AND COALESCE(is_deleted, 0) = 0', (patient_id,)).fetchone()
    if patient is None:
        return "Patient not found", 404

    patient = dict(patient)
    patient.setdefault('has_questionnaire_tab', 0)
    patient.setdefault('questionnaires_file_id', None)
    patient.setdefault('questionnaires_file_url', None)
    patient.setdefault('questionnaires_selected', None)

    # Fetch user account if exists
    user = db.execute('SELECT * FROM users WHERE patient_id = ?', (patient_id,)).fetchone()

    notes = _get_patient_notes(db, patient_id)
    patient_logs = db.execute(
        '''SELECT * FROM patient_logs WHERE patient_id = ?
           ORDER BY COALESCE(encounter_date, substr(created_at, 1, 10)) DESC, id DESC
           LIMIT 100''',
        (patient_id,)
    ).fetchall()
    files = db.execute('SELECT * FROM files WHERE patient_id = ? ORDER BY created_at DESC LIMIT 100', (patient_id,)).fetchall()
    receipts = db.execute('SELECT * FROM receipts WHERE patient_id = ? ORDER BY created_at DESC LIMIT 100', (patient_id,)).fetchall()
    appointments = db.execute('SELECT * FROM appointments WHERE patient_id = ? ORDER BY appointment_date DESC, appointment_time DESC LIMIT 200', (patient_id,)).fetchall()
    next_items = build_patient_upcoming_events(db, patient_id, days_ahead=120, limit=1)
    next_appointment = next_items[0] if next_items else None
    followup_status = _get_patient_followup_status(db, patient_id, next_appointment)

    group_attendance_rows, group_membership_rows, group_arrived_count = _get_patient_group_data(db, patient_id)

    messages, unread_messages_count = _get_patient_messages(db, user, current_user.id)

    # Get resources for assignment
    all_resources = db.execute('SELECT * FROM resources WHERE is_public = 0 ORDER BY title ASC LIMIT 200').fetchall()

    # Get assigned resources
    assigned_resources = db.execute('''
        SELECT r.*
        FROM resources r
        JOIN patient_resources pr ON r.id = pr.resource_id
        WHERE pr.patient_id = ?
    ''', (patient_id,)).fetchall()

    behavior_options, latest_behavior = _get_patient_behavior_info(notes)
    chart_data = _build_patient_chart_data(notes)

    active_tab = request.args.get('tab', 'info')
    intake_enabled = patient['patient_type'] in ('initial-intake', 'diagnosee') or int(patient['has_intake_tab'] or 0) == 1
    questionnaire_enabled = patient['patient_type'] == 'diagnosee' or int(patient.get('has_questionnaire_tab') or 0) == 1
    if active_tab == 'intake' and not intake_enabled:
        active_tab = 'info'
    if active_tab == 'questionnaires' and not questionnaire_enabled:
        active_tab = 'info'

    if user and active_tab == 'messages' and unread_messages_count:
        db.execute(
            'UPDATE messages SET is_read = 1 WHERE sender_id = ? AND recipient_id = ? AND COALESCE(is_read, 0) = 0',
            (user['id'], current_user.id)
        )
        db.commit()
        unread_messages_count = 0

    latest_note = notes[0] if notes else None
    intake_form_data = parse_intake_questionnaire(patient['intake_questionnaire'], patient['intake_assessment'])
    next_session_row = db.execute('''
        SELECT COALESCE(MAX(CAST(COALESCE(session_number, '0') AS INTEGER)), 0) AS max_session
        FROM notes
        WHERE patient_id = ?
    ''', (patient_id,)).fetchone()
    suggested_session_number = int(next_session_row['max_session'] or 0) + 1
    suggested_note_date = datetime.now().date().isoformat()

    supervisions = db.execute(
        'SELECT * FROM supervisions WHERE patient_id = ? ORDER BY supervision_date DESC, created_at DESC LIMIT 100',
        (patient_id,)
    ).fetchall()

    diagnosis_documents = db.execute(
        'SELECT * FROM diagnosis_documents WHERE patient_id = ? ORDER BY created_at DESC, id DESC LIMIT 100',
        (patient_id,)
    ).fetchall()

    goals = db.execute(
        'SELECT * FROM goals WHERE patient_id = ? ORDER BY created_at ASC LIMIT 50',
        (patient_id,)
    ).fetchall()

    selected_questionnaire_titles = []
    raw_selected_questionnaires = patient.get('questionnaires_selected')
    if raw_selected_questionnaires:
        try:
            parsed_questionnaires = json.loads(raw_selected_questionnaires)
            if isinstance(parsed_questionnaires, list):
                selected_questionnaire_titles = [str(item).strip() for item in parsed_questionnaires if str(item).strip()]
        except Exception:
            selected_questionnaire_titles = []

    source_questionnaire_titles = []
    source_questionnaire_error = None
    source_questionnaire_activation_url = None
    source_questionnaire_sheet_url = ''
    if questionnaire_enabled:
        settings = get_site_settings(db)
        source_questionnaire_sheet_url = (settings.get('questionnaires_source_sheet_url') or '').strip()
        source_tabs, source_questionnaire_error = _list_questionnaire_tabs(db)
        source_questionnaire_activation_url = _extract_google_activation_url(source_questionnaire_error)
        source_questionnaire_titles = [str(item.get('title')).strip() for item in source_tabs if str(item.get('title') or '').strip()]

    available_questionnaire_titles = []
    questionnaire_tabs_error = None
    questionnaire_tabs_activation_url = None
    if questionnaire_enabled:
        linked_sheet_id = _extract_google_sheet_id(patient.get('questionnaires_file_id') or patient.get('questionnaires_file_url'))
        if linked_sheet_id:
            available_questionnaire_titles, questionnaire_tabs_error = _list_spreadsheet_tab_titles(db, linked_sheet_id)
            questionnaire_tabs_activation_url = _extract_google_activation_url(questionnaire_tabs_error)
    if not available_questionnaire_titles and selected_questionnaire_titles:
        available_questionnaire_titles = selected_questionnaire_titles

    service_types = db.execute('SELECT * FROM service_types WHERE is_active = 1 ORDER BY name ASC').fetchall()

    return render_template('patient_detail.html', patient=patient, notes=notes, patient_logs=patient_logs, files=files, receipts=receipts, user=user, appointments=appointments, next_appointment=next_appointment, followup_status=followup_status, messages=messages, all_resources=all_resources, assigned_resources=assigned_resources, active_tab=active_tab, behavior_options=behavior_options, latest_behavior=latest_behavior, latest_note=latest_note, suggested_session_number=suggested_session_number, suggested_note_date=suggested_note_date, intake_form_data=intake_form_data, unread_messages_count=unread_messages_count, group_attendance_rows=group_attendance_rows, group_membership_rows=group_membership_rows, group_arrived_count=group_arrived_count, supervisions=supervisions, diagnosis_documents=diagnosis_documents, goals=goals, chart_data=chart_data, questionnaire_enabled=questionnaire_enabled, selected_questionnaire_titles=selected_questionnaire_titles, available_questionnaire_titles=available_questionnaire_titles, questionnaire_tabs_error=questionnaire_tabs_error, source_questionnaire_titles=source_questionnaire_titles, source_questionnaire_error=source_questionnaire_error, source_questionnaire_activation_url=source_questionnaire_activation_url, questionnaire_tabs_activation_url=questionnaire_tabs_activation_url, source_questionnaire_sheet_url=source_questionnaire_sheet_url, service_types=service_types)


@app.route('/patient/<int:patient_id>/encounter-log', methods=['POST'])
@login_required
def add_patient_encounter_log(patient_id):
    if current_user.role != 'admin':
        return 'Unauthorized', 403

    db = get_db()
    patient = db.execute('SELECT id FROM patients WHERE id = ? AND COALESCE(is_deleted, 0) = 0', (patient_id,)).fetchone()
    if not patient:
        return 'Patient not found', 404

    encounter_date = (request.form.get('encounter_date') or '').strip() or None
    title = (request.form.get('title') or '').strip() or None
    content = (request.form.get('content') or '').strip()
    link_url = (request.form.get('link_url') or '').strip() or None

    if not content:
        flash('Encounter note content is required.', 'error')
        return redirect(url_for('patient_detail', patient_id=patient_id, tab='notes'))

    db.execute(
        'INSERT INTO patient_logs (patient_id, encounter_date, title, content, link_url) VALUES (?, ?, ?, ?, ?)',
        (patient_id, encounter_date, title, content, link_url)
    )
    db.execute(
        'INSERT INTO notes (patient_id, note_date, content, link_url) VALUES (?, ?, ?, ?)',
        (patient_id, encounter_date, f"{title}: {content}" if title else content, link_url)
    )
    db.commit()
    flash('Encounter note added.', 'success')
    return redirect(url_for('patient_detail', patient_id=patient_id, tab='notes'))


@app.route('/patient/<int:patient_id>/encounter-log/<int:log_id>/delete', methods=['POST'])
@login_required
def delete_patient_encounter_log(patient_id, log_id):
    if current_user.role != 'admin':
        return 'Unauthorized', 403

    db = get_db()
    deleted = db.execute(
        'DELETE FROM patient_logs WHERE id = ? AND patient_id = ?',
        (log_id, patient_id)
    ).rowcount
    db.commit()
    if deleted:
        flash('Encounter note deleted.', 'success')
    else:
        flash('Encounter note not found.', 'error')
    return redirect(url_for('patient_detail', patient_id=patient_id, tab='notes'))


@app.route('/admin/patient/<int:patient_id>/portal_preview')
@login_required
def admin_portal_preview(patient_id):
    if current_user.role != 'admin':
        return 'Unauthorized', 403
    flash('Portal preview was removed from this workflow.', 'info')
    return redirect(url_for('patient_detail', patient_id=patient_id))


def get_primary_admin_user(db):
    return db.execute(
        "SELECT id, COALESCE(display_name, username) AS name FROM users WHERE role = 'admin' AND COALESCE(is_active, 1) = 1 ORDER BY id ASC LIMIT 1"
    ).fetchone()


def format_lead_time_for_notice(target_dt, reference_dt=None):
    reference_dt = reference_dt or datetime.now()
    delta_seconds = int((target_dt - reference_dt).total_seconds())
    if delta_seconds <= 0:
        return 'after the meeting time'

    days, remainder = divmod(delta_seconds, 86400)
    hours, remainder = divmod(remainder, 3600)
    minutes = remainder // 60
    parts = []
    if days:
        parts.append(f'{days} day' + ('s' if days != 1 else ''))
    if hours:
        parts.append(f'{hours} hour' + ('s' if hours != 1 else ''))
    if minutes or not parts:
        parts.append(f'{minutes} minute' + ('s' if minutes != 1 else ''))
    return ', '.join(parts)


def add_patient_chat_request(db, patient_user_id, patient_id, admin_message, patient_ack_message, audit_action=None, audit_details=None):
    admin_user = get_primary_admin_user(db)
    if not admin_user:
        return False

    db.execute(
        'INSERT INTO messages (sender_id, recipient_id, content) VALUES (?, ?, ?)',
        (patient_user_id, admin_user['id'], admin_message)
    )
    db.execute(
        'INSERT INTO messages (sender_id, recipient_id, content) VALUES (?, ?, ?)',
        (admin_user['id'], patient_user_id, patient_ack_message)
    )
    db.execute('INSERT INTO notifications (message, is_read) VALUES (?, 0)', (admin_message,))
    if audit_action:
        db.execute(
            'INSERT INTO audit_logs (patient_id, action, details) VALUES (?, ?, ?)',
            (patient_id, audit_action, audit_details or admin_message)
        )
    return True


# ── Patient supervision ───────────────────────────────────────────────────────

@app.route('/patient/<int:patient_id>/supervision', methods=['POST'])
@login_required
def add_patient_supervision(patient_id):
    if current_user.role != 'admin':
        return 'Unauthorized', 403
    db = get_db()
    if not db.execute('SELECT id FROM patients WHERE id = ? AND COALESCE(is_deleted,0)=0', (patient_id,)).fetchone():
        return 'Patient not found', 404
    sup_date = (request.form.get('supervision_date') or '').strip()
    supervisor = (request.form.get('supervisor_name') or '').strip()
    content = (request.form.get('content') or '').strip()
    if not sup_date or not content:
        flash('Date and content are required.')
        return redirect(url_for('patient_detail', patient_id=patient_id, tab='supervision'))
    db.execute(
        'INSERT INTO supervisions (patient_id, supervision_date, supervisor_name, content) VALUES (?,?,?,?)',
        (patient_id, sup_date, supervisor or None, content)
    )
    db.commit()
    flash('Supervision record added.')
    return redirect(url_for('patient_detail', patient_id=patient_id, tab='supervision'))


@app.route('/patient/<int:patient_id>/supervision/<int:sup_id>/delete', methods=['POST'])
@login_required
def delete_patient_supervision(patient_id, sup_id):
    if current_user.role != 'admin':
        return 'Unauthorized', 403
    db = get_db()
    db.execute('DELETE FROM supervisions WHERE id = ? AND patient_id = ?', (sup_id, patient_id))
    db.commit()
    flash('Supervision record deleted.')
    return redirect(url_for('patient_detail', patient_id=patient_id, tab='supervision'))


@app.route('/patient/<int:patient_id>/diagnosis_documents/add', methods=['POST'])
@login_required
def add_diagnosis_document(patient_id):
    if current_user.role != 'admin':
        return 'Unauthorized', 403

    if 'diagnosis_file' not in request.files:
        flash('No file selected.')
        return redirect(url_for('patient_detail', patient_id=patient_id, tab='intake'))

    uploaded = request.files['diagnosis_file']
    if uploaded.filename == '':
        flash('No file selected.')
        return redirect(url_for('patient_detail', patient_id=patient_id, tab='intake'))

    category = (request.form.get('category') or 'test_document').strip().lower()
    if category not in {'test_document', 'final_result'}:
        category = 'test_document'

    title = (request.form.get('title') or '').strip() or None
    notes = (request.form.get('notes') or '').strip() or None
    original_filename = secure_filename(uploaded.filename)
    ext = os.path.splitext(original_filename)[1].lower()
    if not ext or ext not in ALLOWED_DIAGNOSIS_EXTENSIONS:
        flash('File type not allowed. Accepted: pdf, docx, png, jpg, jpeg, tiff.')
        return redirect(url_for('patient_detail', patient_id=patient_id, tab='intake'))
    stored_filename = f"diag_{patient_id}_{secrets.token_hex(8)}{ext}"

    diagnosis_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'diagnosis', str(patient_id))
    os.makedirs(diagnosis_dir, exist_ok=True)
    uploaded.save(os.path.join(diagnosis_dir, stored_filename))

    db = get_db()
    patient = db.execute('SELECT id FROM patients WHERE id = ? AND COALESCE(is_deleted, 0) = 0', (patient_id,)).fetchone()
    if not patient:
        return 'Patient not found', 404

    db.execute('''
        INSERT INTO diagnosis_documents (patient_id, category, title, original_filename, stored_filename, notes)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', (patient_id, category, title, original_filename, stored_filename, notes))
    db.commit()

    flash('Diagnostic document uploaded successfully.')
    return redirect(url_for('patient_detail', patient_id=patient_id, tab='intake'))


@app.route('/patient/<int:patient_id>/diagnosis_documents/<int:doc_id>/download', methods=['GET'])
@login_required
def download_diagnosis_document(patient_id, doc_id):
    db = get_db()
    doc = db.execute('''
        SELECT * FROM diagnosis_documents
        WHERE id = ? AND patient_id = ?
    ''', (doc_id, patient_id)).fetchone()
    if not doc:
        return 'Document not found', 404

    if current_user.role == 'patient' and current_user.patient_id != patient_id:
        return 'Unauthorized', 403

    diagnosis_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'diagnosis', str(patient_id))
    return send_from_directory(
        diagnosis_dir,
        doc['stored_filename'],
        as_attachment=True,
        download_name=doc['original_filename']
    )


@app.route('/patient/<int:patient_id>/diagnosis_documents/<int:doc_id>/delete', methods=['POST'])
@login_required
def delete_diagnosis_document(patient_id, doc_id):
    if current_user.role != 'admin':
        return 'Unauthorized', 403

    db = get_db()
    doc = db.execute('''
        SELECT * FROM diagnosis_documents
        WHERE id = ? AND patient_id = ?
    ''', (doc_id, patient_id)).fetchone()
    if not doc:
        return 'Document not found', 404

    diagnosis_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'diagnosis', str(patient_id))
    os.remove(os.path.join(diagnosis_dir, doc['stored_filename'])) if os.path.exists(os.path.join(diagnosis_dir, doc['stored_filename'])) else None
    db.execute('DELETE FROM diagnosis_documents WHERE id = ? AND patient_id = ?', (doc_id, patient_id))
    db.commit()

    flash('Diagnostic document deleted.')
    return redirect(url_for('patient_detail', patient_id=patient_id, tab='intake'))


def build_external_public_url(endpoint, **values):
    path = url_for(endpoint, _external=False, **values)
    configured_base = (app.config.get('PUBLIC_BASE_URL') or '').strip()
    if configured_base:
        return f"{configured_base.rstrip('/')}{path}"

    forwarded_proto = (request.headers.get('X-Forwarded-Proto') or '').split(',')[0].strip()
    forwarded_host = (request.headers.get('X-Forwarded-Host') or '').split(',')[0].strip()
    forwarded_port = (request.headers.get('X-Forwarded-Port') or '').split(',')[0].strip()
    forwarded_prefix = (request.headers.get('X-Forwarded-Prefix') or '').strip()

    scheme = forwarded_proto or request.scheme
    host = forwarded_host or request.host
    if forwarded_port and forwarded_host and ':' not in forwarded_host and forwarded_port not in ('80', '443'):
        host = f'{forwarded_host}:{forwarded_port}'

    if forwarded_prefix:
        if not forwarded_prefix.startswith('/'):
            forwarded_prefix = f'/{forwarded_prefix}'
        forwarded_prefix = forwarded_prefix.rstrip('/')

    return f'{scheme}://{host}{forwarded_prefix}{path}'





@app.route('/patient/<int:patient_id>/toggle_self_booking', methods=('POST',))
@login_required
def toggle_self_booking(patient_id):
    if current_user.role != 'admin':
        return 'Unauthorized', 403

    db = get_db()
    patient = db.execute('SELECT id, name, can_self_schedule FROM patients WHERE id = ?', (patient_id,)).fetchone()
    if not patient:
        return 'Patient not found', 404

    new_value = 0 if int(patient['can_self_schedule'] or 0) == 1 else 1
    db.execute('UPDATE patients SET can_self_schedule = ? WHERE id = ?', (new_value, patient_id))
    db.commit()
    flash(f"Self-booking {'enabled' if new_value == 1 else 'disabled'} for {patient['name']}.")
    return redirect_to_patient_tab(patient_id, 'info')


def parse_date_safe(value):
    if not value:
        return None
    try:
        return datetime.strptime(value, '%Y-%m-%d').date()
    except ValueError:
        return None


def parse_time_safe(value):
    if not value:
        return None
    raw = value.strip()
    formats = ['%H:%M', '%H:%M:%S']
    for fmt in formats:
        try:
            return datetime.strptime(raw, fmt).time()
        except ValueError:
            continue
    return None


def custom_weekday(date_obj):
    # 0=Sunday, 6=Saturday
    return (date_obj.weekday() + 1) % 7


def combine_dt(date_obj, time_str):
    parsed_time = parse_time_safe((time_str or '').strip()[:5])
    if not parsed_time:
        parsed_time = datetime.strptime('00:00', '%H:%M').time()
    return datetime.combine(date_obj, parsed_time)


def daterange(start_date, end_date):
    current = start_date
    while current <= end_date:
        yield current
        current += timedelta(days=1)



def recurring_occurrences_for_week(appt, week_start, week_end):
    return recurring_occurrences_between(appt, week_start, week_end)





def build_patient_upcoming_events(db, patient_id, days_ahead=120, limit=20):
    """Return upcoming patient-facing events from both appointments and active group meetings."""
    today = datetime.now().date()
    range_end = today + timedelta(days=days_ahead)

    rows = db.execute('''
        SELECT *
        FROM appointments
        WHERE patient_id = ?
          AND COALESCE(status, 'scheduled') = 'scheduled'
          AND ((COALESCE(is_recurring, 0) = 0 AND appointment_date BETWEEN ? AND ?)
               OR (COALESCE(is_recurring, 0) = 1 AND appointment_date <= ?))
        ORDER BY appointment_date ASC, appointment_time ASC
    ''', (patient_id, today.isoformat(), range_end.isoformat(), range_end.isoformat())).fetchall()

    upcoming = []
    for appt in rows:
        is_recurring = int(appt['is_recurring'] or 0) == 1
        if is_recurring:
            occ_dates = recurring_occurrences_between(appt, today, range_end)
        else:
            occ = parse_date_safe(appt['appointment_date'])
            occ_dates = [occ] if occ else []

        for occ_date in occ_dates:
            upcoming.append({
                'id': appt['id'],
                'appointment_date': occ_date.isoformat(),
                'appointment_time': appt['appointment_time'],
                'duration_minutes': int(appt['duration_minutes'] or 60),
                'meeting_type': appt['meeting_type'] or 'in-person',
                'meeting_link': appt['meeting_link'] or '',
                'meeting_title': appt['meeting_title'] or '',
                'notes': appt['meeting_title'] or '',
                'status': appt['status'] or 'scheduled',
                'is_recurring': is_recurring,
                'source': 'appointment'
            })

    group_rows = db.execute('''
        SELECT gs.id,
               gs.session_date,
               gs.session_time,
               gs.duration_minutes,
               gs.meeting_type,
               gs.meeting_link,
               gs.title,
               gs.session_summary,
               gs.status,
               gs.series_id,
               g.name AS group_name
        FROM group_sessions gs
        JOIN groups g ON g.id = gs.group_id
        JOIN group_members gm
          ON gm.group_id = gs.group_id
         AND gm.patient_id = ?
        WHERE COALESCE(gs.status, 'scheduled') = 'scheduled'
          AND gs.session_date BETWEEN ? AND ?
          AND date(COALESCE(gm.joined_at, gs.session_date)) <= date(gs.session_date)
          AND (gm.left_at IS NULL OR date(gm.left_at) >= date(gs.session_date))
        ORDER BY gs.session_date ASC, gs.session_time ASC, gs.id ASC
    ''', (patient_id, today.isoformat(), range_end.isoformat())).fetchall()

    for session in group_rows:
        group_label = session['group_name'] or 'Group Session'
        title = session['title'] or group_label
        upcoming.append({
            'id': f"group-session-{session['id']}",
            'appointment_date': session['session_date'],
            'appointment_time': session['session_time'],
            'duration_minutes': int(session['duration_minutes'] or 60),
            'meeting_type': session['meeting_type'] or 'in-person',
            'meeting_link': session['meeting_link'] or '',
            'meeting_title': title,
            'notes': session['session_summary'] or group_label,
            'status': session['status'] or 'scheduled',
            'is_recurring': bool(session['series_id']),
            'source': 'group_session'
        })

    now = datetime.now()
    upcoming = [
        row for row in upcoming
        if datetime.fromisoformat(f"{row['appointment_date']}T{row['appointment_time'] or '00:00'}") >= now
    ]

    seen_keys: set = set()
    deduped: list = []
    for row in upcoming:
        k = (
            row['appointment_date'],
            (row['appointment_time'] or '')[:5],
            row.get('source', 'appointment'),
            (row.get('meeting_title') or '').strip().lower()
        )
        if k in seen_keys:
            continue
        seen_keys.add(k)
        deduped.append(row)
    deduped.sort(key=lambda row: (row['appointment_date'], row['appointment_time'], row.get('meeting_title') or ''))
    return deduped[:limit]


def send_appointment_reminders(db):
    """Return appointments scheduled for today or tomorrow for admin reminder display."""
    today = datetime.now().date()
    tomorrow = today + timedelta(days=1)

    rows = db.execute('''
        SELECT a.id, a.appointment_date, a.appointment_time, a.duration_minutes,
               a.meeting_type, a.is_recurring,
               p.id AS patient_id, p.name AS patient_name
        FROM appointments a
        JOIN patients p ON p.id = a.patient_id
        WHERE COALESCE(a.status, 'scheduled') = 'scheduled'
          AND COALESCE(p.is_deleted, 0) = 0
          AND a.appointment_date IN (?, ?)
        ORDER BY a.appointment_date ASC, a.appointment_time ASC
    ''', (today.isoformat(), tomorrow.isoformat())).fetchall()

    reminders = []
    for row in rows:
        appt_date = parse_date_safe(row['appointment_date'])
        if appt_date is None:
            continue
        days_away = (appt_date - today).days
        reminders.append({
            'appointment_id': row['id'],
            'patient_id': row['patient_id'],
            'patient_name': row['patient_name'],
            'appointment_date': row['appointment_date'],
            'appointment_time': row['appointment_time'],
            'duration_minutes': int(row['duration_minutes'] or 60),
            'meeting_type': row['meeting_type'] or 'in-person',
            'is_today': days_away == 0,
            'is_tomorrow': days_away == 1,
        })
    return reminders


def _sms_settings_summary():
    sid = app.config.get('TWILIO_ACCOUNT_SID', '')
    token = app.config.get('TWILIO_AUTH_TOKEN', '')
    configured = bool(sid and token)
    return {
        'configured': configured,
        'from_number': app.config.get('TWILIO_FROM_NUMBER', ''),
    }


def send_sms(phone, message):
    """Send an SMS via Twilio if configured, otherwise log to sms_logs.
    Returns (status, gateway_response)."""
    settings = _sms_settings_summary()
    status = 'pending'
    gateway_response = None

    if settings['configured']:
        try:
            from twilio.rest import Client
            client = Client(settings['sid'], settings['token'])
            twilio_msg = client.messages.create(
                body=message,
                from_=settings['from_number'],
                to=phone,
            )
            status = 'sent'
            gateway_response = twilio_msg.sid
        except Exception as exc:
            status = 'failed'
            gateway_response = str(exc)
    else:
        gateway_response = 'Twilio not configured; SMS logged only'
        status = 'pending'

    db = get_db()
    db.execute(
        'INSERT INTO sms_logs (recipient_phone, message_body, status, gateway_response) VALUES (?, ?, ?, ?)',
        (phone, message, status, gateway_response))
    db.commit()
    return status, gateway_response


def _send_appointment_sms_reminders(db):
    """Send SMS reminders for upcoming appointments within the reminder window."""
    from_number = _sms_settings_summary()['from_number']
    look_ahead = 7
    today_str = datetime.now().strftime('%Y-%m-%d')
    look_ahead_date = (datetime.now() + timedelta(days=look_ahead)).strftime('%Y-%m-%d')

    candidates = db.execute('''
        SELECT a.*, p.name AS patient_name, p.phone AS patient_phone
        FROM appointments a
        JOIN patients p ON p.id = a.patient_id
        WHERE COALESCE(a.status, 'scheduled') = 'scheduled'
          AND COALESCE(p.is_deleted, 0) = 0
          AND COALESCE(p.reminder_sms_enabled, 0) = 1
          AND p.phone IS NOT NULL AND p.phone != ''
          AND a.appointment_date BETWEEN ? AND ?
        ORDER BY a.appointment_date ASC, a.appointment_time ASC
    ''', (today_str, look_ahead_date)).fetchall()

    sent_count = 0

    for row in candidates:
        appt_date = str(row['appointment_date'] or '')
        appt_time = str(row['appointment_time'] or '')[:5]
        try:
            appt_dt = datetime.strptime(f'{appt_date} {appt_time}', '%Y-%m-%d %H:%M')
        except (ValueError, TypeError):
            continue
        hours_before = int(row['reminder_hours_before'] if row['reminder_hours_before'] else 24)
        reminder_start = appt_dt - timedelta(hours=hours_before + 1)
        reminder_end = appt_dt - timedelta(hours=hours_before - 1)
        now_dt = datetime.now()
        if not (reminder_start <= now_dt <= reminder_end):
            continue

        meeting_type = row['meeting_type'] or 'in-person'
        meeting_title = row['meeting_title'] or ''
        message = (
            f'Reminder: {row["patient_name"]}, you have a {meeting_type} appointment'
            f'{f" ({meeting_title})" if meeting_title else ""}'
            f' on {appt_date} at {appt_time}.'
            f' Please contact the clinic if you need to reschedule.'
        )
        if from_number:
            message = message + f' From: {from_number}'

        phone = row['patient_phone']
        status, _ = send_sms(phone, message)
        if status == 'sent':
            sent_count += 1

    return sent_count


def _send_appointment_email_reminders(db):
    """Send email reminders for upcoming appointments within the reminder window."""
    settings = _smtp_settings_summary()
    if not settings['configured']:
        return 0

    tpl = db.execute(
        "SELECT hours_before, subject_template, body_template FROM email_reminder_templates "
        "WHERE event_type = 'appointment_reminder' AND enabled = 1"
    ).fetchone()

    if tpl:
        global_hours_before = int(tpl['hours_before'] or 24)
        use_custom = True
    else:
        global_hours_before = int(app.config.get('REMINDER_HOURS_BEFORE', 24) or 24)
        use_custom = False

    look_ahead = 7
    today_str = datetime.now().strftime('%Y-%m-%d')
    look_ahead_date = (datetime.now() + timedelta(days=look_ahead)).strftime('%Y-%m-%d')

    candidates = db.execute('''
        SELECT a.*, p.name AS patient_name, p.email AS patient_email,
               u.email AS user_email
        FROM appointments a
        JOIN patients p ON p.id = a.patient_id
        LEFT JOIN users u ON u.patient_id = p.id AND u.role = 'patient' AND COALESCE(u.is_active, 1) = 1
        WHERE COALESCE(a.status, 'scheduled') = 'scheduled'
          AND COALESCE(p.is_deleted, 0) = 0
          AND COALESCE(p.reminder_email_enabled, 1) = 1
          AND a.reminder_sent_at IS NULL
          AND a.appointment_date BETWEEN ? AND ?
        ORDER BY a.appointment_date ASC, a.appointment_time ASC
    ''', (today_str, look_ahead_date)).fetchall()

    non_recurring = []
    recurring_candidates = []

    for row in candidates:
        actual_hours = int(row['reminder_hours_before'] if row['reminder_hours_before'] else global_hours_before)
        appt_date = str(row['appointment_date'] or '')
        appt_time = str(row['appointment_time'] or '')[:5]
        try:
            appt_dt = datetime.strptime(f'{appt_date} {appt_time}', '%Y-%m-%d %H:%M')
        except (ValueError, TypeError):
            continue
        reminder_start = appt_dt - timedelta(hours=actual_hours + 1)
        reminder_end = appt_dt - timedelta(hours=actual_hours - 1)
        now_dt = datetime.now()
        if reminder_start <= now_dt <= reminder_end:
            if row['is_recurring']:
                recurring_candidates.append(row)
            else:
                non_recurring.append(row)

    sent_count = 0
    log_entries = []

    def _fill_vars(text, **kw):
        for k, v in kw.items():
            text = text.replace('{{ ' + k + ' }}', v).replace('{{' + k + '}}', v)
        return text

    def process_row(row):
        nonlocal sent_count
        recipient = (row['user_email'] or '').strip() or (row['patient_email'] or '').strip()
        if not recipient:
            return
        meeting_title = row['meeting_title'] or ''
        meeting_type = row['meeting_type'] or 'in-person'
        meeting_link = row['meeting_link'] or ''
        time_str = str(row['appointment_time'] or '')[:5]
        date_str = str(row['appointment_date'] or '')

        patient_name = row['patient_name']
        vars_dict = dict(
            patient_name=patient_name,
            date=date_str,
            time=time_str,
            meeting_type=meeting_type,
            meeting_title=meeting_title,
            meeting_link=meeting_link,
            clinic_name='Private Clinic',
        )

        if use_custom:
            subject = _fill_vars(tpl['subject_template'], **vars_dict)
            text_body = _fill_vars(tpl['body_template'], **vars_dict)
            html_body = None
        else:
            try:
                html_body = render_template('emails/reminder_en.html',
                    patient_name=patient_name, date_str=date_str, time_str=time_str,
                    meeting_type=meeting_type, meeting_title=meeting_title,
                    meeting_link=meeting_link, clinic_name='Private Clinic')
                text_body = render_template('emails/reminder.txt',
                    patient_name=patient_name, date_str=date_str, time_str=time_str,
                    meeting_type=meeting_type, meeting_title=meeting_title,
                    meeting_link=meeting_link, clinic_name='Private Clinic')
            except Exception:
                text_body = (
                    f'Hello {patient_name},\n\n'
                    f'This is a reminder about your upcoming appointment:\n'
                    f'  Date: {date_str}\n'
                    f'  Time: {time_str}\n'
                    f'  Type: {meeting_type}'
                    f'{f" ({meeting_title})" if meeting_title else ""}'
                    f'{f"\n\nJoin link: {meeting_link}" if meeting_link else ""}\n\n'
                    f'If you need to reschedule or cancel, please contact the clinic.\n\n'
                    f'Private Clinic'
                )
                html_body = None
            subject = f'Appointment Reminder: {patient_name} on {date_str}'
        success, msg = _send_smtp_email(recipient, subject, text_body, html_body)

        log_entries.append({
            'appointment_id': row['id'],
            'patient_id': row['patient_id'],
            'recipient_email': recipient,
            'status': 'sent' if success else 'failed',
            'error_message': msg if not success else None,
        })

        if success:
            db.execute('UPDATE appointments SET reminder_sent_at = ? WHERE id = ?',
                       (datetime.now().isoformat(), row['id']))
            sent_count += 1

    for row in non_recurring:
        process_row(row)

    if recurring_candidates:
        for row in recurring_candidates:
            date_str = str(row['appointment_date'] or '')
            try:
                occurrences = recurring_occurrences_between(
                    row, date_str, date_str
                )
            except Exception:
                occurrences = []
            row_date = str(row['appointment_date'] or '')
            if any(str(o) == row_date for o in occurrences):
                process_row(row)

    if log_entries:
        for entry in log_entries:
            try:
                db.execute('''INSERT INTO reminder_log
                    (appointment_id, patient_id, recipient_email, status, error_message)
                    VALUES (?, ?, ?, ?, ?)''',
                    (entry['appointment_id'], entry['patient_id'],
                     entry['recipient_email'], entry['status'], entry['error_message']))
            except Exception as exc:
                app.logger.error('Failed to write reminder_log entry: %s', exc)
        db.commit()

    return sent_count


def _scheduler_reminder_job():
    try:
        with app.app_context():
            db = get_db()
            _send_appointment_email_reminders(db)
            _send_appointment_sms_reminders(db)
    except Exception:
        app.logger.exception('Scheduled reminder job failed')


def _scheduler_incoming_email_job():
    try:
        from clinic_app.incoming_email import poll_incoming_email
        poll_incoming_email(app)
    except Exception:
        app.logger.exception('Incoming email polling job failed')


def _scheduler_security_scan_job():
    try:
        with app.app_context():
            db = get_db()
            _run_automated_security_scan(db, force=False)
            db.commit()
    except Exception:
        app.logger.exception('Scheduled security scan job failed')


def ensure_appointment_reminder_worker_started():
    if app.config.get('TESTING'):
        return
    if scheduler.get_job('appointment_reminder'):
        return
    interval = app.config.get('REMINDER_SCHEDULER_INTERVAL', 300)
    scheduler.add_job(
        _scheduler_reminder_job,
        IntervalTrigger(seconds=interval),
        id='appointment_reminder',
        name='Send appointment email reminders',
        replace_existing=True,
    )
    incoming_interval = int(os.environ.get('IMAP_POLL_INTERVAL', '300') or 300)
    scheduler.add_job(
        _scheduler_incoming_email_job,
        IntervalTrigger(seconds=incoming_interval),
        id='incoming_email_poller',
        name='Poll IMAP for incoming email replies',
        replace_existing=True,
    )
    scheduler.add_job(
        _scheduler_security_scan_job,
        IntervalTrigger(seconds=3600),
        id='security_scan_checker',
        name='Check and run automated security scan',
        replace_existing=True,
    )
    if not scheduler.running:
        scheduler.start()
        app.logger.info('APScheduler started: reminders every %ss, incoming email every %ss',
                        interval, incoming_interval)



def _notify_patient_appointment_change(change_type, db, appointment, patient, old_details=None):
    """Send email + internal message about an appointment change.

    change_type: 'cancelled', 'rescheduled', 'new'
    """
    recipient = (patient['email'] or '').strip()
    if not recipient:
        return False

    date_str = str(appointment['appointment_date'] or '')
    time_str = str(appointment['appointment_time'] or '')[:5]
    meeting_type = str(appointment['meeting_type'] or 'in-person')

    if change_type == 'cancelled':
        subject = f'Appointment Cancelled: {date_str} at {time_str}'
        text_body = (
            f'Hello {patient["name"]},\n\n'
            f'Your appointment on {date_str} at {time_str} ({meeting_type}) has been cancelled.\n\n'
            f'If you have any questions, please contact the clinic.\n\n'
            f'Private Clinic'
        )
    elif change_type == 'rescheduled':
        old_date = old_details.get('date', '') if old_details else ''
        old_time = old_details.get('time', '') if old_details else ''
        subject = f'Appointment Rescheduled: was {old_date} {old_time}'
        text_body = (
            f'Hello {patient["name"]},\n\n'
            f'Your appointment has been rescheduled:\n'
            f'  Old: {old_date} at {old_time}\n'
            f'  New: {date_str} at {time_str} ({meeting_type})\n\n'
            f'If you have any questions, please contact the clinic.\n\n'
            f'Private Clinic'
        )
    else:
        subject = f'New Appointment: {date_str} at {time_str}'
        text_body = (
            f'Hello {patient["name"]},\n\n'
            f'A new appointment has been scheduled:\n'
            f'  Date: {date_str}\n'
            f'  Time: {time_str}\n'
            f'  Type: {meeting_type}\n\n'
            f'If you need to reschedule or cancel, please contact the clinic.\n\n'
            f'Private Clinic'
        )

    _send_smtp_email(recipient, subject, text_body)

    user = db.execute('SELECT id FROM users WHERE patient_id = ?', (patient['id'],)).fetchone()
    if user:
        admin = db.execute("SELECT id FROM users WHERE role = 'admin' ORDER BY id ASC LIMIT 1").fetchone()
        sender_id = admin['id'] if admin else user['id']
        db.execute(
            'INSERT INTO messages (sender_id, recipient_id, content) VALUES (?, ?, ?)',
            (sender_id, user['id'], f'{subject}\n\n{text_body}')
        )
        db.commit()

    return True



def _process_calendar_follow_ups(db, today):
    follow_up_alerts = []
    # Candidate with a past one-time session and no future booking needs a decision.
    follow_up_rows = db.execute('''
        SELECT p.id AS patient_id, p.name, p.status, MAX(a.appointment_date) AS last_date
        FROM patients p
        JOIN appointments a ON a.patient_id = p.id
        WHERE p.status = 'candidate'
          AND a.is_recurring = 0
          AND COALESCE(a.status, 'scheduled') = 'scheduled'
          AND a.appointment_date < ?
          AND NOT EXISTS (
              SELECT 1 FROM appointments a2
              WHERE a2.patient_id = p.id
                AND a2.appointment_date >= ?
                AND COALESCE(a2.status, 'scheduled') = 'scheduled'
          )
        GROUP BY p.id, p.name, p.status
    ''', (today.isoformat(), today.isoformat())).fetchall()

    for row in follow_up_rows:
        has_future = db.execute('''
            SELECT 1 FROM appointments
            WHERE patient_id = ? AND appointment_date >= ?
              AND COALESCE(status, 'scheduled') = 'scheduled'
            LIMIT 1
        ''', (row['patient_id'], today.isoformat())).fetchone()
        if not has_future:
            follow_up_alerts.append({
                'patient_id': row['patient_id'],
                'patient_name': row['name'],
                'status': row['status'],
                'last_meeting_date': row['last_date'],
                'message': 'Initial one-time meeting has passed with no next booking. Further decision is needed.'
            })
    return follow_up_alerts

def _process_calendar_appointments(appointment_rows, user, week_start, week_end, events, occupied, emitted_appointment_keys):
    for appt in appointment_rows:
        is_recurring = int(appt['is_recurring'] or 0) == 1
        occ_dates = recurring_occurrences_for_week(appt, week_start, week_end) if is_recurring else [parse_date_safe(appt['appointment_date'])]
        occ_dates = [d for d in occ_dates if d is not None]

        for occ_date in occ_dates:
            start_dt = combine_dt(occ_date, appt['appointment_time'])
            duration = int(appt['duration_minutes'] or 60)
            end_dt = start_dt + timedelta(minutes=duration)

            # Patients should not see other patients' bookings, only their own.
            if user.role == 'patient' and appt['patient_id'] != user.patient_id:
                occupied.append((start_dt, end_dt))
                continue

            title = appt['patient_name']

            is_own = (user.role == 'patient' and appt['patient_id'] == user.patient_id)
            can_delete = user.role == 'admin' or is_own

            # Prevent duplicate renders when legacy data has multiple recurring rows
            # that resolve to the same patient+time occurrence in the same week.
            appointment_key = (appt['patient_id'], start_dt.isoformat(), end_dt.isoformat())
            if appointment_key in emitted_appointment_keys:
                continue
            emitted_appointment_keys.add(appointment_key)

            event_color = '#2563eb' if appt['patient_status'] == 'ongoing' else '#f59e0b'
            if appt['patient_status'] == 'archived':
                event_color = '#6b7280'

            platform = (appt['meeting_platform'] or '') if 'meeting_platform' in appt.keys() else ''
            meeting_title = (appt['meeting_title'] or '') if 'meeting_title' in appt.keys() else ''
            save_to_google = int(appt['save_to_google'] or 0) if 'save_to_google' in appt.keys() else 0
            events.append({
                'id': f"appointment-{appt['id']}-{occ_date.isoformat()}",
                'appointment_id': appt['id'],
                'patient_id': appt['patient_id'],
                'title': title,
                'start': start_dt.isoformat(),
                'end': end_dt.isoformat(),
                'editable': False,
                'color': event_color,
                'meta': {
                    'type': 'appointment',
                    'appointment_id': appt['id'],
                    'patient_id': appt['patient_id'],
                    'patient_name': appt['patient_name'],
                    'patient_status': appt['patient_status'],
                    'is_recurring': is_recurring,
                    'meeting_type': appt['meeting_type'],
                    'meeting_link': appt['meeting_link'],
                    'meeting_platform': platform,
                    'meeting_title': meeting_title,
                    'save_to_google': save_to_google,
                    'can_delete': can_delete,
                    'can_edit': can_delete
                }
            })
            occupied.append((start_dt, end_dt))


def _process_calendar_group_sessions(group_sessions, user, events, occupied):
    for group_session in group_sessions:
        session_date = parse_date_safe(group_session['session_date'])
        if not session_date:
            continue

        session_start = combine_dt(session_date, group_session['session_time'])
        session_duration = int(group_session['duration_minutes'] or 60)
        session_end = session_start + timedelta(minutes=session_duration)

        # Keep group slots occupied for availability math, but hide group events from patients.
        if user.role != 'admin':
            occupied.append((session_start, session_end))
            continue

        detail_url = url_for('group_detail', group_id=group_session['group_id'], show_upcoming='all') + f"#session-record-{group_session['id']}"

        events.append({
            'id': f"group-session-{group_session['id']}",
            'group_session_id': group_session['id'],
            'group_id': group_session['group_id'],
            'title': f"Group: {group_session['group_name']}",
            'start': session_start.isoformat(),
            'end': session_end.isoformat(),
            'editable': False,
            'color': '#8b5cf6',
            'meta': {
                'type': 'group_session',
                'group_session_id': group_session['id'],
                'session_date': group_session['session_date'],
                'session_time': group_session['session_time'],
                'duration_minutes': session_duration,
                'title': group_session['title'] or '',
                'facilitator': group_session['facilitator'] or '',
                'group_name': group_session['group_name'],
                'meeting_type': group_session['meeting_type'],
                'meeting_link': group_session['meeting_link'],
                'detail_url': detail_url,
                'can_delete': user.role == 'admin',
                'can_edit': user.role == 'admin'
            }
        })
        occupied.append((session_start, session_end))


def _process_calendar_blocks(blocks, user, events, occupied, weekend_specials):
    for block in blocks:
        block_date = parse_date_safe(block['blocked_date'])
        if not block_date:
            continue
        start_dt = combine_dt(block_date, block['blocked_time'])
        duration = int(block['duration_minutes'] or 60)
        end_dt = start_dt + timedelta(minutes=duration)
        is_private = int(block['is_private'] or 0) == 1
        block_type = (block['block_type'] or 'blocked').strip().lower()
        if block_type != 'blocked':
            block_type = 'blocked'
        raw_title = block['title'] or 'Blocked Slot'
        visible_title = raw_title if (user.role == 'admin' or not is_private) else 'Unavailable'

        # Always mark blocked/special slots as occupied so they don't appear in available_slots.
        occupied.append((start_dt, end_dt))

        # Blocked durations are only shown to admin; patients should not see them at all.
        if user.role != 'admin':
            continue

        events.append({
            'id': f"block-{block['id']}",
            'block_id': block['id'],
            'title': visible_title,
            'start': start_dt.isoformat(),
            'end': end_dt.isoformat(),
            'editable': False,
            'color': '#dc2626',
            'meta': {
                'type': 'block',
                'block_id': block['id'],
                'title': raw_title,
                'blocked_date': block['blocked_date'],
                'blocked_time': block['blocked_time'],
                'duration_minutes': duration,
                'block_type': block_type,
                'is_private': is_private,
                'can_edit': user.role == 'admin',
                'can_delete': user.role == 'admin'
            }
        })

        day_code = custom_weekday(block_date)
        if day_code == 5:
            weekend_specials['friday'].append({
                'id': block['id'],
                'title': visible_title,
                'time': block['blocked_time'],
                'duration': duration,
                'type': block_type
            })
        if day_code == 6:
            weekend_specials['saturday'].append({
                'id': block['id'],
                'title': visible_title,
                'time': block['blocked_time'],
                'duration': duration,
                'type': block_type
            })


def _process_calendar_vacancies(db, week_start, week_end, user, events, occupied):
    rows = db.execute('''
        SELECT id, slot_date, slot_time, duration_minutes, recurrence, weekday
        FROM availability
        WHERE (slot_date BETWEEN ? AND ?)
           OR (recurrence IS NOT NULL AND recurrence != '')
        ORDER BY slot_date ASC, slot_time ASC
    ''', (week_start.isoformat(), week_end.isoformat())).fetchall()
    virtual_vacancies = []
    for row in rows:
        if row['slot_date']:
            virtual_vacancies.append({
                'source_kind': 'one-time',
                'source_id': row['id'],
                'slot_date': row['slot_date'],
                'slot_time': row['slot_time'],
                'duration_minutes': row['duration_minutes'],
            })
        elif row['recurrence'] == 'weekly' and row['weekday'] is not None:
            weekday = int(row['weekday'])
            for day in daterange(week_start, week_end):
                if custom_weekday(day) != weekday:
                    continue
                virtual_vacancies.append({
                    'source_kind': 'weekly',
                    'source_id': row['id'],
                    'slot_date': day.isoformat(),
                    'slot_time': row['slot_time'],
                    'duration_minutes': row['duration_minutes'],
                })
    available_slots = []
    seen_slots = set()
    for row in virtual_vacancies:
        day = parse_date_safe(row['slot_date'])
        if not day:
            continue
        slot_time = (row['slot_time'] or '').strip()
        parsed = parse_time_safe(slot_time)
        if not parsed:
            continue
        duration = int(row['duration_minutes'] or 60)
        if duration <= 0:
            duration = 60

        slot_start = datetime.combine(day, parsed)
        slot_end = slot_start + timedelta(minutes=duration)
        slot_key = (day.isoformat(), slot_start.strftime('%H:%M'), duration)
        if slot_key in seen_slots:
            continue
        seen_slots.add(slot_key)

        if not any(overlaps(slot_start, slot_end, occ_start, occ_end) for occ_start, occ_end in occupied):
            available_slots.append({
                'date': day.isoformat(),
                'time': slot_start.strftime('%H:%M'),
                'duration_minutes': duration
            })
            if user.role == 'admin':
                events.append({
                    'id': f"vacancy-{day.isoformat()}-{slot_start.strftime('%H:%M')}",
                    'title': f"Vacant ({duration}min)",
                    'start': slot_start.isoformat(),
                    'end': slot_end.isoformat(),
                    'editable': False,
                    'color': '#10b981',
                    'meta': {
                        'type': 'vacancy',
                        'slot_id': row['source_id'] if row['source_kind'] == 'one-time' else None,
                        'slot_kind': row['source_kind'],
                        'recurring_id': row['source_id'] if row['source_kind'] == 'weekly' else None,
                        'duration_minutes': duration,
                        'can_delete': True,
                    }
                })
    return available_slots


def _process_calendar_external_events(db, week_start, week_end, user):
    external_events = []
    if gcal and gcal.GOOGLE_LIBS_AVAILABLE and user.role == 'admin':
        try:
            all_gcal = gcal.list_events_for_week(db, week_start.isoformat(), week_end.isoformat())
            our_event_ids = {
                row['google_event_id']
                for row in db.execute(
                    'SELECT google_event_id FROM appointments WHERE google_event_id IS NOT NULL'
                ).fetchall()
            }
            our_event_ids.update(
                row['google_event_id']
                for row in db.execute(
                    'SELECT google_event_id FROM group_sessions WHERE google_event_id IS NOT NULL'
                ).fetchall()
            )
            for evt in all_gcal:
                if evt['google_event_id'] and evt['google_event_id'] not in our_event_ids:
                    external_events.append(evt)
        except Exception:
            pass
    return external_events


def build_week_calendar_snapshot(db, week_start, user):
    week_end = week_start + timedelta(days=6)
    today = datetime.now().date()

    patients = {
        row['id']: row for row in db.execute('SELECT id, name, status, can_self_schedule FROM patients').fetchall()
    }

    appointment_rows = db.execute('''
        SELECT a.*, p.name AS patient_name, p.status AS patient_status, p.patient_type AS patient_type
        FROM appointments a
        JOIN patients p ON p.id = a.patient_id
        WHERE (a.is_recurring = 0 AND a.appointment_date BETWEEN ? AND ?)
           OR (a.is_recurring = 1 AND a.appointment_date <= ?)
        ORDER BY a.appointment_date ASC, a.appointment_time ASC
    ''', (week_start.isoformat(), week_end.isoformat(), week_end.isoformat())).fetchall()

    blocks = db.execute('''
        SELECT * FROM blocked_slots
        WHERE blocked_date BETWEEN ? AND ?
        ORDER BY blocked_date ASC, blocked_time ASC
    ''', (week_start.isoformat(), week_end.isoformat())).fetchall()

    group_sessions = db.execute('''
        SELECT gs.*, g.name AS group_name
        FROM group_sessions gs
        JOIN groups g ON g.id = gs.group_id
        WHERE gs.session_date BETWEEN ? AND ?
          AND COALESCE(g.is_active, 1) = 1
          AND COALESCE(gs.status, 'scheduled') = 'scheduled'
        ORDER BY gs.session_date ASC, gs.session_time ASC
    ''', (week_start.isoformat(), week_end.isoformat())).fetchall()

    events = []
    occupied = []
    emitted_appointment_keys = set()
    weekend_specials = {'friday': [], 'saturday': []}

    follow_up_alerts = _process_calendar_follow_ups(db, today)
    _process_calendar_appointments(appointment_rows, user, week_start, week_end, events, occupied, emitted_appointment_keys)
    _process_calendar_group_sessions(group_sessions, user, events, occupied)
    _process_calendar_blocks(blocks, user, events, occupied, weekend_specials)
    available_slots = _process_calendar_vacancies(db, week_start, week_end, user, events, occupied)
    external_events = _process_calendar_external_events(db, week_start, week_end, user)

    return {
        'week_start': week_start.isoformat(),
        'week_end': week_end.isoformat(),
        'events': events,
        'external_events': external_events,
        'weekend_specials': weekend_specials,
        'available_slots': available_slots,
        'follow_up_alerts': follow_up_alerts
    }


def collect_public_available_slots(db, weeks_ahead=10):
    today = datetime.now().date()
    week_start = today - timedelta(days=custom_weekday(today))
    proxy_user = User(0, 'public', 'admin', None, 'public')
    seen = set()
    slots = []

    for offset in range(max(1, weeks_ahead)):
        target_week = week_start + timedelta(days=7 * offset)
        snapshot = build_week_calendar_snapshot(db, target_week, proxy_user)
        for slot in snapshot['available_slots']:
            slot_date = parse_date_safe(slot.get('date'))
            slot_time = parse_time_safe(slot.get('time'))
            duration = int(slot.get('duration_minutes') or 60)
            if not slot_date or not slot_time:
                continue
            if slot_date < today:
                continue
            key = (slot_date.isoformat(), slot_time.strftime('%H:%M'), duration)
            if key in seen:
                continue
            seen.add(key)
            end_dt = datetime.combine(slot_date, slot_time) + timedelta(minutes=duration)
            slots.append({
                'date': slot_date.isoformat(),
                'time': slot_time.strftime('%H:%M'),
                'duration_minutes': duration,
                'end_time': end_dt.strftime('%H:%M'),
                'label': f"{slot_date.isoformat()} {slot_time.strftime('%H:%M')} - {end_dt.strftime('%H:%M')} ({duration} min)"
            })

    slots.sort(key=lambda s: (s['date'], s['time']))
    return slots


def _nearest_calendar_anchor_date(db, user):
    """Pick the best initial date for calendar view so users land on visible events."""
    today = datetime.now().date()
    params = []
    patient_clause = ''
    if user.role == 'patient' and user.patient_id:
        patient_clause = ' AND patient_id = ?'
        params.append(user.patient_id)

    future_appt = db.execute(
        f'''
        SELECT appointment_date
        FROM appointments
                WHERE appointment_date >= ?
          {patient_clause}
        ORDER BY appointment_date ASC, appointment_time ASC
        LIMIT 1
        ''',
        [today.isoformat(), *params]
    ).fetchone()
    if future_appt and parse_date_safe(future_appt['appointment_date']):
        return parse_date_safe(future_appt['appointment_date'])

    past_appt = db.execute(
        f'''
        SELECT appointment_date
        FROM appointments
        WHERE appointment_date < ?
          {patient_clause}
        ORDER BY appointment_date DESC, appointment_time DESC
        LIMIT 1
        ''',
        [today.isoformat(), *params]
    ).fetchone()
    if past_appt and parse_date_safe(past_appt['appointment_date']):
        return parse_date_safe(past_appt['appointment_date'])

    # For admin, fall back to other calendar entities if no appointments exist.
    if user.role == 'admin':
        future_group = db.execute(
            '''
            SELECT session_date AS day
            FROM group_sessions
            WHERE session_date >= ?
            ORDER BY session_date ASC, session_time ASC
            LIMIT 1
            ''',
            (today.isoformat(),)
        ).fetchone()
        if future_group and parse_date_safe(future_group['day']):
            return parse_date_safe(future_group['day'])

        past_group = db.execute(
            '''
            SELECT session_date AS day
            FROM group_sessions
            WHERE session_date < ?
            ORDER BY session_date DESC, session_time DESC
            LIMIT 1
            ''',
            (today.isoformat(),)
        ).fetchone()
        if past_group and parse_date_safe(past_group['day']):
            return parse_date_safe(past_group['day'])

        future_block = db.execute(
            '''
            SELECT blocked_date AS day
            FROM blocked_slots
            WHERE blocked_date >= ?
            ORDER BY blocked_date ASC, blocked_time ASC
            LIMIT 1
            ''',
            (today.isoformat(),)
        ).fetchone()
        if future_block and parse_date_safe(future_block['day']):
            return parse_date_safe(future_block['day'])

        past_block = db.execute(
            '''
            SELECT blocked_date AS day
            FROM blocked_slots
            WHERE blocked_date < ?
            ORDER BY blocked_date DESC, blocked_time DESC
            LIMIT 1
            ''',
            (today.isoformat(),)
        ).fetchone()
        if past_block and parse_date_safe(past_block['day']):
            return parse_date_safe(past_block['day'])

    return today


def _week_start_for_date(day_obj):
    return day_obj - timedelta(days=custom_weekday(day_obj))


def archive_patient_record(db, patient_id):
    patient = db.execute('SELECT id, name FROM patients WHERE id = ?', (patient_id,)).fetchone()
    if not patient:
        return None

    db.execute('''
        UPDATE patients
        SET is_deleted = 1,
            deleted_at = CURRENT_TIMESTAMP,
            status = 'archived'
        WHERE id = ?
    ''', (patient_id,))
    db.execute('UPDATE users SET is_active = 0 WHERE patient_id = ?', (patient_id,))
    return patient


def delete_patient_files(patient_id):
    upload_root = app.config.get('UPLOAD_FOLDER') or ''
    if not upload_root:
        return

    base_dir = Path(upload_root)
    treatment_dir = base_dir / 'treatments' / str(patient_id)
    if treatment_dir.exists():
        shutil.rmtree(treatment_dir, ignore_errors=True)


def permanently_delete_patient_record(db, patient_id):
    patient = db.execute('SELECT id, name FROM patients WHERE id = ?', (patient_id,)).fetchone()
    if not patient:
        return None

    file_rows = db.execute('SELECT filename FROM files WHERE patient_id = ?', (patient_id,)).fetchall()
    user_ids = [int(row['id']) for row in db.execute('SELECT id FROM users WHERE patient_id = ?', (patient_id,)).fetchall()]

    upload_root = Path(app.config.get('UPLOAD_FOLDER') or '.')
    for row in file_rows:
        filename = (row['filename'] or '').strip()
        if not filename:
            continue
        for candidate in (
            upload_root / filename,
            upload_root / 'treatments' / str(patient_id) / filename,
        ):
            try:
                candidate.unlink(missing_ok=True)
            except OSError:
                pass

    delete_patient_files(patient_id)

    if user_ids:
        placeholders = ','.join(['?'] * len(user_ids))
        db.execute(f'DELETE FROM messages WHERE sender_id IN ({placeholders}) OR recipient_id IN ({placeholders})', tuple(user_ids + user_ids))

    db.execute('DELETE FROM group_session_attendance WHERE patient_id = ?', (patient_id,))
    db.execute('DELETE FROM group_member_history WHERE patient_id = ?', (patient_id,))
    db.execute('DELETE FROM group_members WHERE patient_id = ?', (patient_id,))
    db.execute('DELETE FROM patient_resources WHERE patient_id = ?', (patient_id,))
    db.execute('DELETE FROM goals WHERE patient_id = ?', (patient_id,))
    db.execute('DELETE FROM notes WHERE patient_id = ?', (patient_id,))
    db.execute('DELETE FROM receipts WHERE patient_id = ?', (patient_id,))
    db.execute('DELETE FROM appointments WHERE patient_id = ?', (patient_id,))
    db.execute('DELETE FROM files WHERE patient_id = ?', (patient_id,))
    db.execute('DELETE FROM audit_logs WHERE patient_id = ?', (patient_id,))
    db.execute('DELETE FROM users WHERE patient_id = ?', (patient_id,))
    db.execute('DELETE FROM patients WHERE id = ?', (patient_id,))
    return patient


def build_group_session_collections(group_sessions, show_all_past=False, show_all_upcoming=False):
    now = datetime.now()
    past_sessions = []
    upcoming_sessions = []

    for row in group_sessions:
        session = dict(row)
        session_date = parse_date_safe(session['session_date'])
        session_time = parse_time_safe(session['session_time'])
        if session_date and session_time:
            start_dt = datetime.combine(session_date, session_time)
            end_dt = start_dt + timedelta(minutes=int(session['duration_minutes'] or 60))
        else:
            start_dt = None
            end_dt = None
        session['starts_at'] = start_dt
        session['ends_at'] = end_dt

        if end_dt and end_dt < now:
            past_sessions.append(session)
        else:
            upcoming_sessions.append(session)

    past_sessions.sort(key=lambda item: item['starts_at'] or datetime.min, reverse=True)
    upcoming_sessions.sort(key=lambda item: item['starts_at'] or datetime.max)

    visible_past = past_sessions if show_all_past else past_sessions[:2]
    visible_upcoming = upcoming_sessions if show_all_upcoming else upcoming_sessions[:2]

    return {
        'past_sessions_all': past_sessions,
        'upcoming_sessions_all': upcoming_sessions,
        'visible_past_sessions': visible_past,
        'visible_upcoming_sessions': visible_upcoming,
        'hidden_past_count': max(0, len(past_sessions) - len(visible_past)),
        'hidden_upcoming_count': max(0, len(upcoming_sessions) - len(visible_upcoming)),
        'show_all_past': bool(show_all_past),
        'show_all_upcoming': bool(show_all_upcoming),
    }


def get_group_members_for_session(db, group_id, session_date_iso):
    """Resolve members by membership periods active on the session date."""
    rows = db.execute('''
        SELECT p.id AS patient_id,
               p.name AS patient_name,
               h.joined_at,
               h.left_at,
               COALESCE(h.role, 'member') AS role
        FROM group_member_history h
        JOIN patients p ON p.id = h.patient_id
        WHERE h.group_id = ?
          AND date(h.joined_at) <= date(?)
          AND (h.left_at IS NULL OR date(h.left_at) >= date(?))
          AND COALESCE(p.is_deleted, 0) = 0
        ORDER BY p.name ASC
    ''', (group_id, session_date_iso, session_date_iso)).fetchall()
    members = [dict(row) for row in rows]
    if members:
        return members

    fallback = db.execute('''
        SELECT p.id AS patient_id,
               p.name AS patient_name,
               gm.joined_at,
               gm.left_at,
               COALESCE(gm.role, 'member') AS role
        FROM group_members gm
        JOIN patients p ON p.id = gm.patient_id
        WHERE gm.group_id = ?
          AND date(gm.joined_at) <= date(?)
          AND (gm.left_at IS NULL OR date(gm.left_at) >= date(?))
          AND COALESCE(p.is_deleted, 0) = 0
        ORDER BY p.name ASC
    ''', (group_id, session_date_iso, session_date_iso)).fetchall()
    return [dict(row) for row in fallback]


@app.route('/groups', methods=['GET', 'POST'])
@login_required
def groups_dashboard():
    if current_user.role != 'admin':
        return redirect(url_for('patient_home'))

    db = get_db()
    if request.method == 'POST':
        name = (request.form.get('name') or '').strip()
        group_type = (request.form.get('group_type') or 'support').strip()
        description = (request.form.get('description') or '').strip()
        if not name:
            flash('Group name is required.')
        else:
            db.execute('INSERT INTO groups (name, group_type, description) VALUES (?, ?, ?)',
                       (name, group_type or 'support', description or None))
            db.commit()
            flash('Group created.')
        return redirect(url_for('groups_dashboard'))

    groups = db.execute('''
        SELECT g.*, COUNT(gm.patient_id) AS member_count,
               (
                 SELECT COUNT(*)
                 FROM group_sessions gs
                 WHERE gs.group_id = g.id
               ) AS session_count,
               (
                 SELECT MIN(gs2.session_date)
                 FROM group_sessions gs2
                 WHERE gs2.group_id = g.id AND COALESCE(gs2.status, 'scheduled') = 'scheduled'
               ) AS next_session_date
        FROM groups g
        LEFT JOIN group_members gm ON gm.group_id = g.id AND gm.left_at IS NULL
        GROUP BY g.id
        ORDER BY g.created_at DESC, g.name ASC
    ''').fetchall()

    return render_template('groups_overview.html', groups=groups)



def _get_group_member_history(db, group_id):
    member_history_rows = db.execute('''
         SELECT h.id,
             h.group_id,
               h.patient_id,
               p.name AS patient_name,
               h.joined_at,
               h.left_at,
               COALESCE(h.role, 'member') AS role
        FROM group_member_history h
        JOIN patients p ON p.id = h.patient_id
                WHERE h.group_id = ?
                    AND COALESCE(p.is_deleted, 0) = 0
        ORDER BY h.group_id ASC, h.joined_at DESC
        ''', (group_id,)).fetchall()

    member_history_rows = [dict(row) for row in member_history_rows]
    now_date = datetime.now().date()
    for row in member_history_rows:
        joined_date = parse_date_safe((row.get('joined_at') or '')[:10])
        left_date = parse_date_safe((row.get('left_at') or '')[:10]) if row.get('left_at') else None
        if joined_date:
            end_date = left_date or now_date
            row['membership_days'] = max(0, (end_date - joined_date).days)
        else:
            row['membership_days'] = None
    return member_history_rows


def _get_group_attendance_data(db, group_sessions, member_history_rows, group_members):
    session_member_map = {}
    attendance_by_session = {}
    session_ids = [int(row['id']) for row in group_sessions]

    for gs_row in group_sessions:
        session_date_iso = gs_row['session_date']

        members = []
        for row in member_history_rows:
            joined_date = row['joined_at'][:10] if row['joined_at'] else ''
            left_date = row['left_at'][:10] if row['left_at'] else None

            if joined_date <= session_date_iso and (left_date is None or left_date >= session_date_iso):
                members.append(dict(row))

        if not members:
            for row in group_members:
                joined_date = row['joined_at'][:10] if row['joined_at'] else ''
                left_date = row['left_at'][:10] if row['left_at'] else None

                if joined_date <= session_date_iso and (left_date is None or left_date >= session_date_iso):
                    members.append(dict(row))

        session_member_map[int(gs_row['id'])] = members

    if session_ids:
        marks = db.execute(f'''
            SELECT session_id, patient_id, attendance_status, absence_reason, notified_on_time, attendance_note
            FROM group_session_attendance
            WHERE session_id IN ({','.join(['?'] * len(session_ids))})
        ''', session_ids).fetchall()
        for row in marks:
            session_key = int(row['session_id'])
            attendance_by_session.setdefault(session_key, {})[int(row['patient_id'])] = {
                'attendance_status': row['attendance_status'] or 'pending',
                'absence_reason': row['absence_reason'] or '',
                'notified_on_time': int(row['notified_on_time'] or 0),
                'attendance_note': row['attendance_note'] or ''
            }

    return session_member_map, attendance_by_session


def _get_patient_arrived_counts(db):
    arrived_rows = db.execute('''
        SELECT patient_id, COUNT(*) AS arrived_count
        FROM group_session_attendance
        WHERE attendance_status = 'present'
        GROUP BY patient_id
    ''').fetchall()
    return {int(row['patient_id']): int(row['arrived_count'] or 0) for row in arrived_rows}


def _get_available_group_patients(db):
    return db.execute('''
        SELECT id, name
        FROM patients
        WHERE COALESCE(is_deleted, 0) = 0
          AND COALESCE(patient_type, 'private') = 'group'
        ORDER BY name ASC
    ''').fetchall()


def build_group_detail_payload(db, group_id, show_all_past=False, show_all_upcoming=False):
    group = db.execute('SELECT * FROM groups WHERE id = ?', (group_id,)).fetchone()
    if not group:
        return None

    group_members = [dict(row) for row in db.execute('''
        SELECT gm.group_id, p.id AS patient_id, p.name AS patient_name,
               gm.joined_at, gm.left_at, p.status,
               (SELECT COUNT(*) FROM appointments a WHERE a.patient_id = p.id AND a.is_recurring = 1 AND COALESCE(a.status, 'scheduled') = 'scheduled') as has_recurring
        FROM group_members gm
        JOIN patients p ON p.id = gm.patient_id
                WHERE gm.group_id = ?
                    AND gm.left_at IS NULL
                    AND COALESCE(p.is_deleted, 0) = 0
        ORDER BY p.name ASC
        ''', (group_id,)).fetchall()]

    group_sessions = db.execute('''
        SELECT gs.*, g.name AS group_name,
               ss.recurrence_interval_weeks,
               ss.recurrence_end_date,
               ss.recurrence_count
        FROM group_sessions gs
        JOIN groups g ON g.id = gs.group_id
        LEFT JOIN group_session_series ss ON ss.id = gs.series_id
        WHERE gs.group_id = ?
        ORDER BY gs.session_date ASC, gs.session_time ASC
    ''', (group_id,)).fetchall()

    member_history_rows = _get_group_member_history(db, group_id)
    session_member_map, attendance_by_session = _get_group_attendance_data(db, group_sessions, member_history_rows, group_members)
    arrived_count_map = _get_patient_arrived_counts(db)
    patients = _get_available_group_patients(db)

    session_collections = build_group_session_collections(
        group_sessions,
        show_all_past=show_all_past,
        show_all_upcoming=show_all_upcoming
    )

    group_supervisions = db.execute(
        'SELECT * FROM supervisions WHERE group_id = ? ORDER BY supervision_date DESC, created_at DESC',
        (group_id,)
    ).fetchall()

    return {
        'group': group,
        'group_members': group_members,
        'group_sessions': group_sessions,
        'patients': patients,
        'member_history_rows': member_history_rows,
        'session_member_map': session_member_map,
        'attendance_by_session': attendance_by_session,
        'arrived_count_map': arrived_count_map,
        'group_supervisions': group_supervisions,
        **session_collections
    }


@app.route('/groups/<int:group_id>', methods=['GET'])
@login_required
def group_detail(group_id):
    if current_user.role != 'admin':
        return redirect(url_for('patient_home'))

    db = get_db()
    show_all_past = (request.args.get('show_past') or '').strip().lower() == 'all'
    show_all_upcoming = (request.args.get('show_upcoming') or '').strip().lower() == 'all'
    payload = build_group_detail_payload(
        db,
        group_id,
        show_all_past=show_all_past,
        show_all_upcoming=show_all_upcoming
    )
    if payload is None:
        flash('Group not found.')
        return redirect(url_for('groups_dashboard'))
    return render_template('groups.html', **payload)


@app.route('/groups/<int:group_id>/delete', methods=['POST'])
@login_required
def delete_group(group_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    delete_mode = (request.form.get('delete_mode') or 'archive').strip().lower()
    return_to = (request.form.get('return_to') or 'dashboard').strip().lower()

    db = get_db()
    group = db.execute('SELECT id, name FROM groups WHERE id = ?', (group_id,)).fetchone()
    if not group:
        flash('Group not found.')
        return redirect(url_for('groups_dashboard'))

    if delete_mode == 'delete':
        session_ids = [int(row['id']) for row in db.execute('SELECT id FROM group_sessions WHERE group_id = ?', (group_id,)).fetchall()]
        if session_ids:
            placeholders = ','.join(['?'] * len(session_ids))
            db.execute(f'DELETE FROM group_session_attendance WHERE session_id IN ({placeholders})', session_ids)
        db.execute('DELETE FROM group_sessions WHERE group_id = ?', (group_id,))
        db.execute('DELETE FROM group_session_series WHERE group_id = ?', (group_id,))
        db.execute('DELETE FROM group_member_history WHERE group_id = ?', (group_id,))
        db.execute('DELETE FROM group_members WHERE group_id = ?', (group_id,))
        db.execute('DELETE FROM groups WHERE id = ?', (group_id,))
        db.commit()
        flash('Group and all related data deleted.')
        return redirect(url_for('groups_dashboard'))

    db.execute('UPDATE groups SET is_active = 0 WHERE id = ?', (group_id,))
    db.commit()
    flash('Group moved to history.')
    if return_to == 'detail':
        return redirect(url_for('groups_dashboard'))
    return redirect(url_for('groups_dashboard'))


@app.route('/groups/<int:group_id>/update', methods=['POST'])
@login_required
def update_group_info(group_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    name = (request.form.get('name') or '').strip()
    group_type = (request.form.get('group_type') or 'support').strip() or 'support'
    description = (request.form.get('description') or '').strip()
    is_active = 1 if request.form.get('is_active') in ('1', 'true', 'on') else 0
    return_to = (request.form.get('return_to') or 'detail').strip().lower()

    if not name:
        flash('Group name is required.')
        if return_to == 'dashboard':
            return redirect(url_for('groups_dashboard'))
        return redirect(url_for('group_detail', group_id=group_id))

    db = get_db()
    db.execute('''
        UPDATE groups
        SET name = ?, group_type = ?, description = ?, is_active = ?
        WHERE id = ?
    ''', (name, group_type, description or None, is_active, group_id))
    db.commit()
    flash('Group information updated.')
    if return_to == 'dashboard':
        return redirect(url_for('groups_dashboard'))
    return redirect(url_for('group_detail', group_id=group_id))


# ── Group supervision ─────────────────────────────────────────────────────────

@app.route('/groups/<int:group_id>/supervision', methods=['POST'])
@login_required
def add_group_supervision(group_id):
    if current_user.role != 'admin':
        return 'Unauthorized', 403
    db = get_db()
    if not db.execute('SELECT id FROM groups WHERE id = ?', (group_id,)).fetchone():
        return 'Group not found', 404
    sup_date = (request.form.get('supervision_date') or '').strip()
    supervisor = (request.form.get('supervisor_name') or '').strip()
    content = (request.form.get('content') or '').strip()
    if not sup_date or not content:
        flash('Date and content are required.')
        return redirect(url_for('group_detail', group_id=group_id))
    db.execute(
        'INSERT INTO supervisions (group_id, supervision_date, supervisor_name, content) VALUES (?,?,?,?)',
        (group_id, sup_date, supervisor or None, content)
    )
    db.commit()
    flash('Supervision record added.')
    return redirect(url_for('group_detail', group_id=group_id))


@app.route('/groups/<int:group_id>/supervision/<int:sup_id>/delete', methods=['POST'])
@login_required
def delete_group_supervision(group_id, sup_id):
    if current_user.role != 'admin':
        return 'Unauthorized', 403
    db = get_db()
    db.execute('DELETE FROM supervisions WHERE id = ? AND group_id = ?', (sup_id, group_id))
    db.commit()
    flash('Supervision record deleted.')
    return redirect(url_for('group_detail', group_id=group_id))


@app.route('/groups/<int:group_id>/members', methods=['POST'])
@login_required
def add_group_member(group_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    patient_id_raw = (request.form.get('patient_id') or '').strip()
    if not patient_id_raw.isdigit():
        flash('Valid patient is required.')
        return redirect(url_for('group_detail', group_id=group_id))

    db = get_db()
    patient_id = int(patient_id_raw)
    patient_row = db.execute('SELECT id, patient_type FROM patients WHERE id = ? AND COALESCE(is_deleted, 0) = 0', (patient_id,)).fetchone()
    if not patient_row:
        flash('Patient not found.')
        return redirect(url_for('group_detail', group_id=group_id))
    if (patient_row['patient_type'] or 'private') != 'group':
        flash('Only group-type patients can be added to groups.')
        return redirect(url_for('group_detail', group_id=group_id))

    existing_active = db.execute('''
        SELECT 1
        FROM group_members
        WHERE group_id = ? AND patient_id = ? AND left_at IS NULL
        LIMIT 1
    ''', (group_id, patient_id)).fetchone()
    if existing_active:
        flash('Patient is already an active member in this group.')
        return redirect(url_for('group_detail', group_id=group_id))

    db.execute('''
        INSERT INTO group_members (group_id, patient_id, joined_at, left_at, role)
        VALUES (?, ?, CURRENT_TIMESTAMP, NULL, 'member')
        ON CONFLICT(group_id, patient_id)
        DO UPDATE SET joined_at = CURRENT_TIMESTAMP, left_at = NULL, role = 'member'
    ''', (group_id, patient_id))

    existing_open_history = db.execute('''
        SELECT id
        FROM group_member_history
        WHERE group_id = ? AND patient_id = ? AND left_at IS NULL
        ORDER BY joined_at DESC
        LIMIT 1
    ''', (group_id, patient_id)).fetchone()
    if not existing_open_history:
        db.execute('''
            INSERT INTO group_member_history (group_id, patient_id, joined_at, role)
            VALUES (?, ?, CURRENT_TIMESTAMP, 'member')
        ''', (group_id, patient_id))

    db.commit()
    flash('Patient added to group.')
    return redirect(url_for('group_detail', group_id=group_id))


@app.route('/groups/<int:group_id>/members/<int:patient_id>/remove', methods=['POST'])
@login_required
def remove_group_member(group_id, patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    removal_mode = (request.form.get('removal_mode') or 'keep').strip().lower()

    db = get_db()
    group = db.execute('SELECT id, name FROM groups WHERE id = ?', (group_id,)).fetchone()
    patient = db.execute('SELECT id, name FROM patients WHERE id = ?', (patient_id,)).fetchone()
    active_membership = db.execute('''
        SELECT 1
        FROM group_members
        WHERE group_id = ? AND patient_id = ? AND left_at IS NULL
    ''', (group_id, patient_id)).fetchone()

    if not group or not patient:
        flash('Group member not found.')
        return redirect(url_for('group_detail', group_id=group_id))
    if not active_membership:
        flash('Patient is not an active member in this group.')
        return redirect(url_for('group_detail', group_id=group_id))

    db.execute('''
        UPDATE group_members
        SET left_at = CURRENT_TIMESTAMP
        WHERE group_id = ? AND patient_id = ? AND left_at IS NULL
    ''', (group_id, patient_id))

    open_history = db.execute('''
        SELECT id
        FROM group_member_history
        WHERE group_id = ? AND patient_id = ? AND left_at IS NULL
        ORDER BY joined_at DESC
        LIMIT 1
    ''', (group_id, patient_id)).fetchone()
    if open_history:
        db.execute('UPDATE group_member_history SET left_at = CURRENT_TIMESTAMP WHERE id = ?', (open_history['id'],))

    if removal_mode == 'archive':
        archive_patient_record(db, patient_id)
        message = 'Patient removed from group and moved to archived records.'
    elif removal_mode == 'delete':
        permanently_delete_patient_record(db, patient_id)
        message = 'Patient removed from group and deleted with all related data.'
    else:
        message = 'Patient removed from group.'

    db.commit()
    flash(message)
    return redirect(url_for('group_detail', group_id=group_id))


def sync_group_member_current_record(db, group_id, patient_id):
    """Keep group_members aligned with the latest history state."""
    open_row = db.execute('''
        SELECT joined_at
        FROM group_member_history
        WHERE group_id = ? AND patient_id = ? AND left_at IS NULL
        ORDER BY joined_at DESC
        LIMIT 1
    ''', (group_id, patient_id)).fetchone()

    if open_row:
        db.execute('''
            INSERT INTO group_members (group_id, patient_id, joined_at, left_at, role)
            VALUES (?, ?, ?, NULL, 'member')
            ON CONFLICT(group_id, patient_id)
            DO UPDATE SET joined_at = excluded.joined_at, left_at = NULL, role = 'member'
        ''', (group_id, patient_id, open_row['joined_at']))
        return

    latest_row = db.execute('''
        SELECT joined_at, left_at
        FROM group_member_history
        WHERE group_id = ? AND patient_id = ?
        ORDER BY joined_at DESC
        LIMIT 1
    ''', (group_id, patient_id)).fetchone()
    if latest_row:
        db.execute('''
            INSERT INTO group_members (group_id, patient_id, joined_at, left_at, role)
            VALUES (?, ?, ?, ?, 'member')
            ON CONFLICT(group_id, patient_id)
            DO UPDATE SET joined_at = excluded.joined_at, left_at = excluded.left_at, role = 'member'
        ''', (group_id, patient_id, latest_row['joined_at'], latest_row['left_at']))


@app.route('/groups/history/<int:history_id>/dates', methods=['POST'])
@login_required
def update_group_member_history_dates(history_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    joined_date_raw = (request.form.get('joined_date') or '').strip()
    left_date_raw = (request.form.get('left_date') or '').strip()
    return_patient_id_raw = (request.form.get('return_patient_id') or '').strip()
    return_group_id_raw = (request.form.get('return_group_id') or '').strip()

    def redirect_target():
        if return_patient_id_raw.isdigit():
            return redirect_to_patient_tab(int(return_patient_id_raw), 'info')
        if return_group_id_raw.isdigit():
            return redirect(url_for('group_detail', group_id=int(return_group_id_raw)))
        return redirect(url_for('groups_dashboard'))

    joined_date = parse_date_safe(joined_date_raw)
    left_date = parse_date_safe(left_date_raw) if left_date_raw else None
    if not joined_date:
        flash('Joined date is required and must be valid.')
        return redirect_target()
    if left_date and left_date < joined_date:
        flash('Left date cannot be before joined date.')
        return redirect_target()

    db = get_db()
    history = db.execute('''
        SELECT id, group_id, patient_id
        FROM group_member_history
        WHERE id = ?
    ''', (history_id,)).fetchone()
    if not history:
        flash('Membership history row not found.')
        return redirect_target()

    joined_ts = f"{joined_date.isoformat()} 00:00:00"
    left_ts = f"{left_date.isoformat()} 23:59:59" if left_date else None
    db.execute('''
        UPDATE group_member_history
        SET joined_at = ?, left_at = ?
        WHERE id = ?
    ''', (joined_ts, left_ts, history_id))

    sync_group_member_current_record(db, int(history['group_id']), int(history['patient_id']))
    db.commit()
    flash('Membership dates updated.')
    return redirect_target()


def _parse_group_session_form(form):
    return {
        'session_date': (form.get('session_date') or '').strip(),
        'session_time': (form.get('session_time') or '').strip(),
        'end_time_raw': (form.get('end_time') or '').strip(),
        'title': (form.get('title') or '').strip(),
        'facilitator': (form.get('facilitator') or '').strip(),
        'meeting_type': (form.get('meeting_type') or 'in-person').strip(),
        'meeting_link': (form.get('meeting_link') or '').strip(),
        'recurrence_mode': (form.get('recurrence_mode') or 'one-time').strip().lower(),
        'recurrence_interval_raw': (form.get('recurrence_interval_weeks') or '1').strip(),
        'recurrence_end_mode': (form.get('recurrence_end_mode') or 'count').strip().lower(),
        'recurrence_end_raw': (form.get('recurrence_end_date') or '').strip(),
        'recurrence_count_raw': (form.get('recurrence_count') or '').strip()
    }


def _calculate_group_session_duration(parsed_time, parsed_end):
    duration = 60
    if parsed_end:
        start_minutes = parsed_time.hour * 60 + parsed_time.minute
        end_minutes = parsed_end.hour * 60 + parsed_end.minute
        if end_minutes > start_minutes:
            duration = end_minutes - start_minutes
    return duration


def _resolve_group_recurrence_params(recurrence_mode, recurrence_interval_raw, recurrence_end_mode, recurrence_end_raw, recurrence_count_raw):
    try:
        recurrence_interval_weeks = max(1, int(recurrence_interval_raw or '1'))
    except ValueError:
        recurrence_interval_weeks = 1

    recurrence_end_date = parse_date_safe(recurrence_end_raw) if recurrence_end_raw else None
    recurrence_count = None
    if recurrence_count_raw:
        try:
            recurrence_count = max(1, min(104, int(recurrence_count_raw)))
        except ValueError:
            recurrence_count = None

    error_msg = None
    if recurrence_mode == 'weekly':
        if recurrence_end_mode == 'date':
            recurrence_count = None
            if not recurrence_end_date:
                error_msg = 'Please choose an end date for the recurring meetings.'
        else:
            recurrence_end_date = None
            if recurrence_count is None:
                error_msg = 'Please choose how many meetings to create.'

    return recurrence_interval_weeks, recurrence_end_date, recurrence_count, error_msg


def _insert_group_sessions(db, group_id, parsed_date, parsed_time, duration, recurrence_dates, recurrence_mode, recurrence_interval_weeks, recurrence_end_date, recurrence_count, title, facilitator, meeting_type, meeting_link):
    series_id = None
    if recurrence_mode == 'weekly' and len(recurrence_dates) > 1:
        cur = db.execute('''
            INSERT INTO group_session_series (
                group_id, start_date, start_time, duration_minutes,
                recurrence_interval_weeks, recurrence_end_date, recurrence_count,
                title, facilitator, meeting_type, meeting_link, is_active
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 1)
        ''', (
            group_id,
            parsed_date.isoformat(),
            parsed_time.strftime('%H:%M'),
            duration,
            recurrence_interval_weeks,
            recurrence_end_date.isoformat() if recurrence_end_date else None,
            recurrence_count,
            title or None,
            facilitator or None,
            meeting_type or 'in-person',
            meeting_link or None
        ))
        series_id = cur.lastrowid

    last_session_id = None
    for idx, date_item in enumerate(recurrence_dates, start=1):
        cur = db.execute('''
            INSERT INTO group_sessions
                (group_id, session_date, session_time, duration_minutes, title, facilitator, meeting_type, meeting_link, series_id, occurrence_index, status)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'scheduled')
        ''', (
            group_id,
            date_item.isoformat(),
            parsed_time.strftime('%H:%M'),
            duration,
            title or None,
            facilitator or None,
            meeting_type or 'in-person',
            meeting_link or None,
            series_id,
            idx if series_id else None
        ))
        last_session_id = cur.lastrowid

    return series_id, last_session_id


@app.route('/groups/<int:group_id>/sessions', methods=['POST'])
@login_required
def add_group_session(group_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    data = _parse_group_session_form(request.form)

    parsed_date = parse_date_safe(data['session_date'])
    parsed_time = parse_time_safe(data['session_time'])
    parsed_end = parse_time_safe(data['end_time_raw'])

    if not parsed_date or not parsed_time:
        flash('Valid session date and start time are required.')
        return redirect(url_for('group_detail', group_id=group_id))

    duration = _calculate_group_session_duration(parsed_time, parsed_end)
    db = get_db()

    recurrence_interval_weeks, recurrence_end_date, recurrence_count, error_msg = _resolve_group_recurrence_params(
        data['recurrence_mode'], data['recurrence_interval_raw'],
        data['recurrence_end_mode'], data['recurrence_end_raw'], data['recurrence_count_raw']
    )

    if error_msg:
        flash(error_msg)
        return redirect(url_for('group_detail', group_id=group_id))

    recurrence_dates = [parsed_date]
    if data['recurrence_mode'] == 'weekly':
        recurrence_dates = build_group_recurrence_dates(
            parsed_date,
            recurrence_interval_weeks=recurrence_interval_weeks,
            recurrence_end_date=recurrence_end_date,
            recurrence_count=recurrence_count
        )

    for date_item in recurrence_dates:
        start_at = datetime.combine(date_item, parsed_time)
        end_at = start_at + timedelta(minutes=duration)
        conflict_message = has_time_conflict(db, date_item, start_at, end_at)
        if conflict_message:
            flash(f'{conflict_message} ({date_item.isoformat()})')
            return redirect(url_for('group_detail', group_id=group_id))

    series_id, last_session_id = _insert_group_sessions(
        db, group_id, parsed_date, parsed_time, duration, recurrence_dates,
        data['recurrence_mode'], recurrence_interval_weeks, recurrence_end_date, recurrence_count,
        data['title'], data['facilitator'], data['meeting_type'], data['meeting_link']
    )

    db.commit()

    if series_id:
        flash(f"Group recurrence added ({len(recurrence_dates)} sessions).")
    else:
        flash('Group session added.')

    destination = url_for('group_detail', group_id=group_id, show_upcoming='all')
    if last_session_id:
        destination = f'{destination}#session-record-{last_session_id}'
    return redirect(destination)


def _build_group_session_summary_text(session_row, members, attendance_payload, raw_content):
    text = (raw_content or '').strip()
    lowered = text.lower()
    if any(token in lowered for token in ('participants', 'משתתפים')) and any(token in lowered for token in ('content', 'תוכן')):
        return text

    is_he = session.get('lang') == 'he'
    session_date = parse_date_safe(session_row['session_date'])
    formatted_date = session_date.strftime('%d/%m/%y') if session_date else (session_row['session_date'] or '')
    title_value = (session_row['title'] or '').strip()

    heading_prefix = 'פגישה' if is_he else 'Meeting'
    participants_label = 'משתתפים' if is_he else 'Participants'
    missing_label = 'חסרים' if is_he else 'Missing'
    content_label = 'תוכן' if is_he else 'Content'
    none_label = 'אין' if is_he else 'None'
    no_content_label = 'לא נרשם תוכן מפגש.' if is_he else 'No session content recorded.'
    missed_fallback = 'לא הגיע לקבוצה ולא הודיע' if is_he else 'did not attend and did not notify'

    heading = f"{heading_prefix} {title_value or session_row['id']} - {formatted_date}"
    present_names = []
    missing_lines = []

    for member in members:
        patient_id = int(member['patient_id'])
        payload = attendance_payload.get(patient_id, {})
        status_value = (payload.get('attendance_status') or 'pending').strip().lower()
        patient_name = member.get('patient_name') or member.get('name') or f'Patient {patient_id}'
        if status_value == 'present':
            present_names.append(patient_name)
        elif status_value == 'missed':
            reason_text = (payload.get('absence_reason') or '').strip() or missed_fallback
            if is_he:
                missing_lines.append(f"{patient_name} {reason_text}")
            else:
                missing_lines.append(f"{patient_name} — {reason_text}")

    sections = [
        heading,
        participants_label,
        ', '.join(present_names) if present_names else none_label,
        '',
        missing_label,
        '\n'.join(missing_lines) if missing_lines else none_label,
        '',
        content_label,
        text or no_content_label,
    ]
    return '\n'.join(sections).strip()


@app.route('/groups/sessions/<int:session_id>/record', methods=['POST'])
@login_required
def record_group_session(session_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    session_row = db.execute('SELECT * FROM group_sessions WHERE id = ?', (session_id,)).fetchone()
    if not session_row:
        flash('Group session not found.')
        return redirect(url_for('groups_dashboard'))

    raw_session_summary = (request.form.get('session_summary') or '').strip()
    session_status = (request.form.get('session_status') or 'completed').strip().lower()
    if session_status not in ('scheduled', 'completed', 'cancelled'):
        session_status = 'completed'

    members = get_group_members_for_session(db, int(session_row['group_id']), session_row['session_date'])
    attendance_payload = {}

    def upsert_group_session_note(pid, status_value, attendance_note_text, missed_reason_text):
        marker = f"[Group Session #{session_id}]"
        if status_value not in ('present', 'missed'):
            db.execute('DELETE FROM notes WHERE patient_id = ? AND content LIKE ?', (pid, f'{marker}%'))
            return
        status_label = 'Missed' if status_value == 'missed' else 'Present'
        note_parts = [
            f"{marker} {status_label} group session on {session_row['session_date']} ({session_row['session_time']})."
        ]
        if missed_reason_text:
            note_parts.append(f"Reason: {missed_reason_text}")
        if raw_session_summary:
            note_parts.append(f"Content: {raw_session_summary}")
        if attendance_note_text:
            note_parts.append(f"Member note: {attendance_note_text}")
        note_content = ' '.join(part for part in note_parts if part).strip()
        existing_note = db.execute('''
            SELECT id
            FROM notes
            WHERE patient_id = ?
              AND content LIKE ?
            ORDER BY id DESC
            LIMIT 1
        ''', (pid, f'{marker}%')).fetchone()
        if existing_note:
            db.execute('''
                UPDATE notes
                SET note_date = ?, content = ?, is_missed_meeting = ?, missed_reason = ?, updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            ''', (session_row['session_date'], note_content, 1 if status_value == 'missed' else 0, missed_reason_text or None, existing_note['id']))
        else:
            db.execute('''
                INSERT INTO notes (patient_id, note_date, content, is_missed_meeting, missed_reason)
                VALUES (?, ?, ?, ?, ?)
            ''', (pid, session_row['session_date'], note_content, 1 if status_value == 'missed' else 0, missed_reason_text or None))

    for member in members:
        pid = int(member['patient_id'])
        status_value = (request.form.get(f'attendance_{pid}') or 'pending').strip().lower()
        if status_value not in ('present', 'missed', 'pending'):
            status_value = 'pending'
        absence_reason = (request.form.get(f'absence_reason_{pid}') or '').strip()
        notified_on_time = 1 if request.form.get(f'notified_on_time_{pid}') in ('1', 'true', 'on') else 0
        attendance_note = (request.form.get(f'attendance_note_{pid}') or '').strip()
        if status_value != 'missed':
            absence_reason = ''
            notified_on_time = 0

        attendance_payload[pid] = {
            'attendance_status': status_value,
            'absence_reason': absence_reason,
            'notified_on_time': notified_on_time,
            'attendance_note': attendance_note,
        }

        db.execute('''
            INSERT INTO group_session_attendance (session_id, patient_id, attendance_status, absence_reason, notified_on_time, attendance_note, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
            ON CONFLICT(session_id, patient_id)
            DO UPDATE SET attendance_status = excluded.attendance_status,
                          absence_reason = excluded.absence_reason,
                          notified_on_time = excluded.notified_on_time,
                          attendance_note = excluded.attendance_note,
                          updated_at = CURRENT_TIMESTAMP
        ''', (session_id, pid, status_value, absence_reason or None, notified_on_time, attendance_note or None))

        if session_status == 'completed':
            upsert_group_session_note(pid, status_value, attendance_note, absence_reason)

    formatted_summary = _build_group_session_summary_text(session_row, members, attendance_payload, raw_session_summary)

    db.execute('''
        UPDATE group_sessions
        SET session_summary = ?, status = ?
        WHERE id = ?
    ''', (formatted_summary or None, session_status, session_id))

    db.commit()
    flash('Session record saved.')
    destination_args = {'group_id': int(session_row['group_id'])}
    session_date = parse_date_safe(session_row['session_date'])
    session_time = parse_time_safe(session_row['session_time'])
    if session_date and session_time:
        session_end = datetime.combine(session_date, session_time) + timedelta(minutes=int(session_row['duration_minutes'] or 60))
        if session_end < datetime.now():
            destination_args['show_past'] = 'all'
        else:
            destination_args['show_upcoming'] = 'all'
    destination = url_for('group_detail', **destination_args)
    return redirect(f'{destination}#session-record-{session_id}')


@app.route('/patient/<int:patient_id>/group_attendance/<int:session_id>/update', methods=['POST'])
@login_required
def update_patient_group_attendance(patient_id, session_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    session_row = db.execute('SELECT * FROM group_sessions WHERE id = ?', (session_id,)).fetchone()
    patient_row = db.execute('SELECT id FROM patients WHERE id = ? AND COALESCE(is_deleted, 0) = 0', (patient_id,)).fetchone()
    if not session_row or not patient_row:
        flash('Attendance row could not be updated.')
        return redirect_to_patient_tab(patient_id, 'info')

    status_value = (request.form.get('attendance_status') or 'pending').strip().lower()
    if status_value not in ('present', 'missed', 'pending'):
        status_value = 'pending'
    absence_reason = (request.form.get('absence_reason') or '').strip()
    notified_on_time = 1 if request.form.get('notified_on_time') in ('1', 'true', 'on') else 0
    attendance_note = (request.form.get('attendance_note') or '').strip()
    session_summary = (request.form.get('session_summary') or '').strip()

    if status_value != 'missed':
        absence_reason = ''
        notified_on_time = 0

    db.execute('''
        UPDATE group_sessions
        SET session_summary = ?
        WHERE id = ?
    ''', (session_summary or None, session_id))

    db.execute('''
        INSERT INTO group_session_attendance (session_id, patient_id, attendance_status, absence_reason, notified_on_time, attendance_note, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
        ON CONFLICT(session_id, patient_id)
        DO UPDATE SET attendance_status = excluded.attendance_status,
                      absence_reason = excluded.absence_reason,
                      notified_on_time = excluded.notified_on_time,
                      attendance_note = excluded.attendance_note,
                      updated_at = CURRENT_TIMESTAMP
    ''', (session_id, patient_id, status_value, absence_reason or None, notified_on_time, attendance_note or None))

    if status_value == 'missed':
        marker = f"[Group Session #{session_id}]"
        note_content = f"{marker} Missed group session on {session_row['session_date']} ({session_row['session_time']})."
        if absence_reason:
            note_content = f"{note_content} Reason: {absence_reason}"
        existing_note = db.execute('''
            SELECT id
            FROM notes
            WHERE patient_id = ?
              AND content LIKE ?
            ORDER BY id DESC
            LIMIT 1
        ''', (patient_id, f'{marker}%')).fetchone()
        if existing_note:
            db.execute('''
                UPDATE notes
                SET note_date = ?, content = ?, is_missed_meeting = 1, missed_reason = ?, updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            ''', (session_row['session_date'], note_content, absence_reason or None, existing_note['id']))
        else:
            db.execute('''
                INSERT INTO notes (patient_id, note_date, content, is_missed_meeting, missed_reason)
                VALUES (?, ?, ?, 1, ?)
            ''', (patient_id, session_row['session_date'], note_content, absence_reason or None))

    db.commit()
    flash('Patient group attendance updated.')
    return redirect_to_patient_tab(patient_id, 'info')




@app.route('/patient/<int:patient_id>/quick_book', methods=('POST',))
@login_required
def quick_book_patient_appointment(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    date_raw = (request.form.get('date') or '').strip()
    time_raw = (request.form.get('time') or '').strip()
    end_time_raw = (request.form.get('end_time') or '').strip()
    meeting_type = (request.form.get('meeting_type') or 'in-person').strip() or 'in-person'
    recurrence_mode = (request.form.get('recurrence_mode') or 'auto').strip().lower()
    meeting_link = (request.form.get('meeting_link') or '').strip()
    meeting_title = (request.form.get('meeting_title') or '').strip()
    save_to_google = 1 if request.form.get('save_to_google') in ('1', 'true', 'on') else 0

    booking_date = parse_date_safe(date_raw)
    booking_time = parse_time_safe(time_raw)
    booking_end = parse_time_safe(end_time_raw)
    if not booking_date or not booking_time:
        flash('Valid date and time are required.', 'error')
        return redirect_to_patient_tab(patient_id, 'info')

    duration = 60
    if booking_end:
        start_minutes = booking_time.hour * 60 + booking_time.minute
        end_minutes = booking_end.hour * 60 + booking_end.minute
        if end_minutes > start_minutes:
            duration = end_minutes - start_minutes

    db = get_db()
    patient_row = db.execute('''
        SELECT id, status, patient_type
        FROM patients
        WHERE id = ? AND COALESCE(is_deleted, 0) = 0
    ''', (patient_id,)).fetchone()
    if not patient_row:
        flash('Patient not found.', 'error')
        return redirect(url_for('crm_dashboard'))

    start_dt = datetime.combine(booking_date, booking_time)
    end_dt = start_dt + timedelta(minutes=duration)
    conflict_message = has_time_conflict(db, booking_date, start_dt, end_dt)
    if conflict_message:
        flash(conflict_message, 'error')
        return redirect_to_patient_tab(patient_id, 'info')

    patient_type = (patient_row['patient_type'] or 'private').strip().lower()
    patient_status = (patient_row['status'] or '').strip().lower()
    default_recurring = 1 if patient_status == 'ongoing' and patient_type not in ('initial-intake', 'diagnosee') else 0
    if recurrence_mode not in ('auto', 'one-time', 'recurring'):
        recurrence_mode = 'auto'

    if recurrence_mode == 'one-time':
        is_recurring = 0
    elif recurrence_mode == 'recurring':
        if patient_type in ('initial-intake', 'diagnosee'):
            flash('Initial-intake patients can only be booked as one-time meetings.', 'error')
            return redirect_to_patient_tab(patient_id, 'info')
        is_recurring = 1
    else:
        is_recurring = default_recurring

    recurrence_interval = 1 if is_recurring else None
    recurrence_days = str(custom_weekday(booking_date)) if is_recurring else None
    recurrence_end_date = (booking_date + timedelta(days=365)).isoformat() if is_recurring else None
    recurrence_group_id = build_recurrence_group_id() if is_recurring else None
    meeting_platform = meeting_type if meeting_type in ('zoom', 'google-meet') else None

    db.execute('''
        INSERT INTO appointments (
            patient_id, appointment_date, appointment_time, duration_minutes,
            meeting_type, meeting_link, meeting_platform, meeting_title,
            save_to_google, status, is_recurring, recurrence_interval,
            recurrence_days, recurrence_end_date, recurrence_group_id
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'scheduled', ?, ?, ?, ?, ?)
    ''', (
        patient_id,
        booking_date.isoformat(),
        booking_time.strftime('%H:%M'),
        duration,
        meeting_type,
        meeting_link or None,
        meeting_platform,
        meeting_title or None,
        save_to_google,
        is_recurring,
        recurrence_interval,
        recurrence_days,
        recurrence_end_date,
        recurrence_group_id
    ))
    db.commit()

    if is_recurring:
        flash('Recurring weekly appointment booked for one year.', 'success')
    else:
        flash('Appointment booked.', 'success')
    return redirect_to_patient_tab(patient_id, 'info')

@app.route('/patient/<int:patient_id>/add_note', methods=('POST',))
@login_required
def add_note(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    content = request.form.get('content', '').strip()
    session_number = request.form.get('session_number', '').strip()
    note_date = request.form.get('note_date', '').strip()
    patient_appearance = request.form.get('patient_appearance', '').strip()
    behavior_flags = ','.join(request.form.getlist('behavior_flags'))
    mood_summary = request.form.get('mood_summary', '').strip()
    behavior_notes = request.form.get('behavior_notes', '').strip()
    link_url = request.form.get('link_url', '').strip()
    is_missed_meeting = 1 if request.form.get('is_missed_meeting') in ('1', 'true', 'on') else 0
    missed_reason = request.form.get('missed_reason', '').strip()
    share_with_patient = 1 if request.form.get('share_with_patient') in ('1', 'true', 'on') else 0
    if not is_missed_meeting:
        missed_reason = ''

    if content or is_missed_meeting:
        db = get_db()
        appointment_id = None
        if note_date:
            existing = db.execute(
                'SELECT id FROM appointments WHERE patient_id = ? AND appointment_date = ? LIMIT 1',
                (patient_id, note_date)
            ).fetchone()
            if existing:
                appointment_id = existing['id']

        cur = db.execute(
            '''INSERT INTO notes (patient_id, appointment_id, session_number, note_date, content,
                                  patient_appearance, behavior_checklist, mood_summary, behavior_notes,
                                  is_missed_meeting, missed_reason, link_url, share_with_patient)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
            (
                patient_id,
                appointment_id,
                session_number or None,
                note_date or None,
                content or 'Missed meeting documented.',
                patient_appearance or None,
                behavior_flags or None,
                mood_summary or None,
                behavior_notes or None,
                is_missed_meeting,
                missed_reason or None,
                link_url or None,
                share_with_patient
            )
        )
        note_id = cur.lastrowid
        db.commit()

        files = request.files.getlist('files')
        for file in files:
            if file and file.filename != '':
                filename = secure_filename(file.filename)
                filename = f"{secrets.token_hex(6)}_{filename}"

                patient_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'treatments', str(patient_id))
                if not os.path.exists(patient_dir):
                    os.makedirs(patient_dir)

                filepath = os.path.join(patient_dir, filename)
                file.save(filepath)

                db.execute('INSERT INTO files (patient_id, treatment_id, filename) VALUES (?, ?, ?)', (patient_id, note_id, filename))
                db.commit()

    else:
        flash('Content is required unless this is marked as a missed meeting.')

    return redirect_to_patient_tab(patient_id, 'notes')

@app.route('/note/<int:note_id>/edit', methods=('GET', 'POST'))
@login_required
def edit_note(note_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    note = db.execute('SELECT * FROM notes WHERE id = ?', (note_id,)).fetchone()
    if not note:
        flash('Note not found.', 'error')
        return redirect(url_for('crm_dashboard'))

    if request.method == 'GET':
        patient = db.execute('SELECT * FROM patients WHERE id = ?', (note['patient_id'],)).fetchone()
        return render_template('note_edit.html', note=note, patient=patient)

    content = request.form.get('content', '').strip()
    session_number = request.form.get('session_number', '').strip()
    note_date = request.form.get('note_date', '').strip()
    patient_appearance = request.form.get('patient_appearance', '').strip()
    behavior_flags = ','.join(request.form.getlist('behavior_flags'))
    mood_summary = request.form.get('mood_summary', '').strip()
    behavior_notes = request.form.get('behavior_notes', '').strip()
    link_url = request.form.get('link_url', '').strip()
    is_missed_meeting = 1 if request.form.get('is_missed_meeting') in ('1', 'true', 'on') else 0
    missed_reason = request.form.get('missed_reason', '').strip()
    share_with_patient = 1 if request.form.get('share_with_patient') in ('1', 'true', 'on') else 0
    if not is_missed_meeting:
        missed_reason = ''

    db = get_db()
    note = db.execute('SELECT * FROM notes WHERE id = ?', (note_id,)).fetchone()
    if note:
        db.execute(
            '''UPDATE notes
               SET content = ?, session_number = ?, note_date = ?, patient_appearance = ?,
                   behavior_checklist = ?, mood_summary = ?, behavior_notes = ?,
                   is_missed_meeting = ?, missed_reason = ?, link_url = ?,
                   share_with_patient = ?, updated_at = CURRENT_TIMESTAMP
               WHERE id = ?''',
            (
                content or 'Missed meeting documented.',
                session_number or None,
                note_date or None,
                patient_appearance or None,
                behavior_flags or None,
                mood_summary or None,
                behavior_notes or None,
                is_missed_meeting,
                missed_reason or None,
                link_url or None,
                share_with_patient,
                note_id
            )
        )
        db.commit()
        return redirect_to_patient_tab(note['patient_id'], 'notes')
    flash('Note not found.', 'error')
    return redirect(url_for('crm_dashboard'))


@app.route('/note/<int:note_id>/delete', methods=('POST',))
@login_required
def delete_note(note_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    note = db.execute('SELECT id, patient_id FROM notes WHERE id = ?', (note_id,)).fetchone()
    if note is None:
        flash('Note not found.', 'error')
        return redirect(request.referrer or url_for('crm_dashboard'))

    db.execute('DELETE FROM notes WHERE id = ?', (note_id,))
    db.commit()
    flash('Meeting log deleted.')
    return redirect_to_patient_tab(note['patient_id'], 'notes')

@app.route('/patient/<int:patient_id>/update_treatment_plan', methods=('POST',))
@login_required
def update_treatment_plan(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403
    plan = (request.form.get('treatment_plan') or '').strip()
    db = get_db()
    db.execute('UPDATE patients SET treatment_plan = ? WHERE id = ?', (plan or None, patient_id))
    db.commit()
    flash('Treatment plan updated.')
    return redirect_to_patient_tab(patient_id, 'info')

@app.route('/patient/<int:patient_id>/add_goal', methods=('POST',))
@login_required
def add_goal(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    description = request.form.get('description', '').strip()
    if description:
        db = get_db()
        db.execute('INSERT INTO goals (patient_id, description) VALUES (?, ?)', (patient_id, description))
        db.commit()
    return redirect_to_patient_tab(patient_id, 'info')

@app.route('/goal/<int:goal_id>/toggle_status', methods=('POST',))
@login_required
def toggle_goal_status(goal_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    goal = db.execute('SELECT * FROM goals WHERE id = ?', (goal_id,)).fetchone()
    if goal:
        new_status = 'achieved' if goal['status'] == 'active' else 'active'
        db.execute('UPDATE goals SET status = ? WHERE id = ?', (new_status, goal_id))
        db.commit()
        return redirect_to_patient_tab(goal['patient_id'], 'info')
    return "Goal not found", 404

@app.route('/goal/<int:goal_id>/delete', methods=('POST',))
@login_required
def delete_goal(goal_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    goal = db.execute('SELECT * FROM goals WHERE id = ?', (goal_id,)).fetchone()
    if not goal:
        return "Goal not found", 404
    patient_id = goal['patient_id']
    db.execute('DELETE FROM goals WHERE id = ?', (goal_id,))
    db.commit()
    return redirect_to_patient_tab(patient_id, 'info')

@app.route('/goal/<int:goal_id>/patient_achieve', methods=('POST',))
@login_required
def patient_achieve_goal(goal_id):
    if current_user.role != 'patient':
        return "Unauthorized", 403
    db = get_db()
    goal = db.execute('SELECT * FROM goals WHERE id = ?', (goal_id,)).fetchone()
    if not goal or int(goal['patient_id']) != int(current_user.patient_id or 0):
        return "Goal not found", 404
    new_status = 'achieved' if goal['status'] == 'active' else 'active'
    db.execute('UPDATE goals SET status = ? WHERE id = ?', (new_status, goal_id))
    db.commit()
    if new_status == 'achieved':
        flash('Goal marked as achieved! Great progress.')
    else:
        flash('Goal re-opened.')
    return redirect(url_for('patient_home'))

@app.route('/goal/<int:goal_id>/edit', methods=('POST',))
@login_required
def edit_goal(goal_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    description = (request.form.get('description') or '').strip()
    db = get_db()
    goal = db.execute('SELECT * FROM goals WHERE id = ?', (goal_id,)).fetchone()
    if not goal:
        return "Goal not found", 404
    if description:
        db.execute('UPDATE goals SET description = ? WHERE id = ?', (description, goal_id))
        db.commit()
    else:
        flash('Goal description cannot be empty.')
    return redirect_to_patient_tab(goal['patient_id'], 'info')

@app.route('/patient/<int:patient_id>/add_file', methods=('POST',))
@login_required
def add_file(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    if 'file' not in request.files:
        flash('No file part')
        return redirect_to_patient_tab(patient_id, 'notes')
    file = request.files['file']
    if file.filename == '':
        flash('No selected file')
        return redirect_to_patient_tab(patient_id, 'notes')
    if file:
        filename = secure_filename(file.filename)
        if not _allowed_upload(filename, ALLOWED_UPLOAD_EXTENSIONS):
            flash('File type not allowed. Accepted: docx, pdf, txt, png, jpg, gif, xlsx, csv, webp.')
            return redirect_to_patient_tab(patient_id, 'notes')
        filename = f"{secrets.token_hex(6)}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        db = get_db()
        db.execute('INSERT INTO files (patient_id, filename) VALUES (?, ?)', (patient_id, filename))
        db.commit()

        if filename.endswith('.docx'):
            # Attempt to parse document
            try:
                doc = Document(filepath)
                text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

                # Split the text by meeting header to support multiple entries
                meeting_pattern = re.compile(r'(?:Meeting #|פגישה מספר)[:\s]*\w+', re.IGNORECASE)
                matches = list(meeting_pattern.finditer(text))

                if not matches:
                    blocks = [text]
                else:
                    blocks = []
                    for i, match in enumerate(matches):
                        start_idx = match.start()
                        end_idx = matches[i+1].start() if i + 1 < len(matches) else len(text)
                        blocks.append(text[start_idx:end_idx])

                notes_created = 0
                notes_review = 0

                for block in blocks:
                    parsed_date = None
                    meeting_no_match = re.search(r'(?:Meeting #|פגישה מספר)[:\s]*(\w+)', block, re.IGNORECASE)
                    date_match = re.search(r'(?:Date|תאריך)[:\s]*([\d\./\-]+)', block, re.IGNORECASE)
                    content_match = re.search(r'(?:Content|תוכן)[:\s]*(.*)', block, re.IGNORECASE | re.DOTALL)

                    meeting_no = meeting_no_match.group(1).strip() if meeting_no_match else None
                    date_str = date_match.group(1).strip() if date_match else None
                    content = content_match.group(1).strip() if content_match else block.strip()

                    needs_review = False
                    if not meeting_no or not date_str:
                        needs_review = True

                    appointment_id = None
                    if date_str:
                        try:
                            if '.' in date_str or '/' in date_str:
                                parts = re.split(r'[\./]', date_str)
                                if len(parts) == 3:
                                    d, m, y = int(parts[0]), int(parts[1]), int(parts[2])
                                    if y < 100:
                                        y += 2000
                                    parsed_date = f"{y:04d}-{m:02d}-{d:02d}"
                            if not parsed_date:
                                if re.match(r'\d{4}-\d{2}-\d{2}', date_str):
                                    parsed_date = date_str

                            if parsed_date:
                                appt = db.execute('SELECT id FROM appointments WHERE patient_id = ? AND appointment_date = ?', (patient_id, parsed_date)).fetchone()
                                if appt:
                                    appointment_id = appt['id']
                        except Exception as e:
                            print("Error parsing date:", e)
                            needs_review = True

                    db.execute('''INSERT INTO notes (patient_id, appointment_id, session_number, note_date, needs_review, content)
                                  VALUES (?, ?, ?, ?, ?, ?)''',
                               (patient_id, appointment_id, meeting_no, parsed_date, needs_review, content))
                    if needs_review:
                        notes_review += 1
                    else:
                        notes_created += 1

                db.commit()
                if notes_review > 0:
                    flash(f'DOCX parsed. {notes_created} notes created, {notes_review} marked for review.')
                else:
                    flash(f'DOCX parsed successfully. {notes_created} notes created.')
            except Exception as e:
                print(f"Error parsing DOCX: {e}")
                flash('Error parsing DOCX file.')

    return redirect_to_patient_tab(patient_id, 'notes')

@app.route('/patient/<int:patient_id>/export', methods=('GET',))
@login_required
def export_patient_history(patient_id):

    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    patient = db.execute('SELECT * FROM patients WHERE id = ?', (patient_id,)).fetchone()
    if patient is None:
        return "Patient not found", 404

    appointments = [dict(row) for row in db.execute('SELECT * FROM appointments WHERE patient_id = ? ORDER BY appointment_date ASC, appointment_time ASC', (patient_id,)).fetchall()]
    notes = [dict(row) for row in db.execute('SELECT * FROM notes WHERE patient_id = ? ORDER BY created_at ASC', (patient_id,)).fetchall()]
    receipts = [dict(row) for row in db.execute('SELECT * FROM receipts WHERE patient_id = ? ORDER BY created_at ASC', (patient_id,)).fetchall()]

    data = {
        'patient': dict(patient),
        'appointments': appointments,
        'notes': notes,
        'receipts': receipts
    }

    response = Response(json.dumps(data, indent=4), mimetype='application/json')
    response.headers['Content-Disposition'] = f'attachment; filename=patient_{patient_id}_history.json'
    return response


@app.route('/patient/<int:patient_id>/import', methods=('POST',))
@login_required
def import_patient_history(patient_id):

    if current_user.role != 'admin':
        return "Unauthorized", 403

    if 'file' not in request.files:
        flash('No file part')
        return redirect_to_patient_tab(patient_id, 'notes')

    file = request.files['file']
    if file.filename == '':
        flash('No selected file')
        return redirect_to_patient_tab(patient_id, 'notes')

    if file and file.filename.endswith('.json'):
        try:
            data = json.load(file)
            db = get_db()

            appointments_added = 0
            notes_added = 0
            receipts_added = 0

            if isinstance(data, list):
                appointments_added, notes_added, receipts_added = _import_flat_patient_history(db, patient_id, data)
            else:


                # Import appointments
                appt_id_map = {}
                # Sort appointments by date and time
                sorted_appts = sorted(data.get('appointments', []), key=lambda x: (x.get('appointment_date', ''), x.get('appointment_time', '')))
                for appt in sorted_appts:
                    # Check for existing
                    existing = db.execute('SELECT id FROM appointments WHERE patient_id = ? AND appointment_date = ? AND appointment_time = ?',
                        (patient_id, appt.get('appointment_date'), appt.get('appointment_time'))).fetchone()
                    if not existing:
                        cursor = db.execute('''INSERT INTO appointments
                            (patient_id, appointment_date, appointment_time, cost, duration_minutes, is_recurring, recurrence_interval, recurrence_days, meeting_type, meeting_link, status, recurrence_end_date, recurrence_count)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                            (patient_id, appt.get('appointment_date'), appt.get('appointment_time'), appt.get('cost'), appt.get('duration_minutes'),
                             appt.get('is_recurring'), appt.get('recurrence_interval'), appt.get('recurrence_days'), appt.get('meeting_type'),
                             appt.get('meeting_link'), appt.get('status'), appt.get('recurrence_end_date'), appt.get('recurrence_count')))
                        appt_id_map[appt.get('id')] = cursor.lastrowid
                        appointments_added += 1
                    else:
                        appt_id_map[appt.get('id')] = existing['id']

                # Import notes sorted by date and meeting number.
                sorted_notes = sorted(
                    data.get('notes', []),
                    key=lambda x: (
                        x.get('note_date') or x.get('date') or x.get('created_at', ''),
                        str(x.get('session_number') or x.get('meeting_number') or '')
                    )
                )
                for note in sorted_notes:
                    new_appt_id = appt_id_map.get(note.get('appointment_id')) if note.get('appointment_id') else None
                    session_number = _normalize_session_number(note.get('session_number') or note.get('meeting_number'))
                    note_date = (note.get('note_date') or note.get('date') or '').strip() or None
                    content_text = (note.get('content') or '').strip()
                    appearance_text = (note.get('patient_appearance') or '').strip()
                    checklist_text = note.get('behavior_checklist')
                    if isinstance(checklist_text, list):
                        checklist_text = ','.join([str(i).strip() for i in checklist_text if str(i).strip()])
                    checklist_text = (checklist_text or '').strip()
                    mood_summary = (note.get('mood_summary') or '').strip()
                    behavior_notes = (note.get('behavior_notes') or '').strip()

                    if not session_number and not _has_meaningful_note_information(
                        content_text,
                        mood_summary,
                        behavior_notes,
                        appearance_text,
                        checklist_text,
                    ):
                        continue

                    if not content_text:
                        content_text = mood_summary or behavior_notes or appearance_text
                    if not content_text:
                        continue

                    db.execute('''INSERT INTO notes
                        (patient_id, appointment_id, session_number, note_date, content, patient_appearance,
                         behavior_checklist, mood_summary, behavior_notes, needs_review, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                        (
                            patient_id,
                            new_appt_id,
                            session_number,
                            note_date,
                            content_text,
                            appearance_text,
                            checklist_text,
                            mood_summary,
                            behavior_notes,
                            note.get('needs_review'),
                            note.get('created_at')
                        ))
                    notes_added += 1

                # Import receipts
                receipts_data = data.get('receipts', [])
                if receipts_data:
                    receipt_tuples = [
                        (patient_id, r.get('amount'), r.get('description'), r.get('created_at'))
                        for r in receipts_data
                    ]
                    db.executemany('''INSERT INTO receipts
                        (patient_id, amount, description, created_at)
                        VALUES (?, ?, ?, ?)''', receipt_tuples)
                    receipts_added += len(receipts_data)
                appointments_added, notes_added, receipts_added = _import_structured_patient_history(db, patient_id, data)

            db.commit()
            flash(f'History imported: {appointments_added} appointments, {notes_added} notes, {receipts_added} receipts added.')
        except Exception as e:
            print("Import error:", e)
            flash('Error parsing JSON file.')
    else:
        flash('Please upload a JSON file.')

    return redirect_to_patient_tab(patient_id, 'notes')


@app.route('/uploads/<name>')
@login_required
def download_file(name):
    # Check if user has access to this file.
    # For now, allow admin and the patient who owns the file.
    # But finding the owner of a file from filename is hard if filenames aren't unique or mapped.
    # The 'files' table maps filename to patient_id.
    safe_name = secure_filename(name)
    if safe_name != name:
        return "Invalid filename", 400
    name = safe_name

    db = get_db()
    file_record = db.execute('SELECT patient_id, treatment_id FROM files WHERE filename = ?', (name,)).fetchone()

    if not file_record:
        profile_owner = db.execute('SELECT id FROM patients WHERE profile_image = ?', (name,)).fetchone()
        if current_user.role == 'admin':
            return send_from_directory(app.config['UPLOAD_FOLDER'], name)
        if profile_owner and current_user.role == 'patient' and current_user.patient_id == profile_owner['id']:
            return send_from_directory(app.config['UPLOAD_FOLDER'], name)
        return "File not found or access denied", 403

    if current_user.role == 'admin' or (current_user.role == 'patient' and current_user.patient_id == file_record['patient_id']):
        if file_record['treatment_id']:
            patient_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'treatments', str(file_record['patient_id']))
            return send_from_directory(patient_dir, name)
        return send_from_directory(app.config['UPLOAD_FOLDER'], name)

    return "Access denied", 403




@app.route('/groups/sessions/<int:session_id>/delete', methods=['POST'])
@login_required
def delete_group_session(session_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    existing = db.execute('SELECT id, group_id FROM group_sessions WHERE id = ?', (session_id,)).fetchone()
    if not existing:
        flash('Group session not found.')
        return redirect(url_for('groups_dashboard'))

    db.execute('DELETE FROM group_session_attendance WHERE session_id = ?', (session_id,))
    db.execute('DELETE FROM group_sessions WHERE id = ?', (session_id,))
    db.commit()
    flash('Group session deleted.')
    return redirect(url_for('group_detail', group_id=int(existing['group_id'])))


@app.route('/api/groups/sessions/<int:session_id>/delete', methods=['POST'])
@login_required
def api_delete_group_session(session_id):
    if current_user.role != 'admin':
        return jsonify({'status': 'error', 'message': 'Unauthorized'}), 403

    db = get_db()
    existing = db.execute('SELECT id FROM group_sessions WHERE id = ?', (session_id,)).fetchone()
    if not existing:
        return jsonify({'status': 'error', 'message': 'Group session not found.'}), 404

    db.execute('DELETE FROM group_session_attendance WHERE session_id = ?', (session_id,))
    db.execute('DELETE FROM group_sessions WHERE id = ?', (session_id,))
    db.commit()
    return jsonify({'status': 'success'})


@app.route('/api/groups/sessions/<int:session_id>/link_supervision', methods=['POST'])
@login_required
def api_link_group_session_supervision(session_id):
    if current_user.role != 'admin':
        return jsonify({'status': 'error', 'message': 'Unauthorized'}), 403

    db = get_db()
    session_row = db.execute('SELECT id, group_id FROM group_sessions WHERE id = ?', (session_id,)).fetchone()
    if not session_row:
        return jsonify({'status': 'error', 'message': 'Group session not found.'}), 404

    sup_id_raw = (request.form.get('supervision_id') or '').strip()
    if not sup_id_raw:
        db.execute('UPDATE group_sessions SET supervision_id = NULL WHERE id = ?', (session_id,))
        db.commit()
        return jsonify({'status': 'success'})

    try:
        sup_id = int(sup_id_raw)
    except ValueError:
        return jsonify({'status': 'error', 'message': 'Invalid supervision id.'}), 400

    supervision_row = db.execute(
        'SELECT id FROM supervisions WHERE id = ? AND group_id = ?',
        (sup_id, int(session_row['group_id']))
    ).fetchone()
    if not supervision_row:
        return jsonify({'status': 'error', 'message': 'Supervision record not found for this group.'}), 404

    db.execute('UPDATE group_sessions SET supervision_id = ? WHERE id = ?', (sup_id, session_id))
    db.commit()
    return jsonify({'status': 'success'})







@app.route('/patient/<int:patient_id>/convert', methods=('POST',))
@login_required
def convert_patient(patient_id):

    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()

    start_date = request.form.get('start_date', '').strip()
    time = request.form.get('time', '').strip()
    duration = request.form.get('duration', 60)
    interval = request.form.get('interval', 1)
    cost = request.form.get('cost', 0)
    meeting_type = request.form.get('meeting_type', 'in-person')
    meeting_link = request.form.get('meeting_link', '')

    # Validate required fields
    if not start_date or not time:
        flash('Start date and time are required!', 'error')
        return redirect_to_patient_tab(patient_id, 'info')

    # Validate date and time formats
    try:
        datetime.fromisoformat(start_date)
        datetime.strptime(time, '%H:%M')
    except ValueError as e:
        flash(f'Invalid date or time format: {str(e)}', 'error')
        return redirect_to_patient_tab(patient_id, 'info')

    # Convert types
    try:
        duration = int(duration)
        interval = int(interval)
        cost = float(cost) if cost else 0
    except (ValueError, TypeError):
        flash('Invalid duration, interval, or cost value!', 'error')
        return redirect_to_patient_tab(patient_id, 'info')

    # Get recurrence limit
    limit_type = request.form.get('recurrence_limit_type')
    recurrence_end_date = None
    recurrence_count = None
    
    if limit_type == 'date':
        recurrence_end_date = request.form.get('recurrence_end_date', '').strip()
        if recurrence_end_date:
            try:
                datetime.fromisoformat(recurrence_end_date)
            except ValueError:
                flash('Invalid recurrence end date!', 'error')
                return redirect_to_patient_tab(patient_id, 'info')
    elif limit_type == 'count':
        try:
            recurrence_count = int(request.form.get('recurrence_count', 12))
            if recurrence_count <= 0:
                recurrence_count = 12
        except ValueError:
            recurrence_count = 12

    # Get checked days (multiple values)
    days_list = request.form.getlist('days')
    days_str = ','.join(str(d) for d in days_list if d.strip().isdigit()) if days_list else None

    try:
        # Update patient status
        db.execute("UPDATE patients SET status = 'ongoing' WHERE id = ?", (patient_id,))

        # Create recurring appointment
        db.execute('''INSERT INTO appointments
                      (patient_id, appointment_date, appointment_time, cost, duration_minutes, 
                       is_recurring, recurrence_interval, recurrence_days, meeting_type, meeting_link, 
                                             recurrence_end_date, recurrence_count, recurrence_group_id)
                                            VALUES (?, ?, ?, ?, ?, 1, ?, ?, ?, ?, ?, ?, ?)''',
                   (patient_id, start_date, time, cost, duration, interval, days_str, 
                                        meeting_type, meeting_link, recurrence_end_date, recurrence_count, build_recurrence_group_id()))

        # Log the action
        patient = db.execute('SELECT name FROM patients WHERE id = ?', (patient_id,)).fetchone()
        if patient:
            details = f"Patient {patient['name']} converted to ongoing status with recurring appointment starting {start_date} at {time}."
            db.execute('INSERT INTO audit_logs (patient_id, action, details) VALUES (?, ?, ?)', 
                       (patient_id, 'convert', details))

        db.commit()
        flash('Patient converted to ongoing successfully with recurring appointments.', 'success')
    except sqlite3.IntegrityError as e:
        db.rollback()
        flash(f'Database error: {str(e)}', 'error')
    except Exception as e:
        db.rollback()
        flash(f'Error converting patient: {str(e)}', 'error')

    return redirect_to_patient_tab(patient_id, 'info')

@app.route('/patient/<int:patient_id>/delete', methods=('POST',))
@login_required
def delete_patient(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    patient = db.execute('SELECT * FROM patients WHERE id = ?', (patient_id,)).fetchone()
    if not patient:
        return "Patient not found", 404

    # Get deletion reason from request
    deletion_reason = request.form.get('deletion_reason', '').strip()
    if not deletion_reason:
        flash('Deletion reason is required.', 'error')
        return redirect(request.referrer or url_for('crm_dashboard', status='all'))

    db.execute('''
        UPDATE patients
        SET is_deleted = 1, deleted_at = CURRENT_TIMESTAMP, deleted_reason = ?, status = 'archived'
        WHERE id = ?
    ''', (deletion_reason, patient_id))
    db.execute('UPDATE users SET is_active = 0 WHERE patient_id = ?', (patient_id,))
    db.execute('INSERT INTO audit_logs (patient_id, action, details) VALUES (?, ?, ?)',
               (patient_id, 'delete', f"Patient {patient['name']} marked as deleted. Reason: {deletion_reason}"))
    db.commit()
    flash(f'Patient "{patient["name"]}" moved to deleted records.')
    return redirect(request.referrer or url_for('crm_dashboard', status='all'))

@app.route('/admin/profile/name', methods=('POST',))
@login_required
def update_admin_profile_name():
    if current_user.role != 'admin':
        return "Unauthorized", 403

    new_name = request.form.get('display_name', '').strip()
    if not new_name:
        flash('Admin name is required.')
        return redirect(request.referrer or url_for('crm_dashboard'))

    db = get_db()
    db.execute('UPDATE users SET display_name = ? WHERE id = ?', (new_name, current_user.id))
    db.commit()
    flash('Admin display name updated.')
    return redirect(request.referrer or url_for('crm_dashboard'))





@app.route('/api/groups/sessions/<int:session_id>/update', methods=['POST'])
@login_required
def api_update_group_session(session_id):
    if current_user.role != 'admin':
        return jsonify({'status': 'error', 'message': 'Unauthorized'}), 403

    db = get_db()
    existing = db.execute('SELECT * FROM group_sessions WHERE id = ?', (session_id,)).fetchone()
    if not existing:
        return jsonify({'status': 'error', 'message': 'Group session not found.'}), 404

    session_date = (request.form.get('session_date') or '').strip()
    session_time = (request.form.get('session_time') or '').strip()
    end_time_raw = (request.form.get('end_time') or '').strip()
    title = (request.form.get('title') or '').strip()
    facilitator = (request.form.get('facilitator') or '').strip()
    meeting_type = (request.form.get('meeting_type') or 'in-person').strip() or 'in-person'
    meeting_link = (request.form.get('meeting_link') or '').strip()
    apply_scope = (request.form.get('apply_scope') or 'single').strip().lower()

    parsed_date = parse_date_safe(session_date)
    parsed_time = parse_time_safe(session_time)
    parsed_end = parse_time_safe(end_time_raw)
    if not parsed_date or not parsed_time:
        return jsonify({'status': 'error', 'message': 'Valid date and start time are required.'}), 400

    duration = int(existing['duration_minutes'] or 60)
    if parsed_end:
        start_minutes = parsed_time.hour * 60 + parsed_time.minute
        end_minutes = parsed_end.hour * 60 + parsed_end.minute
        if end_minutes > start_minutes:
            duration = end_minutes - start_minutes

    existing_date = parse_date_safe(existing['session_date'])
    if not existing_date:
        return jsonify({'status': 'error', 'message': 'Stored session date is invalid.'}), 500

    apply_future = apply_scope == 'future' and existing['series_id']
    target_rows = [existing]
    if apply_future:
        target_rows = db.execute('''
            SELECT *
            FROM group_sessions
            WHERE series_id = ? AND session_date >= ?
            ORDER BY session_date ASC, session_time ASC
        ''', (existing['series_id'], existing['session_date'])).fetchall()

    day_delta = (parsed_date - existing_date).days
    for row in target_rows:
        row_date = parse_date_safe(row['session_date'])
        if not row_date:
            return jsonify({'status': 'error', 'message': 'Existing recurrence row has invalid date.'}), 500
        updated_date = row_date + timedelta(days=day_delta) if apply_future else parsed_date
        start_dt = datetime.combine(updated_date, parsed_time)
        end_dt = start_dt + timedelta(minutes=duration)
        conflict_message = has_time_conflict(
            db,
            updated_date,
            start_dt,
            end_dt,
            exclude_group_session_id=int(row['id'])
        )
        if conflict_message:
            return jsonify({'status': 'error', 'message': f"{conflict_message} ({updated_date.isoformat()})"}), 409

    update_data = []
    for row in target_rows:
        row_date = parse_date_safe(row['session_date'])
        updated_date = row_date + timedelta(days=day_delta) if apply_future else parsed_date
        update_data.append((
            updated_date.isoformat(),
            parsed_time.strftime('%H:%M'),
            duration,
            title or None,
            facilitator or None,
            meeting_type,
            meeting_link or None,
            row['id']
        ))

    if update_data:
        db.executemany('''
            UPDATE group_sessions
            SET session_date = ?, session_time = ?, duration_minutes = ?,
                title = ?, facilitator = ?, meeting_type = ?, meeting_link = ?
            WHERE id = ?
        ''', update_data)

    if apply_future and existing['series_id']:
        db.execute('''
            UPDATE group_session_series
            SET start_date = ?, start_time = ?, duration_minutes = ?,
                title = ?, facilitator = ?, meeting_type = ?, meeting_link = ?
            WHERE id = ?
        ''', (
            parsed_date.isoformat(),
            parsed_time.strftime('%H:%M'),
            duration,
            title or None,
            facilitator or None,
            meeting_type,
            meeting_link or None,
            existing['series_id']
        ))

    db.commit()
    return jsonify({'status': 'success'})

@app.route('/patient/<int:patient_id>/upload-photo', methods=('POST',))
@login_required
def upload_patient_photo(patient_id):
    if current_user.role != 'admin' and not (current_user.role == 'patient' and current_user.patient_id == patient_id):
        return "Unauthorized", 403

    photo = request.files.get('photo')
    if photo is None or not (photo.filename or '').strip():
        flash('Please choose an image file.')
        if current_user.role == 'patient':
            return redirect(url_for('patient_home'))
        return redirect_to_patient_tab(patient_id, 'info')

    extension = os.path.splitext(photo.filename)[1].lower()
    if extension not in {'.png', '.jpg', '.jpeg', '.gif', '.webp'}:
        flash('Please upload a PNG, JPG, GIF, or WEBP image.')
        if current_user.role == 'patient':
            return redirect(url_for('patient_home'))
        return redirect_to_patient_tab(patient_id, 'info')

    filename = secure_filename(f'patient_{patient_id}_{secrets.token_hex(6)}{extension}')
    save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    photo.save(save_path)

    db = get_db()
    existing = db.execute('SELECT profile_image FROM patients WHERE id = ?', (patient_id,)).fetchone()
    old_image = existing['profile_image'] if existing else None
    db.execute('UPDATE patients SET profile_image = ? WHERE id = ?', (filename, patient_id))
    db.commit()

    if old_image:
        old_path = os.path.join(app.config['UPLOAD_FOLDER'], old_image)
        if os.path.exists(old_path):
            try:
                os.remove(old_path)
            except OSError:
                pass

    flash('Profile picture updated.')
    if current_user.role == 'patient':
        return redirect(url_for('patient_home'))
    return redirect_to_patient_tab(patient_id, 'info')


@app.route('/patient/<int:patient_id>/edit_info', methods=('POST',))
@login_required
def update_patient_info(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    background = request.form.get('background')
    treatment_info = request.form.get('treatment_info')

    db = get_db()
    existing = db.execute(
        'SELECT background, treatment_info, intake_questionnaire, intake_assessment FROM patients WHERE id = ?',
        (patient_id,)
    ).fetchone()
    existing_intake_data = {}
    if existing is not None:
        existing_intake_data = parse_intake_questionnaire(
            existing['intake_questionnaire'],
            existing['intake_assessment']
        )
        if background is None:
            background = existing['background'] or ''
        if treatment_info is None:
            treatment_info = existing['treatment_info'] or ''

    intake_data = intake_data_from_request(request.form, existing_data=existing_intake_data)
    serialized_intake = json.dumps(intake_data, ensure_ascii=False, indent=2) if intake_data is not None else None
    serialized_assessment = serialize_intake_assessment(intake_data) if intake_data is not None else None

    background = background or ''
    treatment_info = treatment_info or ''

    if intake_data is None:
        db.execute('UPDATE patients SET background = ?, treatment_info = ? WHERE id = ?',
                   (background, treatment_info, patient_id))
    else:
        db.execute('''
            UPDATE patients
            SET background = ?, treatment_info = ?, intake_assessment = ?, intake_questionnaire = ?
            WHERE id = ?
        ''', (background, treatment_info, serialized_assessment or None, serialized_intake or None, patient_id))
    db.commit()
    flash('Patient information updated.')
    return redirect_to_patient_tab(patient_id, 'info')


@app.route('/patient/<int:patient_id>/intake_docx', methods=('GET',))
@login_required
def export_patient_intake_docx(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    patient = db.execute('SELECT id, name, intake_questionnaire, intake_assessment FROM patients WHERE id = ?', (patient_id,)).fetchone()
    if patient is None:
        return "Patient not found", 404

    intake_data = parse_intake_questionnaire(patient['intake_questionnaire'], patient['intake_assessment'])
    if not intake_data:
        flash('No intake form data found for export.')
        return redirect_to_patient_tab(patient_id, 'intake')

    language = (request.args.get('lang') or 'en').strip().lower()
    if language not in {'en', 'he'}:
        language = 'en'

    document = build_intake_docx(patient['name'], intake_data, language=language)
    safe_name = secure_filename(patient['name'] or f'patient_{patient_id}')
    if not safe_name:
        safe_name = f'patient_{patient_id}'
    output_name = f'intake_{safe_name}.docx'
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name=output_name,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/patient/<int:patient_id>/generate_background', methods=('POST',))
@login_required
def generate_patient_background(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    patient = db.execute('SELECT id, name FROM patients WHERE id = ?', (patient_id,)).fetchone()
    if patient is None:
        return "Patient not found", 404

    background = build_patient_background_from_notes(db, patient_id, patient['name'])
    db.execute('UPDATE patients SET background = ? WHERE id = ?', (background, patient_id))
    db.commit()
    flash('AI background generated.')
    return redirect_to_patient_tab(patient_id, 'info')

@app.route('/patient/<int:patient_id>/edit', methods=('GET', 'POST'))
@login_required
def edit_patient(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    patient = db.execute('SELECT * FROM patients WHERE id = ?', (patient_id,)).fetchone()
    if patient is None:
        return "Patient not found", 404

    if request.method == 'POST':
        name = request.form['name']
        status = _normalize_patient_status(request.form['status'])
        email = request.form.get('email')
        phone = request.form.get('phone')
        birth_date = request.form.get('birth_date') or None
        id_number = (request.form.get('id_number') or '').strip() or None
        can_self_schedule = 1 if request.form.get('can_self_schedule') else 0
        reminder_email_enabled = 1 if request.form.get('reminder_email_enabled') else 0
        reminder_sms_enabled = 1 if request.form.get('reminder_sms_enabled') else 0
        patient_type = request.form.get('patient_type', 'private')
        if patient_type not in ('private', 'residency', 'initial-intake', 'diagnosee', 'group'):
            patient_type = 'private'
        has_intake_tab = int(patient['has_intake_tab'] or 0)
        has_questionnaire_tab = int(patient['has_questionnaire_tab'] or 0)
        if patient_type in ('initial-intake', 'diagnosee'):
            has_intake_tab = 1
        if patient_type == 'diagnosee':
            has_questionnaire_tab = 1
        if patient_type in ('initial-intake', 'diagnosee'):
            intake_assessment = request.form.get('intake_assessment')
            intake_questionnaire = request.form.get('intake_questionnaire')
            if intake_assessment is None:
                intake_assessment = patient['intake_assessment'] or ''
            else:
                intake_assessment = intake_assessment.strip()
            if intake_questionnaire is None:
                intake_questionnaire = patient['intake_questionnaire'] or ''
            else:
                intake_questionnaire = intake_questionnaire.strip()
        else:
            intake_assessment = ''
            intake_questionnaire = ''

        if not name:
            flash('Name is required!')
        else:
            field_errors = _validate_patient_fields(name, phone=phone, birth_date=birth_date, email=email)
            for err in field_errors:
                flash(err)
            if not field_errors:
                treatment_method = request.form.get('treatment_method', '').strip() or None
                db.execute('''UPDATE patients
                              SET name = ?, status = ?, email = ?, phone = ?, birth_date = ?, id_number = ?, can_self_schedule = ?,
                                  reminder_email_enabled = ?, reminder_sms_enabled = ?,
                                  patient_type = ?, has_intake_tab = ?, has_questionnaire_tab = ?, intake_assessment = ?, intake_questionnaire = ?,
                                  treatment_method = ?
                              WHERE id = ?''',
                             (name, status, email, phone, birth_date, id_number, can_self_schedule,
                              reminder_email_enabled, reminder_sms_enabled, patient_type,
                              has_intake_tab, has_questionnaire_tab, intake_assessment or None, intake_questionnaire or None, treatment_method, patient_id))
                db.commit()
                flash('Patient updated successfully.')
                return redirect(url_for('patient_detail', patient_id=patient_id))

    db_r = get_db()
    treatment_method_options = [r['label'] for r in db_r.execute(
        'SELECT label FROM treatment_method_options ORDER BY display_order ASC, label ASC'
    ).fetchall()]
    return render_template('edit_patient.html', patient=patient, treatment_method_options=treatment_method_options)

@app.route('/patient/<int:patient_id>/access', methods=('POST',))
@login_required
def manage_access(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    username = request.form['username']
    password = request.form['password']

    if not username or not password:
        flash('Username and password are required.')
        return redirect_to_patient_tab(patient_id, 'info')

    db = get_db()

    patient = db.execute('SELECT * FROM patients WHERE id = ?', (patient_id,)).fetchone()
    if patient is None:
        flash('Patient not found.')
        return redirect_to_patient_tab(patient_id, 'info')

    # Check if user exists for this patient
    existing_user = db.execute('SELECT * FROM users WHERE patient_id = ?', (patient_id,)).fetchone()

    hashed_pw = generate_password_hash(password)
    action = None

    if existing_user:
        try:
            db.execute(
                'UPDATE users SET username = ?, password_hash = ?, force_password_change = 1, is_active = 1 WHERE id = ?',
                (username, hashed_pw, existing_user['id'])
            )
            db.commit()
            flash('User access updated. Patient will be required to change their password on first login.')
            action = 'updated'
        except sqlite3.IntegrityError:
            flash('Username already taken.')
    else:
        try:
            db.execute(
                'INSERT INTO users (username, password_hash, role, patient_id, force_password_change) VALUES (?, ?, ?, ?, 1)',
                (username, hashed_pw, 'patient', patient_id)
            )
            db.commit()
            flash('User access granted. Patient will be required to change their password on first login.')
            action = 'created'
        except sqlite3.IntegrityError:
            flash('Username already taken.')

    if action and patient and patient['email']:
        try:
            subject = 'Access to Patient Portal'
            body = (
                f"Hello {patient['name']},\n\n"
                f"Your access to the patient portal has been {'set up' if action == 'created' else 'updated'}.\n\n"
                f"Username: {username}\n"
                f"Temporary password: {password}\n\n"
                f"You will be asked to change your password when you first sign in.\n\n"
                f"Please keep your credentials secure and do not share them with anyone."
            )
            _send_smtp_email(patient['email'], subject, body)
        except Exception:
            pass  # Email failure is non-fatal

    return redirect_to_patient_tab(patient_id, 'info')


@app.route('/patient/<int:patient_id>/reset_portal_password', methods=('POST',))
@login_required
def reset_portal_password(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    patient = db.execute('SELECT * FROM patients WHERE id = ?', (patient_id,)).fetchone()
    if patient is None:
        flash('Patient not found.')
        return redirect_to_patient_tab(patient_id, 'info')

    user = db.execute('SELECT * FROM users WHERE patient_id = ?', (patient_id,)).fetchone()
    if not user:
        flash('No portal account exists for this patient. Use "Grant Access" to create one first.')
        return redirect_to_patient_tab(patient_id, 'info')

    import secrets as _secrets
    new_password = _secrets.token_urlsafe(12)
    hashed_pw = generate_password_hash(new_password)

    db.execute(
        'UPDATE users SET password_hash = ?, force_password_change = 1, is_active = 1 WHERE id = ?',
        (hashed_pw, user['id'])
    )
    db.commit()
    flash('Portal password has been reset. Patient will be required to change their password on next login.')

    if patient['email']:
        try:
            subject = 'Patient Portal Password Reset'
            body = (
                f"Hello {patient['name']},\n\n"
                f"Your patient portal password has been reset by the clinic admin.\n\n"
                f"Username: {user['username']}\n"
                f"Temporary password: {new_password}\n\n"
                f"You will be required to choose a new password when you sign in.\n\n"
                f"Please keep your credentials secure and do not share them with anyone."
            )
            _send_smtp_email(patient['email'], subject, body)
        except Exception:
            pass  # Email failure is non-fatal

    return redirect_to_patient_tab(patient_id, 'info')


@app.route('/patient/<int:patient_id>/enable_intake_tab', methods=('POST',))
@login_required
def enable_intake_tab(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    db.execute('UPDATE patients SET has_intake_tab = 1 WHERE id = ?', (patient_id,))
    db.commit()
    flash('Intake tab enabled for this patient.')
    return redirect(url_for('patient_detail', patient_id=patient_id, tab='intake'))


@app.route('/patient/<int:patient_id>/enable_questionnaire_tab', methods=('POST',))
@login_required
def enable_questionnaire_tab(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    db.execute('UPDATE patients SET has_questionnaire_tab = 1 WHERE id = ?', (patient_id,))
    db.commit()
    flash('Questionnaire tab enabled for this patient.')
    return redirect(url_for('patient_detail', patient_id=patient_id, tab='questionnaires'))


@app.route('/patient/<int:patient_id>/save_questionnaires', methods=('POST',))
@login_required
def save_patient_questionnaires(patient_id):
    if current_user.role != 'admin':
        return 'Unauthorized', 403

    db = get_db()
    patient = db.execute('SELECT * FROM patients WHERE id = ?', (patient_id,)).fetchone()
    if not patient:
        return 'Patient not found', 404

    patient_type = (patient['patient_type'] or '').strip()
    has_questionnaire_tab = int(patient['has_questionnaire_tab'] or 0) == 1
    if patient_type != 'diagnosee' and not has_questionnaire_tab:
        flash('Questionnaire tab is not enabled for this patient.')
        return redirect(url_for('patient_detail', patient_id=patient_id, tab='info'))

    selected_titles = [
        item.strip() for item in request.form.getlist('questionnaire_titles') if item and item.strip()
    ]
    if not selected_titles:
        flash('Select at least one questionnaire.')
        return redirect(url_for('patient_detail', patient_id=patient_id, tab='questionnaires'))

    linked_sheet_id = _extract_google_sheet_id(patient['questionnaires_file_id'] or patient['questionnaires_file_url'])

    if linked_sheet_id:
        result, sync_err = _copy_questionnaire_tabs_to_spreadsheet(db, linked_sheet_id, selected_titles)
        if sync_err:
            flash(f'Failed to update questionnaires file: {sync_err}')
            return redirect(url_for('patient_detail', patient_id=patient_id, tab='questionnaires'))

        db.execute(
            'UPDATE patients SET questionnaires_selected = ? WHERE id = ?',
            (json.dumps(selected_titles, ensure_ascii=False), patient_id)
        )
        db.commit()

        copied_count = len(result.get('copied_titles') or [])
        skipped_count = len(result.get('skipped_existing_titles') or [])
        missing_titles = result.get('missing_titles') or []
        msg = f'Questionnaires updated. Added {copied_count} new tab(s).'
        if skipped_count:
            msg += f' {skipped_count} already existed.'
        if missing_titles:
            msg += ' Missing in source: ' + ', '.join(missing_titles)
        flash(msg)
        return redirect(url_for('patient_detail', patient_id=patient_id, tab='questionnaires'))

    create_result, create_err = _create_diagnosee_questionnaires_sheet(db, patient['name'], selected_titles)
    if create_err:
        flash(f'Failed to create questionnaires file: {create_err}')
        return redirect(url_for('patient_detail', patient_id=patient_id, tab='questionnaires'))

    db.execute(
        '''
        UPDATE patients
        SET questionnaires_file_id = ?, questionnaires_file_url = ?, questionnaires_selected = ?
        WHERE id = ?
        ''',
        (
            create_result['spreadsheet_id'],
            create_result['spreadsheet_url'],
            json.dumps(create_result['selected_titles'], ensure_ascii=False),
            patient_id,
        )
    )
    db.commit()
    flash('Questionnaires file created and linked successfully.')
    return redirect(url_for('patient_detail', patient_id=patient_id, tab='questionnaires'))

@app.route('/patient/<int:patient_id>/toggle_access', methods=('POST',))
@login_required
def toggle_access(patient_id):
    if current_user.role != 'admin':
        return "Unauthorized", 403

    db = get_db()
    user = db.execute('SELECT * FROM users WHERE patient_id = ?', (patient_id,)).fetchone()
    if user:
        new_status = not user['is_active']
        db.execute('UPDATE users SET is_active = ? WHERE id = ?', (new_status, user['id']))
        db.commit()
        flash(f"Access {'enabled' if new_status else 'disabled'}.")
    else:
        flash('No user account found for this patient.')

    return redirect_to_patient_tab(patient_id, 'info')

def _validate_appointment_datetime(date_str, time_str):
    if not date_str or not time_str:
        return None, 'Date and time are required!'

    try:
        datetime.fromisoformat(date_str)
    except ValueError:
        return None, 'Invalid date format!'

    try:
        time_obj = datetime.strptime(time_str, '%H:%M')
        formatted_time = time_obj.strftime('%H:%M')
        return formatted_time, None
    except ValueError:
        return None, 'Invalid time format! Expected HH:MM'

def _extract_recurrence_data(form):
    recurrence_interval = int(form.get('interval', 1))
    recurrence_days = None
    recurrence_end_date = None
    recurrence_count = None

    limit_type = form.get('recurrence_limit_type')
    if limit_type == 'date':
        recurrence_end_date = form.get('recurrence_end_date', '').strip()
        if recurrence_end_date:
            try:
                datetime.fromisoformat(recurrence_end_date)
            except ValueError:
                return None, None, None, None, 'Invalid recurrence end date!'
    elif limit_type == 'count':
        try:
            recurrence_count = int(form.get('recurrence_count', 12))
            if recurrence_count <= 0:
                recurrence_count = 12
        except ValueError:
            recurrence_count = 12

    days_list = form.getlist('days')
    if days_list:
        recurrence_days = ','.join(str(d) for d in days_list if d.strip().isdigit())

    return recurrence_interval, recurrence_days, recurrence_end_date, recurrence_count, None




def is_port_in_use(port, host='127.0.0.1'):
    """Return True when a TCP port is already occupied."""
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.settimeout(0.2)
        return s.connect_ex((host, port)) == 0


def find_available_port(start_port=5000, max_tries=100):
    """Find an open TCP port, starting from start_port."""
    for port in range(start_port, start_port + max_tries):
        if not is_port_in_use(port):
            return port
    raise RuntimeError(f"No available port found in range {start_port}-{start_port + max_tries - 1}")

if __name__ == '__main__':
    ensure_runtime_paths()
    init_db()
    if not os.environ.get('GOOGLE_CLIENT_ID') or not os.environ.get('GOOGLE_CLIENT_SECRET'):
        print('[WARNING] GOOGLE_CLIENT_ID / GOOGLE_CLIENT_SECRET not set — Google integration disabled.')
    try:
        requested_port = int(os.environ.get('PORT', '5000'))
    except ValueError:
        requested_port = 5000

    port = requested_port
    if is_port_in_use(port):
        port = find_available_port(start_port=port + 1)
        print(f"[WARNING] Port {requested_port} is in use. Falling back to {port}.")

    print(f"[INFO] Starting server on port {port}")
    app.run(host='0.0.0.0', port=port, debug=False)
